import sys
import os
import json
import re
import numpy as np
import pandas as pd
from datetime import datetime

# macOS compatibility fixes
if sys.platform == 'darwin':
    # Set environment variables for Qt on macOS
    os.environ['QT_MAC_WANTS_LAYER'] = '1'
    # Fix for macOS Big Sur and later Qt issues
    os.environ['QT_AUTO_SCREEN_SCALE_FACTOR'] = '1'

# CRITICAL: Set Qt API before any Qt or matplotlib imports to avoid conflicts
os.environ['QT_API'] = 'PySide6'

# Set matplotlib to use Qt backend for PySide6 compatibility
import matplotlib
matplotlib.use('QtAgg')  # QtAgg backend works with PySide6
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
from scipy.optimize import curve_fit
from scipy import integrate, stats
from PySide6.QtWidgets import *
from PySide6.QtCore import *
from PySide6.QtGui import *
from matplotlib_config import apply_theme
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import shutil

# Import FP method module
try:
    from xrf_fp_method import XRFFundamentalParameters, HAS_XRAYLIB
except ImportError:
    HAS_XRAYLIB = False
    XRFFundamentalParameters = None

# Helper function for zero-intercept linear regression
def zero_intercept_regression(x, y):
    """
    Perform linear regression forcing intercept through origin (y = mx)
    
    Parameters:
    x, y: array-like, data points
    
    Returns:
    slope: float, slope of the line
    r_squared: float, coefficient of determination
    std_error: float, standard error of the slope
    """
    x = np.array(x)
    y = np.array(y)
    
    # Calculate slope: slope = sum(x*y) / sum(x^2)
    slope = np.sum(x * y) / np.sum(x * x)
    
    # Calculate predicted values
    y_pred = slope * x
    
    # Calculate R-squared
    ss_res = np.sum((y - y_pred) ** 2)  # Sum of squares of residuals
    ss_tot = np.sum((y - np.mean(y)) ** 2)  # Total sum of squares
    r_squared = 1 - (ss_res / ss_tot) if ss_tot != 0 else 0
    
    # Calculate standard error of slope
    n = len(x)
    if n > 1:
        mse = ss_res / (n - 1)  # Mean squared error (adjusted for 1 parameter)
        std_error = np.sqrt(mse / np.sum(x * x))
    else:
        std_error = 0
    
    return slope, r_squared, std_error

class CalibrationManager:
    """Manages persistent storage and retrieval of element calibrations"""
    
    def __init__(self, calibration_file="xrf_calibrations.json"):
        self.calibration_file = calibration_file
        self.calibrations = self.load_calibrations()
    
    def load_calibrations(self):
        """Load calibrations from file"""
        if os.path.exists(self.calibration_file):
            try:
                with open(self.calibration_file, 'r') as f:
                    data = json.load(f)
                print(f"Loaded calibrations from {self.calibration_file}")
                return data
            except Exception as e:
                print(f"Error loading calibrations: {e}")
                return {}
        return {}
    
    def save_calibrations(self):
        """Save calibrations to file"""
        try:
            with open(self.calibration_file, 'w') as f:
                json.dump(self.calibrations, f, indent=2)
            print(f"Saved calibrations to {self.calibration_file}")
        except Exception as e:
            print(f"Error saving calibrations: {e}")
    
    def update_calibration(self, element, slope, intercept, r_squared=None, standards_used=None, raw_intensities=None, raw_standards=None):
        """Update calibration for an element"""
        if element not in self.calibrations:
            self.calibrations[element] = {}
        
        self.calibrations[element].update({
            'slope': float(slope),
            'intercept': float(intercept),
            'r_squared': float(r_squared) if r_squared is not None else None,
            'standards_used': standards_used if standards_used else [],
            'created_date': datetime.now().isoformat(),
            'equation': f"Concentration = {slope:.4f} × Intensity + {intercept:.4f}",
            'raw_intensities': raw_intensities if raw_intensities else {},  # Individual measurements per standard
            'raw_standards': raw_standards if raw_standards else []  # List of standards for each measurement
        })
        
        self.save_calibrations()
        print(f"Updated calibration for {element}")
    
    def get_calibration(self, element):
        """Get calibration for an element"""
        return self.calibrations.get(element, None)
    
    def has_calibration(self, element):
        """Check if element has a calibration"""
        return element in self.calibrations
    
    def get_all_calibrations(self):
        """Get all calibrations"""
        return self.calibrations.copy()
    
    def delete_calibration(self, element):
        """Delete calibration for an element"""
        if element in self.calibrations:
            del self.calibrations[element]
            self.save_calibrations()
            print(f"Deleted calibration for {element}")
    
    def export_calibrations(self, filename):
        """Export calibrations to a file"""
        try:
            with open(filename, 'w') as f:
                json.dump(self.calibrations, f, indent=2)
            return True
        except Exception as e:
            print(f"Error exporting calibrations: {e}")
            return False
    
    def import_calibrations(self, filename):
        """Import calibrations from a file"""
        try:
            with open(filename, 'r') as f:
                imported = json.load(f)
            self.calibrations.update(imported)
            self.save_calibrations()
            return True
        except Exception as e:
            print(f"Error importing calibrations: {e}")
            return False

# Element definitions for XRF analysis
ELEMENT_DEFINITIONS = {
    'Pb': {
        'name': 'Lead',
        'symbol': 'Pb',
        'primary_energy': 10.55,  # Pb L-alpha
        'secondary_energy': 12.61,  # Pb L-beta
        'peak_region': (10.0, 11.0),
        'integration_region': (9.8, 11.2),
        'default_calibration': {'slope': 13.8913, 'intercept': 0.0},
        # Alternative peak to avoid As interference
        'alternative_peak': {
            'name': 'L-beta',
            'energy': 12.61,
            'peak_region': (12.0, 13.2),
            'integration_region': (11.8, 13.4),
            'note': 'Use to avoid As K-alpha interference (lower intensity)'
        }
    },
    'As': {
        'name': 'Arsenic',
        'symbol': 'As',
        'primary_energy': 10.54,  # As K-alpha
        'secondary_energy': 11.73,  # As K-beta
        'peak_region': (10.0, 11.0),
        'integration_region': (9.8, 11.2),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0},
        # Alternative peak to avoid Pb interference
        'alternative_peak': {
            'name': 'K-beta',
            'energy': 11.73,
            'peak_region': (11.2, 12.3),
            'integration_region': (11.0, 12.5),
            'note': 'Use to avoid Pb L-alpha interference (lower intensity)'
        }
    },
    'Cd': {
        'name': 'Cadmium',
        'symbol': 'Cd',
        'primary_energy': 23.17,  # Cd K-alpha
        'secondary_energy': 26.10,  # Cd K-beta
        'peak_region': (22.5, 24.0),
        'integration_region': (22.0, 24.5),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0}
    },
    'Cr': {
        'name': 'Chromium',
        'symbol': 'Cr',
        'primary_energy': 5.41,  # Cr K-alpha
        'secondary_energy': 5.95,  # Cr K-beta
        'peak_region': (5.0, 6.0),
        'integration_region': (4.8, 6.2),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0}
    },
    'Zn': {
        'name': 'Zinc',
        'symbol': 'Zn',
        'primary_energy': 8.64,  # Zn K-alpha
        'secondary_energy': 9.57,  # Zn K-beta
        'peak_region': (8.2, 9.2),
        'integration_region': (8.0, 9.4),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0}
    },
    'Ni': {
        'name': 'Nickel',
        'symbol': 'Ni',
        'primary_energy': 7.48,  # Ni K-alpha
        'secondary_energy': 8.26,  # Ni K-beta
        'peak_region': (7.0, 8.0),
        'integration_region': (6.8, 8.2),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0}
    },
    'Cu': {
        'name': 'Copper',
        'symbol': 'Cu',
        'primary_energy': 8.05,  # Cu K-alpha
        'secondary_energy': 8.91,  # Cu K-beta
        'peak_region': (7.6, 8.6),
        'integration_region': (7.4, 8.8),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0}
    },
    'Fe': {
        'name': 'Iron',
        'symbol': 'Fe',
        'primary_energy': 6.40,  # Fe K-alpha
        'secondary_energy': 7.06,  # Fe K-beta
        'peak_region': (6.0, 7.0),
        'integration_region': (5.8, 7.2),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0}
    },
    'Se': {
        'name': 'Selenium',
        'symbol': 'Se',
        'primary_energy': 11.22,  # Se K-alpha
        'secondary_energy': 12.50,  # Se K-beta
        'peak_region': (10.8, 11.8),
        'integration_region': (10.6, 12.0),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0}
    },
    'S': {
        'name': 'Sulfur',
        'symbol': 'S',
        'primary_energy': 2.31,  # S K-alpha
        'secondary_energy': 2.46,  # S K-beta
        'peak_region': (2.0, 2.8),
        'integration_region': (1.8, 3.0),
        'default_calibration': {'slope': 1.0, 'intercept': 0.0}
    }
}

# Peak interference definitions
# Format: {element: [list of elements that interfere with it]}
PEAK_INTERFERENCES = {
    'Pb': ['As'],  # Pb L-alpha (10.55) overlaps with As K-alpha (10.54)
    'As': ['Pb'],  # As K-alpha (10.54) overlaps with Pb L-alpha (10.55)
    # Add more as needed:
    # 'Ni': ['Cu'],  # If Ni K-beta overlaps with Cu K-alpha
    # 'Cu': ['Zn'],  # If Cu K-beta overlaps with Zn K-alpha
}

# Interference correction notes
INTERFERENCE_NOTES = {
    ('Pb', 'As'): "Pb L-alpha (10.55 keV) and As K-alpha (10.54 keV) are indistinguishable on most portable XRF. Use deconvolution or alternative peaks.",
    ('As', 'Pb'): "As K-alpha (10.54 keV) and Pb L-alpha (10.55 keV) are indistinguishable on most portable XRF. Use deconvolution or alternative peaks.",
}

# Reference material data from the user's table
REFERENCE_MATERIALS = {
    'Till 1': {
        'S': '<0.05%',  # Special case - below detection limit
        'As': 18,
        'Cd': None,
        'Cr': 65,
        'Zn': 98,
        'Ni': 24,
        'Cu': 47,
        'Pb': 22,
        'Fe': '4.81%',
        'Se': None
    },
    'LKSD 1': {
        'S': '1.57%',
        'As': 40,
        'Cd': None,
        'Cr': 31,
        'Zn': 331,
        'Ni': 16,
        'Cu': 44,
        'Pb': 82,
        'Fe': '2.8%',
        'Se': None
    },
    'PACS 2': {
        'S': '1',  # Assuming this is 1%
        'As': 26.2,
        'Cd': 2.11,
        'Cr': 90.7,
        'Zn': 364,
        'Ni': 39.5,
        'Cu': 310,
        'Pb': 183,
        'Fe': '4.09',  # Assuming this is 4.09%
        'Se': 0.92
    },
    'STDS 2': {
        'S': '0.06%',
        'As': 42,
        'Cd': None,
        'Cr': 116,
        'Zn': 246,
        'Ni': 53,
        'Cu': 47,
        'Pb': 66,
        'Fe': '5.2%',
        'Se': None
    },
    'NIST 2586': {
        'S': '600',  # Assuming ppm
        'As': 8.7,
        'Cd': 2.71,
        'Cr': 301,
        'Zn': 352,
        'Ni': None,
        'Cu': None,
        'Pb': 432,
        'Fe': '51610',  # Assuming ppm
        'Se': None
    },
    'NIST 2587': {
        'S': None,
        'As': 13.7,
        'Cd': 1.92,
        'Cr': 92,
        'Zn': 335.8,
        'Ni': None,
        'Cu': None,
        'Pb': 3242,
        'Fe': '28130',  # Assuming ppm
        'Se': None
    },
    'NBS 1633': {
        'S': None,
        'As': None,
        'Cd': None,
        'Cr': None,
        'Zn': None,
        'Ni': None,
        'Cu': None,
        'Pb': None,
        'Fe': None,
        'Se': None
    }
}

# Load XRF lines database for element identification
def load_xrf_lines_database():
    """Load XRF characteristic lines from CSV file"""
    csv_path = os.path.join(os.path.dirname(__file__), 'data', 'xrf_lines_Na_to_U.csv')
    try:
        df = pd.read_csv(csv_path)
        # Group by element for easy lookup
        xrf_lines_db = {}
        for element in df['Element'].unique():
            element_lines = df[df['Element'] == element]
            xrf_lines_db[element] = element_lines.to_dict('records')
        return xrf_lines_db, df
    except Exception as e:
        print(f"Warning: Could not load XRF lines database: {e}")
        return {}, pd.DataFrame()

XRF_LINES_DB, XRF_LINES_DF = load_xrf_lines_database()

# Import matplotlib configuration
try:
    from matplotlib_config import (
        configure_compact_ui, 
        CompactNavigationToolbar, 
        MiniNavigationToolbar,
        apply_theme,
        get_toolbar_class
    )
    # Apply compact configuration
    configure_compact_ui()
except ImportError:
    print("Warning: matplotlib_config.py not found, using default matplotlib settings")
    CompactNavigationToolbar = None
    MiniNavigationToolbar = None


class FileSortingDialog(QDialog):
    """Dialog for previewing and selecting file sorting options"""
    
    def __init__(self, file_paths, folder_path, parent=None):
        super().__init__(parent)
        self.file_paths = file_paths
        self.folder_path = folder_path
        self.sorted_files = []
        
        self.setWindowTitle("File Sorting Options")
        self.setGeometry(300, 200, 800, 600)
        self.setModal(True)
        
        self.init_ui()
        self.load_files()
    
    def init_ui(self):
        """Initialize the user interface"""
        layout = QVBoxLayout(self)
        
        # Header
        header_label = QLabel("Preview and select file sorting method:")
        header_label.setFont(QFont("Arial", 12, QFont.Weight.Bold))
        layout.addWidget(header_label)
        
        # Extension filtering
        filter_group = QGroupBox("File Type Filter")
        filter_layout = QVBoxLayout(filter_group)
        
        # Extension checkboxes
        self.extension_checkboxes = {}
        common_extensions = ['.txt', '.csv', '.xlsx', '.dat', '.emsa', '.spc']
        
        # Get all unique extensions from the files
        all_extensions = set()
        for file_path in self.file_paths:
            ext = os.path.splitext(file_path)[1].lower()
            if ext:
                all_extensions.add(ext)
        
        # Sort extensions for consistent display
        all_extensions = sorted(all_extensions)
        
        # Create checkboxes for each extension
        ext_layout = QHBoxLayout()
        for ext in all_extensions:
            checkbox = QCheckBox(ext)
            checkbox.setChecked(True)  # Default to checked
            checkbox.stateChanged.connect(self.update_preview)
            self.extension_checkboxes[ext] = checkbox
            ext_layout.addWidget(checkbox)
        
        filter_layout.addLayout(ext_layout)
        
        # Quick selection buttons
        quick_buttons = QHBoxLayout()
        select_all_btn = QPushButton("Select All")
        select_all_btn.clicked.connect(self.select_all_extensions)
        quick_buttons.addWidget(select_all_btn)
        
        deselect_all_btn = QPushButton("Deselect All")
        deselect_all_btn.clicked.connect(self.deselect_all_extensions)
        quick_buttons.addWidget(deselect_all_btn)
        
        # Add common XRF extensions button
        xrf_only_btn = QPushButton("XRF Files Only")
        xrf_only_btn.clicked.connect(self.select_xrf_extensions)
        quick_buttons.addWidget(xrf_only_btn)
        
        quick_buttons.addStretch()
        filter_layout.addLayout(quick_buttons)
        
        layout.addWidget(filter_group)
        
        # Sorting options
        options_group = QGroupBox("Sorting Options")
        options_layout = QHBoxLayout(options_group)
        
        self.sort_method_combo = QComboBox()
        self.sort_method_combo.addItems([
            "Smart Sort (Recommended)",
            "Alphabetical Sort",
            "Date Modified",
            "File Size",
            "Custom Order"
        ])
        self.sort_method_combo.currentTextChanged.connect(self.update_preview)
        options_layout.addWidget(QLabel("Sort Method:"))
        options_layout.addWidget(self.sort_method_combo)
        
        # Pattern detection
        self.pattern_label = QLabel("Detected Pattern: Unknown")
        self.pattern_label.setStyleSheet("color: #666; font-style: italic;")
        options_layout.addWidget(self.pattern_label)
        
        layout.addWidget(options_group)
        
        # Preview area
        preview_group = QGroupBox("File Order Preview")
        preview_layout = QVBoxLayout(preview_group)
        
        # File list
        self.file_list = QTableWidget()
        self.file_list.setColumnCount(3)
        self.file_list.setHorizontalHeaderLabels(["Order", "Filename", "Full Path"])
        self.file_list.horizontalHeader().setStretchLastSection(True)
        self.file_list.setAlternatingRowColors(True)
        preview_layout.addWidget(self.file_list)
        
        # Preview info
        self.preview_info = QLabel("Ready to preview files...")
        preview_layout.addWidget(self.preview_info)
        
        layout.addWidget(preview_group)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        self.refresh_btn = QPushButton("🔄 Refresh Preview")
        self.refresh_btn.clicked.connect(self.update_preview)
        button_layout.addWidget(self.refresh_btn)
        
        button_layout.addStretch()
        
        self.cancel_btn = QPushButton("Cancel")
        self.cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(self.cancel_btn)
        
        self.accept_btn = QPushButton("Use This Order")
        self.accept_btn.clicked.connect(self.accept)
        button_layout.addWidget(self.accept_btn)
        
        layout.addLayout(button_layout)
    
    def load_files(self):
        """Load and display files"""
        self.update_preview()
    
    def update_preview(self):
        """Update the file preview based on selected sorting method and extension filter"""
        # First filter by selected extensions
        selected_extensions = [ext for ext, checkbox in self.extension_checkboxes.items() 
                             if checkbox.isChecked()]
        
        if not selected_extensions:
            # If no extensions selected, show no files
            self.sorted_files = []
            self.pattern_label.setText("No file types selected")
            self.display_files()
            return
        
        # Filter files by selected extensions
        filtered_files = [f for f in self.file_paths 
                         if os.path.splitext(f)[1].lower() in selected_extensions]
        
        if not filtered_files:
            self.sorted_files = []
            self.pattern_label.setText("No files match selected extensions")
            self.display_files()
            return
        
        # Apply sorting method
        method = self.sort_method_combo.currentText()
        
        if method == "Smart Sort (Recommended)":
            self.sorted_files = self.smart_sort_files(filtered_files)
            pattern = self.detect_sorting_pattern(filtered_files)
            self.pattern_label.setText(f"Detected Pattern: {pattern}")
        elif method == "Alphabetical Sort":
            self.sorted_files = sorted(filtered_files)
            self.pattern_label.setText("Detected Pattern: Alphabetical")
        elif method == "Date Modified":
            self.sorted_files = sorted(filtered_files, key=lambda x: os.path.getmtime(x))
            self.pattern_label.setText("Detected Pattern: Date Modified")
        elif method == "File Size":
            self.sorted_files = sorted(filtered_files, key=lambda x: os.path.getsize(x))
            self.pattern_label.setText("Detected Pattern: File Size")
        elif method == "Custom Order":
            # For custom order, we'll keep the original order but allow manual reordering
            self.sorted_files = filtered_files.copy()
            self.pattern_label.setText("Detected Pattern: Custom")
        
        self.display_files()
    
    def smart_sort_files(self, file_paths):
        """Sort files using natural sorting (handles numbers correctly)"""
        def natural_sort_key(filename):
            basename = os.path.basename(filename)
            def convert(text):
                return int(text) if text.isdigit() else text.lower()
            return [convert(c) for c in re.split('([0-9]+)', basename)]
        
        return sorted(file_paths, key=natural_sort_key)
    
    def detect_sorting_pattern(self, file_paths):
        """Detect the naming pattern in the files"""
        if not file_paths:
            return "unknown"
        
        filenames = [os.path.basename(f) for f in file_paths]
        
        patterns = {
            "sample_number": r"sample_(\d+)",
            "sample_number_extended": r"sample_(\d+)_",
            "number_only": r"^(\d+)",
            "letter_number": r"([A-Za-z]+)(\d+)",
            "date_pattern": r"(\d{4})[-_](\d{2})[-_](\d{2})",
            "time_pattern": r"(\d{2})[-:](\d{2})[-:](\d{2})"
        }
        
        for pattern_name, pattern in patterns.items():
            matches = [re.search(pattern, f) for f in filenames]
            if all(matches):
                return pattern_name
        
        return "unknown"
    
    def display_files(self):
        """Display files in the table"""
        self.file_list.setRowCount(len(self.sorted_files))
        
        for i, file_path in enumerate(self.sorted_files):
            # Order number
            order_item = QTableWidgetItem(str(i + 1))
            order_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.file_list.setItem(i, 0, order_item)
            
            # Filename
            filename_item = QTableWidgetItem(os.path.basename(file_path))
            self.file_list.setItem(i, 1, filename_item)
            
            # Full path
            path_item = QTableWidgetItem(file_path)
            path_item.setToolTip(file_path)
            self.file_list.setItem(i, 2, path_item)
        
        # Update info
        self.preview_info.setText(f"Showing {len(self.sorted_files)} files in {self.sort_method_combo.currentText()} order")
        
        # Resize columns
        self.file_list.resizeColumnsToContents()
    
    def get_sorted_files(self):
        """Return the sorted file list"""
        return self.sorted_files
    
    def select_all_extensions(self):
        """Select all file extensions"""
        for checkbox in self.extension_checkboxes.values():
            checkbox.setChecked(True)
    
    def deselect_all_extensions(self):
        """Deselect all file extensions"""
        for checkbox in self.extension_checkboxes.values():
            checkbox.setChecked(False)
    
    def select_xrf_extensions(self):
        """Select only common XRF file extensions"""
        xrf_extensions = ['.txt', '.csv', '.dat', '.emsa', '.spc']
        for ext, checkbox in self.extension_checkboxes.items():
            checkbox.setChecked(ext in xrf_extensions)


class XRFPeakFitter:
    """Core class for XRF peak fitting with background subtraction and Gaussian-A fitting"""
    
    def __init__(self, element='Pb'):
        self.current_element = element
        self.element_data = ELEMENT_DEFINITIONS.get(element, ELEMENT_DEFINITIONS['Pb'])
        
        self.target_energy = self.element_data['primary_energy']
        # Calibration parameters
        self.calibration_slope = self.element_data['default_calibration']['slope']
        self.calibration_intercept = self.element_data['default_calibration']['intercept']
        # Custom calibration name
        self.calibration_name = f"{element} Default Calibration"
        
        # Element-specific calibrations storage
        self.element_calibrations = {element: {'slope': self.calibration_slope, 'intercept': self.calibration_intercept}}
        
        # Track which peak to use (primary or alternative)
        self.use_alternative_peak = {}
    
    def set_element(self, element):
        """Switch to a different element"""
        if element in ELEMENT_DEFINITIONS:
            self.current_element = element
            self.element_data = ELEMENT_DEFINITIONS[element]
            self.target_energy = self.element_data['primary_energy']
            
            # Load calibration for this element
            if element in self.element_calibrations:
                self.calibration_slope = self.element_calibrations[element]['slope']
                self.calibration_intercept = self.element_calibrations[element]['intercept']
            else:
                self.calibration_slope = self.element_data['default_calibration']['slope']
                self.calibration_intercept = self.element_data['default_calibration']['intercept']
                self.element_calibrations[element] = {'slope': self.calibration_slope, 'intercept': self.calibration_intercept}
            
            self.calibration_name = f"{element} Calibration"
    
    def update_element_calibration(self, element, slope, intercept):
        """Update calibration parameters for a specific element"""
        self.element_calibrations[element] = {'slope': slope, 'intercept': intercept}
        if element == self.current_element:
            self.calibration_slope = slope
            self.calibration_intercept = intercept
    
    def set_use_alternative_peak(self, element, use_alternative):
        """Set whether to use alternative peak for an element"""
        self.use_alternative_peak[element] = use_alternative
    
    def get_peak_regions(self, element=None):
        """Get the appropriate peak regions based on peak selection"""
        if element is None:
            element = self.current_element
        
        if element not in ELEMENT_DEFINITIONS:
            return None, None
        
        element_data = ELEMENT_DEFINITIONS[element]
        use_alt = self.use_alternative_peak.get(element, False)
        
        if use_alt and 'alternative_peak' in element_data:
            alt_peak = element_data['alternative_peak']
            return alt_peak['peak_region'], alt_peak['integration_region']
        else:
            return element_data['peak_region'], element_data['integration_region']
        
    def gaussian_a(self, x, a, x0, dx):
        """
        Gaussian-A function with the specified equation:
        sqrt(ln(2)/pi) * (a/dx) * exp(-ln(2) * (x-x0)^2 / dx^2)
        
        Parameters:
        - a: amplitude parameter
        - x0: peak center
        - dx: full width at half maximum (FWHM)
        """
        ln2 = np.log(2)
        coefficient = np.sqrt(ln2 / np.pi)
        amplitude = a / dx
        exponent = -ln2 * ((x - x0) ** 2) / (dx ** 2)
        return coefficient * amplitude * np.exp(exponent)
    
    def linear_background(self, x, m, b):
        """Linear background function"""
        return m * x + b
    
    def combined_model(self, x, a, x0, dx, m, b):
        """Combined peak + background model"""
        return self.gaussian_a(x, a, x0, dx) + self.linear_background(x, m, b)
    
    def estimate_background(self, x, y, peak_region):
        """Estimate linear background excluding peak region"""
        mask = ~((x >= peak_region[0]) & (x <= peak_region[1]))
        x_bg = x[mask]
        y_bg = y[mask]
        
        if len(x_bg) < 2:
            return 0, np.mean(y)
        
        # Linear fit to background points
        p = np.polyfit(x_bg, y_bg, 1)
        return p[0], p[1]
    
    def calculate_integrated_intensity(self, x, y, fit_params, peak_region):
        """
        Calculate integrated intensity of the peak after background subtraction
        
        Parameters:
        - x: energy array
        - y: intensity array
        - fit_params: fitted parameters
        - peak_region: integration region
        
        Returns:
        - integrated_intensity: background-corrected integrated peak area
        """
        # Select integration region
        mask = (x >= peak_region[0]) & (x <= peak_region[1])
        x_int = x[mask]
        y_int = y[mask]
        
        # Calculate background in integration region
        background = self.linear_background(x_int, 
                                          fit_params['background_slope'], 
                                          fit_params['background_intercept'])
        
        # Background-corrected intensity
        y_corrected = y_int - background
        
        # Integrate using trapezoidal rule
        integrated_intensity = integrate.trapz(y_corrected, x_int)
        
        return integrated_intensity
    
    def apply_calibration(self, integrated_intensity):
        """
        Apply NIST calibration curve to convert integrated intensity to concentration
        
        Calibration equation: Concentration = 13.8913 * AVE_I + 0
        
        Parameters:
        - integrated_intensity: background-corrected integrated peak area
        
        Returns:
        - concentration: calibrated concentration value
        """
        concentration = self.calibration_slope * integrated_intensity + self.calibration_intercept
        return concentration
    
    def fit_peak(self, x, y, peak_region=None, background_subtract=True, integration_region=None):
        """
        Fit Gaussian-A peak with optional background subtraction and calculate integrated intensity
        
        Parameters:
        - x: energy array
        - y: intensity array
        - peak_region: tuple (min_energy, max_energy) for fitting region
        - background_subtract: whether to include background in fit
        - integration_region: tuple (min_energy, max_energy) for integration
        
        Returns:
        - fit_params: dictionary with fit parameters
        - fit_curve: fitted curve
        - r_squared: coefficient of determination
        - integrated_intensity: background-corrected integrated peak area
        - concentration: calibrated concentration
        """
        if peak_region is None:
            # Use appropriate peak region based on peak selection
            peak_region, integration_region_default = self.get_peak_regions()
            if integration_region is None:
                integration_region = integration_region_default
        elif integration_region is None:
            # Peak region provided but not integration region
            _, integration_region = self.get_peak_regions()
        
        # Select fitting region
        mask = (x >= peak_region[0]) & (x <= peak_region[1])
        x_fit = x[mask]
        y_fit = y[mask]
        
        if len(x_fit) < 5:
            raise ValueError("Insufficient data points in fitting region")
        
        # Initial parameter estimation
        peak_idx = np.argmax(y_fit)
        x0_init = x_fit[peak_idx]
        a_init = np.max(y_fit)
        dx_init = 0.1  # Initial FWHM estimate
        
        try:
            if background_subtract:
                # Estimate background
                m_init, b_init = self.estimate_background(x_fit, y_fit, peak_region)
                
                # Initial guess
                p0 = [a_init, x0_init, dx_init, m_init, b_init]
                
                # Bounds for parameters [a, x0, dx, m, b]
                bounds = ([0, peak_region[0], 0.01, -np.inf, 0],
                         [np.inf, peak_region[1], 1.0, np.inf, np.inf])
                
                # Fit combined model
                popt, pcov = curve_fit(self.combined_model, x_fit, y_fit, p0=p0, bounds=bounds)
                
                # Calculate fitted curve
                fit_curve = self.combined_model(x_fit, *popt)
                
                # Extract parameters
                fit_params = {
                    'amplitude': popt[0],
                    'center': popt[1],
                    'fwhm': popt[2],
                    'background_slope': popt[3],
                    'background_intercept': popt[4],
                    'actual_peak_area': popt[0] * popt[2] * np.sqrt(np.pi / np.log(2)),  # Proper Gaussian-A area
                    'amplitude_error': np.sqrt(pcov[0,0]),
                    'center_error': np.sqrt(pcov[1,1]),
                    'fwhm_error': np.sqrt(pcov[2,2])
                }
                
            else:
                # Subtract estimated background first
                m_bg, b_bg = self.estimate_background(x_fit, y_fit, peak_region)
                y_bg_sub = y_fit - self.linear_background(x_fit, m_bg, b_bg)
                
                # Initial guess for peak only
                p0 = [a_init, x0_init, dx_init]
                
                # Bounds for parameters [a, x0, dx]
                bounds = ([0, peak_region[0], 0.01],
                         [np.inf, peak_region[1], 1.0])
                
                # Fit peak only
                popt, pcov = curve_fit(self.gaussian_a, x_fit, y_bg_sub, p0=p0, bounds=bounds)
                
                # Calculate fitted curve
                fit_curve = self.gaussian_a(x_fit, *popt) + self.linear_background(x_fit, m_bg, b_bg)
                
                # Extract parameters
                fit_params = {
                    'amplitude': popt[0],
                    'center': popt[1],
                    'fwhm': popt[2],
                    'background_slope': m_bg,
                    'background_intercept': b_bg,
                    'actual_peak_area': popt[0] * popt[2] * np.sqrt(np.pi / np.log(2)),  # Proper Gaussian-A area
                    'amplitude_error': np.sqrt(pcov[0,0]),
                    'center_error': np.sqrt(pcov[1,1]),
                    'fwhm_error': np.sqrt(pcov[2,2])
                }
            
            # Calculate R-squared
            ss_res = np.sum((y_fit - fit_curve) ** 2)
            ss_tot = np.sum((y_fit - np.mean(y_fit)) ** 2)
            r_squared = 1 - (ss_res / ss_tot)
            
            # Calculate integrated intensity
            integrated_intensity = self.calculate_integrated_intensity(x, y, fit_params, integration_region)
            
            # Apply calibration
            concentration = self.apply_calibration(integrated_intensity)
            
            return fit_params, fit_curve, r_squared, x_fit, integrated_intensity, concentration
            
        except Exception as e:
            raise RuntimeError(f"Fitting failed: {str(e)}")
    
    def fit_pb_as_deconvolution(self, x, y):
        """
        Simultaneous deconvolution of Pb and As peaks using multiple characteristic lines.
        
        This method addresses the overlap between Pb Lα1 (10.5515 keV) and As Kα1 (10.5437 keV)
        by fitting all available lines:
        - Pb: Lα1 (10.5515), Lα2 (10.4495), Lβ1 (12.6137)
        - As: Kα1 (10.5437), Kα2 (10.5078), Kβ1 (11.7262)
        
        Uses theoretical intensity ratios as constraints to improve deconvolution accuracy.
        
        Returns:
        - pb_results: dict with Pb fit parameters and concentration
        - as_results: dict with As fit parameters and concentration
        - combined_fit: fitted spectrum
        - r_squared: goodness of fit
        """
        
        # Define characteristic line energies and relative intensities from database
        pb_lines = {
            'La1': {'energy': 10.5515, 'rel_intensity': 100},
            'La2': {'energy': 10.4495, 'rel_intensity': 10},
            'Lb1': {'energy': 12.6137, 'rel_intensity': 75}
        }
        
        as_lines = {
            'Ka1': {'energy': 10.5437, 'rel_intensity': 100},
            'Ka2': {'energy': 10.5078, 'rel_intensity': 50},
            'Kb1': {'energy': 11.7262, 'rel_intensity': 62}
        }
        
        # Define fitting region covering all peaks (10-13 keV)
        fit_region = (9.8, 13.2)
        mask = (x >= fit_region[0]) & (x <= fit_region[1])
        x_fit = x[mask]
        y_fit = y[mask]
        
        if len(x_fit) < 20:
            raise ValueError("Insufficient data points for Pb-As deconvolution")
        
        # Estimate background
        bg_mask = ((x_fit < 10.0) | (x_fit > 13.0))
        if np.sum(bg_mask) > 2:
            m_bg, b_bg = np.polyfit(x_fit[bg_mask], y_fit[bg_mask], 1)
        else:
            m_bg = 0
            b_bg = np.min(y_fit)
        
        # Define multi-peak model
        def multi_peak_model(x, pb_amp, as_amp, fwhm_pb, fwhm_as, m, b):
            """
            Model with all Pb and As characteristic lines.
            Amplitudes are constrained by theoretical intensity ratios.
            """
            model = np.zeros_like(x)
            
            # Pb peaks (using Lα1 amplitude as reference)
            for line_name, line_data in pb_lines.items():
                amplitude = pb_amp * (line_data['rel_intensity'] / 100.0)
                model += self.gaussian_a(x, amplitude, line_data['energy'], fwhm_pb)
            
            # As peaks (using Kα1 amplitude as reference)
            for line_name, line_data in as_lines.items():
                amplitude = as_amp * (line_data['rel_intensity'] / 100.0)
                model += self.gaussian_a(x, amplitude, line_data['energy'], fwhm_as)
            
            # Add background
            model += m * x + b
            
            return model
        
        # Initial parameter estimates
        # Find peak around 10.5 keV (overlapped Pb Lα + As Kα)
        overlap_mask = (x_fit >= 10.3) & (x_fit <= 10.7)
        overlap_height = np.max(y_fit[overlap_mask]) if np.sum(overlap_mask) > 0 else 1000
        
        # Find As Kβ peak around 11.7 keV (well separated)
        as_kb_mask = (x_fit >= 11.5) & (x_fit <= 11.9)
        as_kb_height = np.max(y_fit[as_kb_mask]) if np.sum(as_kb_mask) > 0 else 100
        
        # Find Pb Lβ peak around 12.6 keV (well separated)
        pb_lb_mask = (x_fit >= 12.4) & (x_fit <= 12.8)
        pb_lb_height = np.max(y_fit[pb_lb_mask]) if np.sum(pb_lb_mask) > 0 else 100
        
        # Estimate initial amplitudes based on separated peaks
        # As Kβ1 has 62% intensity relative to Kα1
        as_amp_init = as_kb_height / 0.62 if as_kb_height > 0 else overlap_height * 0.5
        
        # Pb Lβ1 has 75% intensity relative to Lα1
        pb_amp_init = pb_lb_height / 0.75 if pb_lb_height > 0 else overlap_height * 0.5
        
        # Initial guess: [pb_amp, as_amp, fwhm_pb, fwhm_as, m, b]
        p0 = [pb_amp_init, as_amp_init, 0.15, 0.15, m_bg, b_bg]
        
        # Bounds: amplitudes > 0, FWHM between 0.05-0.5 keV
        bounds = (
            [0, 0, 0.05, 0.05, -np.inf, 0],
            [np.inf, np.inf, 0.5, 0.5, np.inf, np.inf]
        )
        
        try:
            # Perform fit
            popt, pcov = curve_fit(multi_peak_model, x_fit, y_fit, p0=p0, bounds=bounds, maxfev=10000)
            
            pb_amp, as_amp, fwhm_pb, fwhm_as, m, b = popt
            
            # Calculate fitted curve
            combined_fit = multi_peak_model(x_fit, *popt)
            
            # Calculate R-squared
            ss_res = np.sum((y_fit - combined_fit) ** 2)
            ss_tot = np.sum((y_fit - np.mean(y_fit)) ** 2)
            r_squared = 1 - (ss_res / ss_tot) if ss_tot > 0 else 0
            
            # Calculate integrated intensities for each element
            # Pb: integrate all Pb peaks
            pb_intensity = 0
            for line_name, line_data in pb_lines.items():
                amplitude = pb_amp * (line_data['rel_intensity'] / 100.0)
                # Gaussian-A area = amplitude * fwhm * sqrt(pi/ln(2))
                pb_intensity += amplitude * fwhm_pb * np.sqrt(np.pi / np.log(2))
            
            # As: integrate all As peaks
            as_intensity = 0
            for line_name, line_data in as_lines.items():
                amplitude = as_amp * (line_data['rel_intensity'] / 100.0)
                as_intensity += amplitude * fwhm_as * np.sqrt(np.pi / np.log(2))
            
            # Apply calibrations
            pb_concentration = None
            as_concentration = None
            
            if 'Pb' in self.element_calibrations:
                pb_cal = self.element_calibrations['Pb']
                pb_concentration = pb_cal['slope'] * pb_intensity + pb_cal['intercept']
            
            if 'As' in self.element_calibrations:
                as_cal = self.element_calibrations['As']
                as_concentration = as_cal['slope'] * as_intensity + as_cal['intercept']
            
            # Prepare results
            pb_results = {
                'amplitude': pb_amp,
                'fwhm': fwhm_pb,
                'integrated_intensity': pb_intensity,
                'concentration': pb_concentration,
                'lines_used': ['Lα1', 'Lα2', 'Lβ1'],
                'amplitude_error': np.sqrt(pcov[0, 0]) if pcov[0, 0] > 0 else 0,
                'fwhm_error': np.sqrt(pcov[2, 2]) if pcov[2, 2] > 0 else 0
            }
            
            as_results = {
                'amplitude': as_amp,
                'fwhm': fwhm_as,
                'integrated_intensity': as_intensity,
                'concentration': as_concentration,
                'lines_used': ['Kα1', 'Kα2', 'Kβ1'],
                'amplitude_error': np.sqrt(pcov[1, 1]) if pcov[1, 1] > 0 else 0,
                'fwhm_error': np.sqrt(pcov[3, 3]) if pcov[3, 3] > 0 else 0
            }
            
            return pb_results, as_results, x_fit, combined_fit, r_squared
            
        except Exception as e:
            raise RuntimeError(f"Pb-As deconvolution failed: {str(e)}")

class SampleGroup:
    """Class to handle groups of spectra from the same sample"""
    
    def __init__(self, sample_name, spectra_data):
        self.sample_name = sample_name
        self.spectra_data = spectra_data  # List of (filename, fit_params, integrated_intensity, concentration)
        self.calculate_statistics()
    
    def calculate_statistics(self):
        """Calculate statistical parameters for the sample group"""
        if not self.spectra_data:
            return
        
        # Extract values
        integrated_intensities = [data[2] for data in self.spectra_data]
        concentrations = [data[3] for data in self.spectra_data]
        
        # Calculate statistics
        self.n_spectra = len(integrated_intensities)
        self.mean_integrated_intensity = np.mean(integrated_intensities)
        self.std_integrated_intensity = np.std(integrated_intensities, ddof=1) if len(integrated_intensities) > 1 else 0
        self.mean_concentration = np.mean(concentrations)
        self.std_concentration = np.std(concentrations, ddof=1) if len(concentrations) > 1 else 0
        
        # Calculate relative standard deviation (RSD)
        self.rsd_integrated_intensity = (self.std_integrated_intensity / self.mean_integrated_intensity * 100) if self.mean_integrated_intensity != 0 else 0
        self.rsd_concentration = (self.std_concentration / self.mean_concentration * 100) if self.mean_concentration != 0 else 0
        
        # Standard error of the mean
        self.sem_integrated_intensity = self.std_integrated_intensity / np.sqrt(self.n_spectra) if self.n_spectra > 0 else 0
        self.sem_concentration = self.std_concentration / np.sqrt(self.n_spectra) if self.n_spectra > 0 else 0

class PlotCanvas(FigureCanvas):
    """Matplotlib canvas for displaying XRF spectra and fits"""
    
    def __init__(self, parent=None):
        self.fig = Figure(figsize=(12, 8))
        super().__init__(self.fig)
        self.setParent(parent)
        
        # Apply compact theme
        try:
            apply_theme('compact')
        except:
            pass
        
        # Initialize subplots immediately
        self.setup_subplots()
        
        # Store current data for zoom updates
        self.current_spectrum_data = None
        
        # Connect to zoom/pan events for auto Y-scaling
        self.setup_zoom_events()
        
    def setup_zoom_events(self):
        """Setup event handlers for zoom/pan to trigger Y-axis auto-scaling"""
        try:
            # Connect to xlim change events (fired when user zooms or pans)
            if hasattr(self, 'ax1'):
                self.ax1.callbacks.connect('xlim_changed', self.on_xlim_changed)
        except Exception as e:
            print(f"Warning: Could not setup zoom events: {e}")
    
    def on_xlim_changed(self, ax):
        """Called when X-axis limits change (zoom/pan)"""
        try:
            if self.current_spectrum_data is not None:
                x, y, fit_x, fit_y, background_x, background_y = self.current_spectrum_data
                x_min, x_max = ax.get_xlim()
                self.update_y_limits_for_zoom(x, y, fit_x, fit_y, background_x, background_y, x_min, x_max)
                self.draw_idle()  # Use draw_idle for better performance
        except Exception as e:
            print(f"Error in xlim_changed handler: {e}")
    
    def setup_subplots(self):
        """Setup subplots for spectrum and calibration"""
        self.fig.clear()
        
        # Main spectrum plot
        self.ax1 = self.fig.add_subplot(2, 2, (1, 2))
        self.ax1.set_xlabel('Energy (keV)')
        self.ax1.set_ylabel('Intensity (counts)')
        self.ax1.grid(True, alpha=0.3)
        
        # Sample statistics plot
        self.ax2 = self.fig.add_subplot(2, 2, 3)
        self.ax2.set_xlabel('Sample Number')
        self.ax2.set_ylabel('Concentration')
        self.ax2.grid(True, alpha=0.3)
        
        # Calibration verification plot
        self.ax3 = self.fig.add_subplot(2, 2, 4)
        self.ax3.set_xlabel('Integrated Intensity')
        self.ax3.set_ylabel('Concentration')
        self.ax3.grid(True, alpha=0.3)
        
        self.fig.tight_layout()
        
    def plot_spectrum(self, x, y, fit_x=None, fit_y=None, background_x=None, background_y=None, 
                     r_squared=None, concentration=None, title="XRF Spectrum"):
        """Plot spectrum with optional fit overlay, background curve, R² value, and Pb concentration"""
        try:
            # Ensure subplots are set up
            if not hasattr(self, 'ax1'):
                self.setup_subplots()
                self.setup_zoom_events()  # Reconnect events after subplot setup

            # Store current spectrum data for zoom event handling
            self.current_spectrum_data = (x, y, fit_x, fit_y, background_x, background_y)

            # 1. Save current zoom window (if any), else use display_min/max or default
            try:
                x_min, x_max = self.ax1.get_xlim()
                # If the current zoom is the default (0, 1), use display_min/max or default
                if x_min == 0.0 and x_max == 1.0:
                    raise Exception()
            except:
                if hasattr(self, 'display_min') and hasattr(self, 'display_max'):
                    x_min, x_max = self.display_min, self.display_max
                else:
                    x_min, x_max = 9.5, 11.5

            # 2. Clear and plot with ORIGINAL intensity values (no normalization)
            self.ax1.clear()
            
            # Reconnect zoom events after clearing (clearing removes callbacks)
            self.ax1.callbacks.connect('xlim_changed', self.on_xlim_changed)
            
            self.ax1.plot(x, y, 'b-', linewidth=1, label='Raw Data', alpha=0.7)
            if background_x is not None and background_y is not None:
                self.ax1.plot(background_x, background_y, 'g--', linewidth=1.5, label='Background', alpha=0.8)
            if fit_x is not None and fit_y is not None:
                self.ax1.plot(fit_x, fit_y, 'r-', linewidth=2, label='Gaussian-A Fit')

            self.ax1.set_xlabel('Energy (keV)')
            self.ax1.set_ylabel('Intensity (counts)')
            if r_squared is not None:
                title_with_r2 = f"{title} (R² = {r_squared:.4f})"
            else:
                title_with_r2 = title
            self.ax1.set_title(title_with_r2)
            self.ax1.grid(True, alpha=0.3)
            if concentration is not None:
                legend_text = f'Pb: {concentration:.2f} ppm'
                self.ax1.text(0.02, 0.98, legend_text, transform=self.ax1.transAxes, 
                             verticalalignment='top', bbox=dict(boxstyle='round', facecolor='lightblue', alpha=0.8),
                             fontsize=10, fontweight='bold')
            self.ax1.legend()
            self.ax1.axvspan(10.0, 11.0, alpha=0.2, color='yellow', label='Pb Peak Region')

            # 3. Restore X zoom window
            self.ax1.set_xlim(x_min, x_max)
            
            # 4. Auto-scale Y-axis based only on data within the current X zoom window
            self.update_y_limits_for_zoom(x, y, fit_x, fit_y, background_x, background_y, x_min, x_max)
            
            self.fig.tight_layout()
            self.draw()
        except Exception as e:
            print(f"Error in plot_spectrum: {e}")
            try:
                self.setup_subplots()
                self.setup_zoom_events()
                self.ax1.clear()
                self.ax1.plot(x, y, 'b-', linewidth=1, label='Raw Data', alpha=0.7)
                self.ax1.set_xlabel('Energy (keV)')
                self.ax1.set_ylabel('Intensity (counts)')
                self.ax1.set_title(title)
                self.ax1.grid(True, alpha=0.3)
                self.ax1.legend()
                self.fig.tight_layout()
                self.draw()
            except Exception as e2:
                print(f"Failed to recover from plotting error: {e2}")
    
    def update_y_limits_for_zoom(self, x, y, fit_x=None, fit_y=None, background_x=None, background_y=None, x_min=None, x_max=None):
        """Update Y-axis limits based on data within the current X zoom window"""
        try:
            # Get current X limits if not provided
            if x_min is None or x_max is None:
                x_min, x_max = self.ax1.get_xlim()
            
            # Collect all Y values within the X zoom window
            all_y_values = []
            
            # Raw data within zoom window
            mask = (x >= x_min) & (x <= x_max)
            if mask.any():
                all_y_values.extend(y[mask])
            
            # Fit data within zoom window
            if fit_x is not None and fit_y is not None:
                fit_mask = (fit_x >= x_min) & (fit_x <= x_max)
                if fit_mask.any():
                    all_y_values.extend(fit_y[fit_mask])
            
            # Background data within zoom window
            if background_x is not None and background_y is not None:
                bg_mask = (background_x >= x_min) & (background_x <= x_max)
                if bg_mask.any():
                    all_y_values.extend(background_y[bg_mask])
            
            # Calculate Y limits with some padding
            if all_y_values:
                y_min = min(all_y_values)
                y_max = max(all_y_values)
                
                # Add 5% padding on both sides
                y_range = y_max - y_min
                if y_range > 0:
                    padding = 0.05 * y_range
                    new_y_min = y_min - padding
                    new_y_max = y_max + padding
                    self.ax1.set_ylim(new_y_min, new_y_max)
                else:
                    # If all values are the same, add some default padding
                    self.ax1.set_ylim(y_min - 1, y_max + 1)
            else:
                # Fallback: use full data range
                self.ax1.set_ylim(y.min(), y.max())
                
        except Exception as e:
            print(f"Error updating Y limits: {e}")
            # Fallback: use automatic scaling
            try:
                self.ax1.relim()
                self.ax1.autoscale_view()
            except:
                pass
    
    def plot_sample_statistics(self, sample_groups):
        """Plot sample statistics and calibration verification"""
        if not sample_groups:
            return
        
        # Clear only the bottom subplots (ax2 and ax3) for statistics
        self.ax2.clear()
        self.ax3.clear()
        
        # Extract data for plotting
        sample_names = [group.sample_name for group in sample_groups]
        mean_concentrations = [group.mean_concentration for group in sample_groups]
        std_concentrations = [group.std_concentration for group in sample_groups]
        mean_intensities = [group.mean_integrated_intensity for group in sample_groups]
        
        # Sample statistics plot (NO error bars)
        x_pos = range(len(sample_names))
        self.ax2.plot(x_pos, mean_concentrations, 'o-', linewidth=2, markersize=6)
        self.ax2.set_xlabel('Sample Number')
        self.ax2.set_ylabel('Pb Concentration')
        self.ax2.set_title('Sample Concentrations')
        self.ax2.set_xticks(x_pos)
        self.ax2.set_xticklabels([f'S{i+1}' for i in x_pos])
        self.ax2.grid(True, alpha=0.3)

        # --- Dynamic x-axis label management ---
        num_labels = len(sample_names)
        if num_labels > 20:
            # Show every Nth label only
            N = max(1, num_labels // 20)
            for i, label in enumerate(self.ax2.get_xticklabels()):
                label.set_visible(i % N == 0)
            # Rotate and align
            import matplotlib.pyplot as plt
            plt.setp(self.ax2.get_xticklabels(), rotation=45, ha='right', fontsize=8)
        else:
            import matplotlib.pyplot as plt
            plt.setp(self.ax2.get_xticklabels(), rotation=0, ha='center', fontsize=10)
        # --------------------------------------
        
        # Calibration verification plot
        self.ax3.scatter(mean_intensities, mean_concentrations, s=60, alpha=0.7, color='blue')
        
        # Plot calibration line using CURRENT calibration parameters
        if mean_intensities:
            x_cal = np.linspace(min(mean_intensities), max(mean_intensities), 100)
            
            # Try to get current calibration parameters from parent GUI
            try:
                # Access the GUI through the parent chain
                gui_parent = self.parent()
                while gui_parent and not hasattr(gui_parent, 'fitter'):
                    gui_parent = gui_parent.parent()
                
                if gui_parent and hasattr(gui_parent, 'fitter'):
                    slope = gui_parent.fitter.calibration_slope
                    intercept = gui_parent.fitter.calibration_intercept
                    cal_name = getattr(gui_parent.fitter, 'calibration_name', None)
                    
                    # Determine calibration type for label
                    if cal_name:
                        cal_label = cal_name
                    elif abs(slope - 13.8913) < 0.0001 and abs(intercept - 0.0) < 0.0001:
                        cal_label = 'NIST Calibration'
                    else:
                        cal_label = 'Custom Calibration'
                else:
                    # Fallback to default NIST values
                    slope = 13.8913
                    intercept = 0.0
                    cal_label = 'NIST Calibration'
            except:
                # Fallback to default NIST values
                slope = 13.8913
                intercept = 0.0
                cal_label = 'NIST Calibration'
            
            y_cal = slope * x_cal + intercept
            self.ax3.plot(x_cal, y_cal, 'r--', linewidth=2, label=cal_label)
            self.ax3.legend()
        
        self.ax3.set_xlabel('Mean Integrated Intensity')
        self.ax3.set_ylabel('Mean Concentration')
        self.ax3.set_title('Calibration Verification')
        self.ax3.grid(True, alpha=0.3)
        
        self.fig.tight_layout()
        self.draw()

class ProcessingThread(QThread):
    """Thread for batch processing XRF files with sample grouping"""
    
    progress = Signal(int)
    error_occurred = Signal(str, str)
    finished = Signal(list, list)
    
    def __init__(self, file_paths, fitting_params, spectra_per_sample):
        super().__init__()
        self.file_paths = file_paths
        self.fitting_params = fitting_params
        self.spectra_per_sample = spectra_per_sample
        self.fitter = XRFPeakFitter()
        
    def run(self):
        """Process all files in the list with sample grouping"""
        results = []
        sample_groups = []
        
        try:
            for i, file_path in enumerate(self.file_paths):
                try:
                    # Read XRF data
                    data = self.read_xrf_file(file_path)
                    
                    if data is None:
                        self.error_occurred.emit(file_path, "Could not read file")
                        continue
                    
                    x, y = data
                    
                    # Fit peak
                    fit_params, fit_curve, r_squared, x_fit, integrated_intensity, concentration = self.fitter.fit_peak(
                        x, y, 
                        peak_region=(self.fitting_params['peak_min'], self.fitting_params['peak_max']),
                        background_subtract=self.fitting_params['background_subtract'],
                        integration_region=(self.fitting_params['integration_min'], self.fitting_params['integration_max'])
                    )
                    
                    # Store results
                    result = {
                        'filename': os.path.basename(file_path),
                        'filepath': file_path,
                        'fit_params': fit_params,
                        'r_squared': r_squared,
                        'integrated_intensity': integrated_intensity,
                        'concentration': concentration,
                        'x_data': x,
                        'y_data': y,
                        'fit_x': x_fit,
                        'fit_y': fit_curve
                    }
                    
                    results.append(result)
                    
                    # Emit progress
                    progress_value = min(98, int((i + 1) / len(self.file_paths) * 98))  # Cap at 98% during file processing
                    self.progress.emit(progress_value)
                    
                except Exception as e:
                    self.error_occurred.emit(file_path, str(e))
                    # Still emit progress even for failed files
                    progress_value = min(98, int((i + 1) / len(self.file_paths) * 98))
                    self.progress.emit(progress_value)
            
            # Group results by sample (this can take some time)
            self.progress.emit(99)  # Show 99% while grouping
            sample_groups = self.group_by_sample(results)
            
            # Emit finished signal (100% will be set in on_batch_finished)
            self.finished.emit(results, sample_groups)
            
        except Exception as e:
            # Handle any unexpected errors in the thread
            self.error_occurred.emit("THREAD_ERROR", f"Unexpected error in processing thread: {str(e)}")
            self.finished.emit(results, sample_groups)  # Still emit finished to reset UI
    
    def group_by_sample(self, results):
        """Group results by sample based on spectra_per_sample"""
        sample_groups = []
        
        for i in range(0, len(results), self.spectra_per_sample):
            sample_data = []
            sample_number = (i // self.spectra_per_sample) + 1
            sample_name = f"Sample_{sample_number}"
            
            # Get spectra for this sample
            for j in range(i, min(i + self.spectra_per_sample, len(results))):
                result = results[j]
                sample_data.append((
                    result['filename'],
                    result['fit_params'],
                    result['integrated_intensity'],
                    result['concentration']
                ))
            
            if sample_data:
                sample_group = SampleGroup(sample_name, sample_data)
                sample_groups.append(sample_group)
        
        return sample_groups
    
    def read_xrf_file(self, file_path):
        """Read XRF data from various file formats using smart detection"""
        try:
            # Use the smart parser that automatically detects file format
            x, y, format_type = parse_xrf_file_smart(file_path)
            
            if x is not None and y is not None:
                print(f"Successfully parsed {os.path.basename(file_path)} as {format_type} format")
                return x, y
            else:
                print(f"Failed to parse {file_path} with smart parser")
                return None
                
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return None

class MultiElementProcessingThread(QThread):
    """Thread for multi-element batch processing of XRF files"""
    
    progress = Signal(int)
    error_occurred = Signal(str, str)
    finished = Signal(list, list)
    
    def __init__(self, file_paths, fitting_params, spectra_per_sample):
        super().__init__()
        self.file_paths = file_paths
        self.fitting_params = fitting_params
        self.spectra_per_sample = spectra_per_sample
        self.selected_elements = fitting_params.get('selected_elements', ['Pb'])
        
        # Create fitters for each selected element
        self.element_fitters = {}
        for element in self.selected_elements:
            self.element_fitters[element] = XRFPeakFitter(element=element)
    
    def run(self):
        """Process all files for all selected elements"""
        results = []
        sample_groups = []
        
        try:
            for i, file_path in enumerate(self.file_paths):
                try:
                    # Read XRF data once
                    data = self.read_xrf_file(file_path)
                    
                    if data is None:
                        self.error_occurred.emit(file_path, "Could not read file")
                        continue
                    
                    x, y = data
                    
                    # Analyze each selected element
                    element_results = {}
                    
                    for element in self.selected_elements:
                        try:
                            fitter = self.element_fitters[element]
                            
                            # Use element-specific regions or fallback to UI parameters
                            element_data = ELEMENT_DEFINITIONS[element]
                            peak_region = element_data['peak_region']
                            integration_region = element_data['integration_region']
                            
                            # Fit peak for this element
                            fit_params, fit_curve, r_squared, x_fit, integrated_intensity, concentration = fitter.fit_peak(
                                x, y, 
                                peak_region=peak_region,
                                background_subtract=self.fitting_params['background_subtract'],
                                integration_region=integration_region
                            )
                            
                            element_results[element] = {
                                'fit_params': fit_params,
                                'r_squared': r_squared,
                                'integrated_intensity': integrated_intensity,
                                'concentration': concentration,
                                'fit_x': x_fit,
                                'fit_y': fit_curve
                            }
                            
                        except Exception as e:
                            # If one element fails, continue with others
                            element_results[element] = {
                                'error': str(e),
                                'fit_params': None,
                                'r_squared': 0,
                                'integrated_intensity': 0,
                                'concentration': 0
                            }
                    
                    # Store multi-element results
                    result = {
                        'filename': os.path.basename(file_path),
                        'filepath': file_path,
                        'x_data': x,
                        'y_data': y,
                        'element_results': element_results,
                        'selected_elements': self.selected_elements
                    }
                    
                    results.append(result)
                    
                    # Emit progress
                    progress_value = min(98, int((i + 1) / len(self.file_paths) * 98))
                    self.progress.emit(progress_value)
                    
                except Exception as e:
                    self.error_occurred.emit(file_path, str(e))
                    progress_value = min(98, int((i + 1) / len(self.file_paths) * 98))
                    self.progress.emit(progress_value)
            
            # Group results by sample
            self.progress.emit(99)
            sample_groups = self.group_by_sample_multi_element(results)
            
            # Emit finished signal
            self.finished.emit(results, sample_groups)
            
        except Exception as e:
            self.error_occurred.emit("THREAD_ERROR", f"Unexpected error in multi-element processing: {str(e)}")
            self.finished.emit(results, sample_groups)
    
    def group_by_sample_multi_element(self, results):
        """Group multi-element results by sample"""
        sample_groups = []
        
        for i in range(0, len(results), self.spectra_per_sample):
            sample_number = (i // self.spectra_per_sample) + 1
            sample_name = f"Sample_{sample_number}"
            
            # Create element-specific sample groups
            element_sample_groups = {}
            
            for element in self.selected_elements:
                sample_data = []
                
                # Get spectra for this sample and element
                for j in range(i, min(i + self.spectra_per_sample, len(results))):
                    result = results[j]
                    element_result = result['element_results'].get(element, {})
                    
                    if 'error' not in element_result:
                        sample_data.append((
                            result['filename'],
                            element_result.get('fit_params'),
                            element_result.get('integrated_intensity', 0),
                            element_result.get('concentration', 0)
                        ))
                
                if sample_data:
                    element_sample_groups[element] = SampleGroup(f"{sample_name}_{element}", sample_data)
            
            if element_sample_groups:
                # Store as a multi-element sample group
                multi_element_group = {
                    'sample_name': sample_name,
                    'element_groups': element_sample_groups,
                    'selected_elements': self.selected_elements
                }
                sample_groups.append(multi_element_group)
        
        return sample_groups
    
    def read_xrf_file(self, file_path):
        """Read XRF data from various file formats using smart detection"""
        try:
            x, y, format_type = parse_xrf_file_smart(file_path)
            
            if x is not None and y is not None:
                print(f"Successfully parsed {os.path.basename(file_path)} as {format_type} format")
                return x, y
            else:
                print(f"Failed to parse {file_path} with smart parser")
                return None
                
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return None

class XRFPeakFittingGUI(QMainWindow):
    """Main GUI application for XRF peak fitting with calibration and sample grouping"""
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle("XRF Peak Fitting with NIST Calibration - Pb L-alpha at 10.5 keV")
        self.setGeometry(100, 100, 1600, 1000)
        
        # Initialize components
        self.calibration_manager = CalibrationManager()
        self.fitter = XRFPeakFitter()  # Default to Pb
        self.peak_fitter = XRFPeakFitter()  # For multi-element calibration UI
        self.current_data = None
        self.processing_thread = None
        self.batch_results = []
        self.sample_groups = []
        
        # Load saved calibrations into fitters
        self.load_saved_calibrations()
        
        self.init_ui()
    
    def load_saved_calibrations(self):
        """Load saved calibrations into the fitters"""
        all_calibrations = self.calibration_manager.get_all_calibrations()
        
        for element, cal_data in all_calibrations.items():
            if element in ELEMENT_DEFINITIONS:
                slope = cal_data.get('slope', ELEMENT_DEFINITIONS[element]['default_calibration']['slope'])
                intercept = cal_data.get('intercept', ELEMENT_DEFINITIONS[element]['default_calibration']['intercept'])
                
                # Update both fitters
                self.fitter.update_element_calibration(element, slope, intercept)
                self.peak_fitter.update_element_calibration(element, slope, intercept)
                
                print(f"Loaded calibration for {element}: {slope:.4f}x + {intercept:.4f}")
        
    def init_ui(self):
        """Initialize the user interface"""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Create main layout
        main_layout = QHBoxLayout(central_widget)
        
        # Create tab widget for left panel
        self.tab_widget = QTabWidget()
        
        # Create main workflow tab
        main_tab = QWidget()
        main_tab_layout = QVBoxLayout(main_tab)
        
        # File selection group
        file_group = QGroupBox("File Selection")
        file_layout = QVBoxLayout(file_group)
        
        # Single file selection
        single_file_layout = QHBoxLayout()
        self.single_file_btn = QPushButton("Select Single File")
        self.single_file_btn.clicked.connect(self.select_single_file)
        self.single_file_label = QLabel("No file selected")
        single_file_layout.addWidget(self.single_file_btn)
        single_file_layout.addWidget(self.single_file_label)
        file_layout.addLayout(single_file_layout)
        
        # Batch file selection
        batch_file_layout = QHBoxLayout()
        self.batch_folder_btn = QPushButton("Select Folder for Batch Processing")
        self.batch_folder_btn.clicked.connect(self.select_batch_folder)
        self.batch_folder_label = QLabel("No folder selected")
        batch_file_layout.addWidget(self.batch_folder_btn)
        batch_file_layout.addWidget(self.batch_folder_label)
        file_layout.addLayout(batch_file_layout)
        
        main_tab_layout.addWidget(file_group)
        
        # Sample grouping parameters
        grouping_group = QGroupBox("Sample Grouping")
        grouping_layout = QGridLayout(grouping_group)
        
        grouping_layout.addWidget(QLabel("Spectra per Sample:"), 0, 0)
        self.spectra_per_sample_spin = QSpinBox()
        self.spectra_per_sample_spin.setRange(1, 20)
        self.spectra_per_sample_spin.setValue(6)
        grouping_layout.addWidget(self.spectra_per_sample_spin, 0, 1)
        
        main_tab_layout.addWidget(grouping_group)
        
        # Element selection group
        element_selection_group = QGroupBox("Elements to Analyze")
        element_selection_layout = QVBoxLayout(element_selection_group)
        
        # Instructions
        element_info = QLabel("Select which elements to analyze during processing:")
        element_info.setStyleSheet("QLabel { color: #666; font-style: italic; }")
        element_selection_layout.addWidget(element_info)
        
        # Create checkboxes for each element
        self.element_checkboxes = {}
        element_grid = QGridLayout()
        
        elements = list(ELEMENT_DEFINITIONS.keys())
        for i, element in enumerate(elements):
            element_data = ELEMENT_DEFINITIONS[element]
            checkbox = QCheckBox(f"{element} ({element_data['name']}) - {element_data['primary_energy']} keV")
            
            # Check Pb by default
            if element == 'Pb':
                checkbox.setChecked(True)
            
            self.element_checkboxes[element] = checkbox
            
            # Arrange in 2 columns
            row = i // 2
            col = i % 2
            element_grid.addWidget(checkbox, row, col)
        
        element_selection_layout.addLayout(element_grid)
        
        # Quick selection buttons
        quick_select_layout = QHBoxLayout()
        
        select_all_elements_btn = QPushButton("Select All")
        select_all_elements_btn.clicked.connect(self.select_all_elements)
        quick_select_layout.addWidget(select_all_elements_btn)
        
        select_none_elements_btn = QPushButton("Select None")
        select_none_elements_btn.clicked.connect(self.select_none_elements)
        quick_select_layout.addWidget(select_none_elements_btn)
        
        select_common_elements_btn = QPushButton("Common Elements (Pb, Zn, Cu, Cr)")
        select_common_elements_btn.clicked.connect(self.select_common_elements)
        quick_select_layout.addWidget(select_common_elements_btn)
        
        quick_select_layout.addStretch()
        element_selection_layout.addLayout(quick_select_layout)
        
        main_tab_layout.addWidget(element_selection_group)
        
        # Processing buttons
        process_group = QGroupBox("Processing")
        process_layout = QVBoxLayout(process_group)
        
        self.fit_single_btn = QPushButton("Fit Single File")
        self.fit_single_btn.clicked.connect(self.fit_single_file)
        self.fit_single_btn.setEnabled(False)
        process_layout.addWidget(self.fit_single_btn)
        
        # Note: Pb-As deconvolution is now automatic when both elements are selected
        # Keeping the button hidden but code available for manual use if needed
        # self.pb_as_deconv_btn = QPushButton("⚛️ Pb-As Deconvolution")
        # self.pb_as_deconv_btn.clicked.connect(self.run_pb_as_deconvolution)
        # self.pb_as_deconv_btn.setEnabled(False)
        # self.pb_as_deconv_btn.setToolTip("Simultaneous fitting of overlapping Pb Lα and As Kα peaks using all characteristic lines")
        # self.pb_as_deconv_btn.setStyleSheet("QPushButton { background-color: #FFE4B5; font-weight: bold; }")
        # process_layout.addWidget(self.pb_as_deconv_btn)
        
        self.fit_batch_btn = QPushButton("Process Batch")
        self.fit_batch_btn.clicked.connect(self.process_batch)
        self.fit_batch_btn.setEnabled(False)
        process_layout.addWidget(self.fit_batch_btn)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        process_layout.addWidget(self.progress_bar)
        
        main_tab_layout.addWidget(process_group)
        
        # Results display
        results_group = QGroupBox("Fit Results")
        results_layout = QVBoxLayout(results_group)
        
        self.results_text = QTextEdit()
        self.results_text.setMaximumHeight(250)
        self.results_text.setFont(QFont("Courier", 9))
        results_layout.addWidget(self.results_text)
        
        main_tab_layout.addWidget(results_group)
        

        

        
        # Spectrum Browser
        browser_group = QGroupBox("Spectrum Browser")
        browser_layout = QVBoxLayout(browser_group)
        
        # Browser controls
        browser_controls = QHBoxLayout()
        self.prev_spectrum_btn = QPushButton("◀ Previous")
        self.prev_spectrum_btn.clicked.connect(self.show_previous_spectrum)
        self.prev_spectrum_btn.setEnabled(False)
        browser_controls.addWidget(self.prev_spectrum_btn)
        
        self.spectrum_info_label = QLabel("No spectra available")
        self.spectrum_info_label.setAlignment(Qt.AlignCenter)
        browser_controls.addWidget(self.spectrum_info_label)
        
        self.next_spectrum_btn = QPushButton("Next ▶")
        self.next_spectrum_btn.clicked.connect(self.show_next_spectrum)
        self.next_spectrum_btn.setEnabled(False)
        browser_controls.addWidget(self.next_spectrum_btn)
        
        browser_layout.addLayout(browser_controls)
        
        # Quick navigation
        nav_layout = QHBoxLayout()
        nav_layout.addWidget(QLabel("Go to:"))
        self.spectrum_spinbox = QSpinBox()
        self.spectrum_spinbox.setMinimum(1)
        self.spectrum_spinbox.setMaximum(1)
        self.spectrum_spinbox.valueChanged.connect(self.go_to_spectrum)
        nav_layout.addWidget(self.spectrum_spinbox)
        nav_layout.addWidget(QLabel("of"))
        self.total_spectra_label = QLabel("0")
        nav_layout.addWidget(self.total_spectra_label)
        nav_layout.addStretch()
        browser_layout.addLayout(nav_layout)
        
        # Filter controls
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel("Filter by R²:"))
        self.r2_filter_spin = QDoubleSpinBox()
        self.r2_filter_spin.setRange(0.0, 1.0)
        self.r2_filter_spin.setValue(0.0)
        self.r2_filter_spin.setSingleStep(0.1)
        self.r2_filter_spin.setDecimals(3)
        self.r2_filter_spin.valueChanged.connect(self.apply_r2_filter)
        filter_layout.addWidget(self.r2_filter_spin)
        filter_layout.addWidget(QLabel("or higher"))
        filter_layout.addStretch()
        browser_layout.addLayout(filter_layout)
        
        main_tab_layout.addWidget(browser_group)
        
        # Plot controls removed - statistics now show automatically after batch processing
        
        # Add zoom controls for top plot
        zoom_group = QGroupBox("Top Plot Zoom (keV)")
        zoom_layout = QHBoxLayout(zoom_group)
        zoom_layout.addWidget(QLabel("Display Min (keV):"))
        self.display_min_spin = QDoubleSpinBox()
        self.display_min_spin.setRange(0.0, 20.0)
        self.display_min_spin.setValue(8.0)
        self.display_min_spin.setSingleStep(0.1)
        zoom_layout.addWidget(self.display_min_spin)
        zoom_layout.addWidget(QLabel("Display Max (keV):"))
        self.display_max_spin = QDoubleSpinBox()
        self.display_max_spin.setRange(0.0, 20.0)
        self.display_max_spin.setValue(13.0)
        self.display_max_spin.setSingleStep(0.1)
        zoom_layout.addWidget(self.display_max_spin)
        main_tab_layout.addWidget(zoom_group)
        # Connect zoom controls to update plot
        self.display_min_spin.valueChanged.connect(self.update_top_plot_zoom)
        self.display_max_spin.valueChanged.connect(self.update_top_plot_zoom)
        
        # Add stretch to push everything to top
        main_tab_layout.addStretch()
        
        # Initialize sorting-related attributes
        self.batch_file_paths = []
        self.sorted_file_paths = []
        
        # Create advanced parameters tab
        advanced_tab = QWidget()
        advanced_tab_layout = QVBoxLayout(advanced_tab)

        # Add a QTabWidget for subtabs in Advanced Parameters
        self.advanced_subtabs = QTabWidget()
        advanced_tab_layout.addWidget(self.advanced_subtabs)

        # Multi-Element Calibrations subtab
        multi_element_tab = QWidget()
        multi_element_layout = QVBoxLayout(multi_element_tab)
        
        # Calibration Status Display
        calibration_status_group = QGroupBox("Current Calibration Status")
        calibration_status_layout = QVBoxLayout(calibration_status_group)
        
        # Status info
        status_info = QLabel("Shows the current calibration curves for each element. Green = calibrated, Red = using defaults.")
        status_info.setStyleSheet("QLabel { color: #666; font-style: italic; margin-bottom: 10px; }")
        calibration_status_layout.addWidget(status_info)
        
        # Calibration status table
        self.calibration_status_table = QTableWidget()
        self.setup_calibration_status_table()
        calibration_status_layout.addWidget(self.calibration_status_table)
        
        # Calibration management buttons - Compact layout
        cal_mgmt_layout = QHBoxLayout()
        cal_mgmt_layout.setSpacing(4)
        
        # Compact button style for management buttons
        mgmt_btn_style = """
            QPushButton {{
                color: white;
                font-weight: bold;
                padding: 4px 8px;
                border-radius: 3px;
                font-size: 10px;
                min-height: 16px;
            }}
            QPushButton:hover {{
                background-color: {hover_color};
            }}
        """
        
        refresh_cal_btn = QPushButton("🔄 Refresh")
        refresh_cal_btn.setToolTip("Refresh calibration status display")
        refresh_cal_btn.clicked.connect(self.refresh_calibration_status)
        refresh_cal_btn.setStyleSheet(
            "QPushButton { background-color: #4CAF50; }" + 
            mgmt_btn_style.format(hover_color="#45a049")
        )
        cal_mgmt_layout.addWidget(refresh_cal_btn)
        
        export_cal_btn = QPushButton("📤 Export")
        export_cal_btn.setToolTip("Export calibrations to file")
        export_cal_btn.clicked.connect(self.export_calibrations)
        export_cal_btn.setStyleSheet(
            "QPushButton { background-color: #2196F3; }" + 
            mgmt_btn_style.format(hover_color="#1976D2")
        )
        cal_mgmt_layout.addWidget(export_cal_btn)
        
        import_cal_btn = QPushButton("📥 Import")
        import_cal_btn.setToolTip("Import calibrations from file")
        import_cal_btn.clicked.connect(self.import_calibrations)
        import_cal_btn.setStyleSheet(
            "QPushButton { background-color: #FF9800; }" + 
            mgmt_btn_style.format(hover_color="#F57C00")
        )
        cal_mgmt_layout.addWidget(import_cal_btn)
        
        reset_cal_btn = QPushButton("🗑️ Reset All")
        reset_cal_btn.setToolTip("Reset all calibrations to defaults")
        reset_cal_btn.clicked.connect(self.reset_all_calibrations)
        reset_cal_btn.setStyleSheet(
            "QPushButton { background-color: #f44336; }" + 
            mgmt_btn_style.format(hover_color="#d32f2f")
        )
        cal_mgmt_layout.addWidget(reset_cal_btn)
        
        cal_mgmt_layout.addStretch()
        calibration_status_layout.addLayout(cal_mgmt_layout)
        
        multi_element_layout.addWidget(calibration_status_group)
        
        # Element selection group with peak selection
        element_selection_group = QGroupBox("Element & Peak Selection")
        element_selection_layout = QVBoxLayout(element_selection_group)
        
        # Element selector row
        elem_row = QHBoxLayout()
        elem_row.addWidget(QLabel("Current Element:"))
        self.element_combo = QComboBox()
        self.element_combo.addItems(list(ELEMENT_DEFINITIONS.keys()))
        self.element_combo.setCurrentText('Pb')  # Default to Pb
        self.element_combo.currentTextChanged.connect(self.on_element_changed)
        elem_row.addWidget(self.element_combo)
        elem_row.addStretch()
        element_selection_layout.addLayout(elem_row)
        
        # Peak selection row (for elements with alternative peaks)
        peak_row = QHBoxLayout()
        peak_row.addWidget(QLabel("Peak to Use:"))
        self.peak_selection_combo = QComboBox()
        self.peak_selection_combo.addItems(['Primary (default)', 'Alternative'])
        self.peak_selection_combo.currentTextChanged.connect(self.on_peak_selection_changed)
        peak_row.addWidget(self.peak_selection_combo)
        
        self.peak_info_label = QLabel("")
        self.peak_info_label.setStyleSheet("QLabel { color: #666; font-style: italic; }")
        self.peak_info_label.setWordWrap(True)
        peak_row.addWidget(self.peak_info_label)
        peak_row.addStretch()
        element_selection_layout.addLayout(peak_row)
        
        multi_element_layout.addWidget(element_selection_group)
        
        # Element properties display
        element_props_group = QGroupBox("Element Properties")
        element_props_layout = QGridLayout(element_props_group)
        
        element_props_layout.addWidget(QLabel("Primary Energy (keV):"), 0, 0)
        self.element_energy_label = QLabel("10.55")
        element_props_layout.addWidget(self.element_energy_label, 0, 1)
        
        element_props_layout.addWidget(QLabel("Peak Region (keV):"), 1, 0)
        self.element_peak_region_label = QLabel("10.0 - 11.0")
        element_props_layout.addWidget(self.element_peak_region_label, 1, 1)
        
        element_props_layout.addWidget(QLabel("Integration Region (keV):"), 2, 0)
        self.element_integration_region_label = QLabel("9.8 - 11.2")
        element_props_layout.addWidget(self.element_integration_region_label, 2, 1)
        
        multi_element_layout.addWidget(element_props_group)
        
        # Element calibration parameters
        element_cal_group = QGroupBox("Element Calibration Parameters")
        element_cal_layout = QGridLayout(element_cal_group)
        
        element_cal_layout.addWidget(QLabel("Slope:"), 0, 0)
        self.element_slope_edit = QLineEdit("13.8913")
        element_cal_layout.addWidget(self.element_slope_edit, 0, 1)
        
        element_cal_layout.addWidget(QLabel("Intercept:"), 1, 0)
        self.element_intercept_edit = QLineEdit("0.0")
        element_cal_layout.addWidget(self.element_intercept_edit, 1, 1)
        
        # Update button for element calibration
        self.update_element_cal_btn = QPushButton("Update Element Calibration")
        self.update_element_cal_btn.clicked.connect(self.update_element_calibration)
        element_cal_layout.addWidget(self.update_element_cal_btn, 2, 0, 1, 2)
        
        multi_element_layout.addWidget(element_cal_group)
        
        # Reference materials table
        ref_materials_group = QGroupBox("Reference Materials Database")
        ref_materials_layout = QVBoxLayout(ref_materials_group)
        
        # Create table for reference materials
        self.ref_materials_table = QTableWidget()
        self.setup_reference_materials_table()
        ref_materials_layout.addWidget(self.ref_materials_table)
        
        # Buttons for reference materials - Compact layout
        ref_buttons_layout = QVBoxLayout()
        ref_buttons_layout.setSpacing(4)  # Reduce spacing between rows
        
        # Shared compact button style
        compact_btn_style = """
            QPushButton {{
                color: white;
                font-weight: bold;
                padding: 5px 8px;
                border-radius: 3px;
                min-height: 16px;
                font-size: 10px;
            }}
            QPushButton:hover {{
                background-color: {hover_color};
            }}
        """
        
        # First row of buttons
        ref_buttons_row1 = QHBoxLayout()
        ref_buttons_row1.setSpacing(4)
        
        self.create_calibration_btn = QPushButton("🧮 Create from Standards")
        self.create_calibration_btn.setToolTip("Create calibration from reference material standards")
        self.create_calibration_btn.clicked.connect(self.create_calibration_from_standards)
        self.create_calibration_btn.setStyleSheet(
            "QPushButton { background-color: #2196F3; }" + 
            compact_btn_style.format(hover_color="#1976D2")
        )
        ref_buttons_row1.addWidget(self.create_calibration_btn)
        
        self.auto_calibration_btn = QPushButton("🤖 Auto Current")
        self.auto_calibration_btn.setToolTip("Auto-calibrate current element from spectra files")
        self.auto_calibration_btn.clicked.connect(self.auto_calibrate_from_spectra)
        self.auto_calibration_btn.setStyleSheet(
            "QPushButton { background-color: #FF9800; }" + 
            compact_btn_style.format(hover_color="#F57C00")
        )
        ref_buttons_row1.addWidget(self.auto_calibration_btn)
        
        # Second row of buttons
        ref_buttons_row2 = QHBoxLayout()
        ref_buttons_row2.setSpacing(4)
        
        self.auto_calibrate_all_btn = QPushButton("🚀 Auto ALL Elements")
        self.auto_calibrate_all_btn.setToolTip("Auto-calibrate all elements simultaneously")
        self.auto_calibrate_all_btn.clicked.connect(self.auto_calibrate_all_elements)
        self.auto_calibrate_all_btn.setStyleSheet(
            "QPushButton { background-color: #E91E63; }" + 
            compact_btn_style.format(hover_color="#C2185B")
        )
        ref_buttons_row2.addWidget(self.auto_calibrate_all_btn)
        
        self.view_standards_btn = QPushButton("📊 View Plot")
        self.view_standards_btn.setToolTip("View standards concentration plot")
        self.view_standards_btn.clicked.connect(self.view_standards_plot)
        self.view_standards_btn.setStyleSheet(
            "QPushButton { background-color: #607D8B; }" + 
            compact_btn_style.format(hover_color="#546E7A")
        )
        ref_buttons_row2.addWidget(self.view_standards_btn)
        
        ref_buttons_layout.addLayout(ref_buttons_row1)
        ref_buttons_layout.addLayout(ref_buttons_row2)
        
        ref_materials_layout.addLayout(ref_buttons_layout)
        
        multi_element_layout.addWidget(ref_materials_group)
        multi_element_layout.addStretch()
        
        self.advanced_subtabs.addTab(multi_element_tab, "Multi-Element Calibrations")

        # Remove protocol subtab - will be replaced with button

        # Add stretch to push everything to top
        advanced_tab_layout.addStretch()
        
        # Create Docs/Export tab
        docs_export_tab = QWidget()
        docs_export_layout = QVBoxLayout(docs_export_tab)
        
        # Documentation section
        docs_group = QGroupBox("Documentation")
        docs_layout = QVBoxLayout(docs_group)
        
        self.view_protocol_btn = QPushButton("📋 View XRF SOP")
        self.view_protocol_btn.clicked.connect(self.show_protocol_dialog)
        docs_layout.addWidget(self.view_protocol_btn)
        
        docs_export_layout.addWidget(docs_group)
        
        # Export section
        export_group = QGroupBox("Export Results")
        export_layout = QVBoxLayout(export_group)
        
        self.export_individual_btn = QPushButton("Export Individual Results to CSV")
        self.export_individual_btn.clicked.connect(self.export_individual_results)
        self.export_individual_btn.setEnabled(False)
        export_layout.addWidget(self.export_individual_btn)
        
        self.export_samples_btn = QPushButton("Export Sample Statistics to CSV")
        self.export_samples_btn.clicked.connect(self.export_sample_statistics)
        self.export_samples_btn.setEnabled(False)
        export_layout.addWidget(self.export_samples_btn)
        
        docs_export_layout.addWidget(export_group)
        
        # Report Generation section
        report_group = QGroupBox("Report Generation")
        report_layout = QVBoxLayout(report_group)
        
        # Report options
        options_layout = QHBoxLayout()
        options_layout.addWidget(QLabel("Report Format:"))
        self.report_format_combo = QComboBox()
        self.report_format_combo.addItems(["PDF", "HTML"])  # Removed 'Word Document'
        options_layout.addWidget(self.report_format_combo)
        options_layout.addStretch()
        report_layout.addLayout(options_layout)
        
        # Report content options
        content_layout = QVBoxLayout()
        
        self.include_protocol_check = QCheckBox("Include Protocol Summary")
        self.include_protocol_check.setChecked(True)
        content_layout.addWidget(self.include_protocol_check)
        
        self.include_spectra_check = QCheckBox("Include Example Spectra Plots")
        self.include_spectra_check.setChecked(True)
        content_layout.addWidget(self.include_spectra_check)
        
        self.include_statistics_check = QCheckBox("Include Detailed Statistics Tables")
        self.include_statistics_check.setChecked(True)
        content_layout.addWidget(self.include_statistics_check)
        
        self.include_calibration_check = QCheckBox("Include Calibration Information")
        self.include_calibration_check.setChecked(True)
        content_layout.addWidget(self.include_calibration_check)
        
        report_layout.addLayout(content_layout)
        
        # Report generation buttons
        report_buttons_layout = QHBoxLayout()
        
        self.generate_all_reports_btn = QPushButton("📊 Generate All Sample Reports")
        self.generate_all_reports_btn.clicked.connect(self.generate_all_sample_reports)
        self.generate_all_reports_btn.setEnabled(False)
        report_buttons_layout.addWidget(self.generate_all_reports_btn)
        
        self.generate_single_report_btn = QPushButton("📋 Generate Single Report")
        self.generate_single_report_btn.clicked.connect(self.generate_single_sample_report)
        self.generate_single_report_btn.setEnabled(False)
        report_buttons_layout.addWidget(self.generate_single_report_btn)
        
        report_layout.addLayout(report_buttons_layout)
        
        docs_export_layout.addWidget(report_group)
        
        # Add stretch to push everything to top
        docs_export_layout.addStretch()
        
        # Create search tab
        search_tab = QWidget()
        self.setup_search_tab(search_tab)
        
        # Create calibration plots tab
        calibration_plots_tab = QWidget()
        self.setup_calibration_plots_tab(calibration_plots_tab)
        
        # Create FP (Fundamental Parameters) tab
        fp_tab = QWidget()
        self.setup_fp_tab(fp_tab)
        
        # Add tabs to tab widget
        self.tab_widget.addTab(search_tab, "Search")
        self.tab_widget.addTab(main_tab, "Quant")
        self.tab_widget.addTab(advanced_tab, "Calibrations")
        self.tab_widget.addTab(calibration_plots_tab, "Calibration Plots")
        self.tab_widget.addTab(fp_tab, "FP Method")
        self.tab_widget.addTab(docs_export_tab, "Docs / Export")
        
        # Create right panel for plot
        right_panel = QVBoxLayout()
        
        # Plot canvas with navigation toolbar
        self.plot_canvas = PlotCanvas()
        
        # Add navigation toolbar if available
        if CompactNavigationToolbar is not None:
            self.plot_toolbar = CompactNavigationToolbar(self.plot_canvas, self)
            right_panel.addWidget(self.plot_toolbar)
        else:
            # Fallback to standard toolbar
            try:
                from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT
                self.plot_toolbar = NavigationToolbar2QT(self.plot_canvas, self)
            except ImportError:
                try:
                    from matplotlib.backends.backend_qtagg import NavigationToolbar2QT
                    self.plot_toolbar = NavigationToolbar2QT(self.plot_canvas, self)
                except ImportError:
                    from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT
                    self.plot_toolbar = NavigationToolbar2QT(self.plot_canvas, self)
            right_panel.addWidget(self.plot_toolbar)
        
        right_panel.addWidget(self.plot_canvas)
        
        # Add panels to main layout
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.addWidget(self.tab_widget)
        left_widget.setMaximumWidth(650)  # Increased from 450 to allow more space
        
        right_widget = QWidget()
        right_widget.setLayout(right_panel)
        
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setSizes([600, 1000])  # Widened left panel from 450 to 600
        
        main_layout.addWidget(splitter)
        
    def select_single_file(self):
        """Select a single XRF file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select XRF File", "", 
            "All Files (*);;CSV Files (*.csv);;Text Files (*.txt);;Excel Files (*.xlsx)"
        )
        
        if file_path:
            self.single_file_label.setText(os.path.basename(file_path))
            self.current_file_path = file_path
            self.fit_single_btn.setEnabled(True)
            # pb_as_deconv_btn removed - deconvolution is now automatic
            
            # Try to load and display the file
            self.load_and_display_file(file_path)
    
    def select_batch_folder(self):
        """Select folder for batch processing"""
        folder_path = QFileDialog.getExistingDirectory(self, "Select Folder with XRF Files")
        
        if folder_path:
            # Get all files in the folder first, then let user filter by extension
            all_files = [os.path.join(folder_path, f) for f in os.listdir(folder_path) 
                        if os.path.isfile(os.path.join(folder_path, f))]
            
            if all_files:
                # Show sorting dialog with extension filtering
                self.show_sorting_dialog(all_files, folder_path)
            else:
                self.batch_folder_label.setText("No files found")
                self.fit_batch_btn.setEnabled(False)
    
    def load_and_display_file(self, file_path):
        """Load and display XRF file"""
        try:
            # Read file
            data = self.read_xrf_file(file_path)
            if data is None:
                QMessageBox.warning(self, "Error", f"Could not read file: {file_path}")
                return
            
            x, y = data
            self.current_data = (x, y)
            
            # Ensure plot canvas is set up
            if not hasattr(self.plot_canvas, 'ax1'):
                self.plot_canvas.setup_subplots()
            
            # Display spectrum
            self.plot_canvas.plot_spectrum(x, y, title=f"XRF Spectrum - {os.path.basename(file_path)}")
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Error loading file: {str(e)}")
    
    def read_xrf_file(self, file_path):
        """Read XRF data from file using smart format detection"""
        try:
            # Use the smart parser that automatically detects file format
            x, y, format_type = parse_xrf_file_smart(file_path)
            
            if x is not None and y is not None:
                print(f"Successfully parsed {os.path.basename(file_path)} as {format_type} format")
                return x, y
            else:
                print(f"Failed to parse {file_path} with smart parser")
                return None
                
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return None
    
    def fit_single_file(self):
        """Fit a single XRF file for all selected elements"""
        if not hasattr(self, 'current_data') or self.current_data is None:
            QMessageBox.warning(self, "Error", "No file loaded. Please select a file first.")
            return
        
        try:
            x, y = self.current_data
            
            # Get selected elements from checkboxes
            selected_elements = []
            for element, checkbox in self.element_checkboxes.items():
                if checkbox.isChecked():
                    selected_elements.append(element)
            
            if not selected_elements:
                QMessageBox.warning(self, "No Elements Selected", 
                                  "Please select at least one element to analyze.")
                return
            
            # Check if both Pb and As are selected - use automatic deconvolution
            has_pb = 'Pb' in selected_elements
            has_as = 'As' in selected_elements
            
            if has_pb and has_as:
                # Automatically use Pb-As deconvolution
                try:
                    pb_results, as_results, x_fit, combined_fit, r_squared = self.peak_fitter.fit_pb_as_deconvolution(x, y)
                    
                    # Display deconvolution results
                    results_text = f"🔬 Pb-As Deconvolution Results:\n"
                    results_text += f"File: {os.path.basename(self.current_file_path)}\n"
                    results_text += "=" * 50 + "\n\n"
                    results_text += f"✨ Automatic deconvolution used (Pb + As detected)\n\n"
                    
                    results_text += f"LEAD (Pb):\n"
                    results_text += f"  Lines Used: {', '.join(pb_results['lines_used'])}\n"
                    results_text += f"  Amplitude: {pb_results['amplitude']:.2f} ± {pb_results['amplitude_error']:.2f}\n"
                    results_text += f"  FWHM: {pb_results['fwhm']:.4f} keV\n"
                    results_text += f"  Integrated Intensity: {pb_results['integrated_intensity']:.1f} cps\n"
                    if pb_results['concentration']:
                        results_text += f"  Concentration: {pb_results['concentration']:.2f} ppm\n\n"
                    else:
                        results_text += f"  Concentration: N/A (no calibration)\n\n"
                    
                    results_text += f"ARSENIC (As):\n"
                    results_text += f"  Lines Used: {', '.join(as_results['lines_used'])}\n"
                    results_text += f"  Amplitude: {as_results['amplitude']:.2f} ± {as_results['amplitude_error']:.2f}\n"
                    results_text += f"  FWHM: {as_results['fwhm']:.4f} keV\n"
                    results_text += f"  Integrated Intensity: {as_results['integrated_intensity']:.1f} cps\n"
                    if as_results['concentration']:
                        results_text += f"  Concentration: {as_results['concentration']:.2f} ppm\n\n"
                    else:
                        results_text += f"  Concentration: N/A (no calibration)\n\n"
                    
                    results_text += f"Fit Quality: R² = {r_squared:.6f}\n"
                    
                    if pb_results['concentration'] and as_results['concentration']:
                        results_text += f"Pb/As Ratio: {pb_results['concentration']/as_results['concentration']:.2f}\n"
                    
                    self.results_text.clear()
                    self.results_text.append(results_text)
                    
                    # Plot deconvolution results
                    self.plot_canvas.figure.clear()
                    ax = self.plot_canvas.figure.add_subplot(111)
                    
                    ax.plot(x, y, 'b-', linewidth=0.8, alpha=0.5, label='Original Spectrum')
                    ax.plot(x_fit, combined_fit, 'r-', linewidth=2, label=f'Combined Fit (R²={r_squared:.4f})')
                    
                    # Mark characteristic lines
                    pb_lines = [10.5515, 10.4495, 12.6137]
                    as_lines = [10.5437, 10.5078, 11.7262]
                    
                    for energy in pb_lines:
                        ax.axvline(energy, color='green', linestyle='--', alpha=0.5, linewidth=1)
                    for energy in as_lines:
                        ax.axvline(energy, color='orange', linestyle='--', alpha=0.5, linewidth=1)
                    
                    ax.plot([], [], 'g--', alpha=0.5, label='Pb L-lines')
                    ax.plot([], [], color='orange', linestyle='--', alpha=0.5, label='As K-lines')
                    
                    ax.set_xlabel('Energy (keV)', fontsize=10)
                    ax.set_ylabel('Counts', fontsize=10)
                    ax.set_title('Pb-As Deconvolution (Automatic)', fontsize=12, fontweight='bold')
                    ax.set_xlim(9.5, 13.5)
                    ax.grid(True, alpha=0.3)
                    ax.legend(loc='best')
                    
                    self.plot_canvas.figure.tight_layout()
                    self.plot_canvas.draw()
                    
                    return  # Exit after deconvolution
                    
                except Exception as e:
                    QMessageBox.warning(self, "Deconvolution Failed", 
                                      f"Pb-As deconvolution failed: {str(e)}\n\nFalling back to individual fits.")
            
            # Standard multi-element fitting (not Pb+As together)
            results_text = f"Multi-Element Fit Results:\n"
            results_text += f"File: {os.path.basename(self.current_file_path)}\n"
            results_text += "=" * 50 + "\n\n"
            
            all_results = {}
            
            for element in selected_elements:
                try:
                    # Create element-specific fitter
                    element_fitter = XRFPeakFitter(element=element)
                    
                    # Check if calibration exists
                    if self.calibration_manager.has_calibration(element):
                        cal = self.calibration_manager.get_calibration(element)
                        element_fitter.update_element_calibration(element, cal['slope'], cal['intercept'])
                    
                    # Perform fit
                    fit_params, fit_curve, r_squared, x_fit, integrated_intensity, concentration = element_fitter.fit_peak(
                        x, y, 
                        peak_region=None,  # Use element defaults
                        background_subtract=True,
                        integration_region=None
                    )
                    
                    # Store results
                    all_results[element] = {
                        'fit_params': fit_params,
                        'r_squared': r_squared,
                        'integrated_intensity': integrated_intensity,
                        'concentration': concentration
                    }
                    
                    # Add to results text
                    element_name = ELEMENT_DEFINITIONS[element]['name']
                    results_text += f"{element} ({element_name}):\n"
                    results_text += f"  Peak Center: {fit_params['center']:.3f} keV\n"
                    results_text += f"  FWHM: {fit_params['fwhm']:.3f} keV\n"
                    results_text += f"  Integrated Intensity: {integrated_intensity:.1f} cps\n"
                    results_text += f"  Concentration: {concentration:.2f} ppm\n"
                    results_text += f"  R²: {r_squared:.4f}\n\n"
                    
                except Exception as e:
                    results_text += f"{element}: Fit failed - {str(e)}\n\n"
            
            # Display results
            self.results_text.clear()
            self.results_text.append(results_text)
            
            # Plot the first element's fit (for visualization)
            if selected_elements and selected_elements[0] in all_results:
                first_element = selected_elements[0]
                result = all_results[first_element]
                
                # Create fitter for plotting
                plot_fitter = XRFPeakFitter(element=first_element)
                if self.calibration_manager.has_calibration(first_element):
                    cal = self.calibration_manager.get_calibration(first_element)
                    plot_fitter.update_element_calibration(first_element, cal['slope'], cal['intercept'])
                
                # Refit for plotting (to get fit curve)
                fit_params, fit_curve, r_squared, x_fit, integrated_intensity, concentration = plot_fitter.fit_peak(
                    x, y, peak_region=None, background_subtract=True, integration_region=None
                )
                
                # Calculate background curve
                background_x = x_fit
                background_y = plot_fitter.linear_background(x_fit, 
                                                            fit_params['background_slope'], 
                                                            fit_params['background_intercept'])
                
                # Ensure plot canvas is initialized
                if not hasattr(self.plot_canvas, 'ax1'):
                    self.plot_canvas.setup_subplots()
                    self.plot_canvas.setup_zoom_events()
                
                # Plot the spectrum with fit
                self.plot_canvas.plot_spectrum(
                    x, y, 
                    fit_x=x_fit, 
                    fit_y=fit_curve,
                    background_x=background_x,
                    background_y=background_y,
                    r_squared=r_squared,
                    concentration=concentration,
                    title=f"XRF Spectrum with Gaussian-A Fit - {os.path.basename(self.current_file_path)} ({first_element})"
                )
            
        except Exception as e:
            QMessageBox.warning(self, "Fitting Error", f"Error during fitting: {str(e)}")
    
    def process_batch(self):
        """Process batch of files with sample grouping"""
        if not hasattr(self, 'batch_file_paths'):
            QMessageBox.warning(self, "Error", "No files selected for batch processing")
            return
        
        # Get selected elements
        selected_elements = self.get_selected_elements()
        if not selected_elements:
            QMessageBox.warning(self, "No Elements Selected", 
                              "Please select at least one element to analyze.")
            return
        
        # Get fitting parameters - now using element-specific regions from ELEMENT_DEFINITIONS
        # Background subtraction is always enabled for multi-element analysis
        fitting_params = {
            'background_subtract': True,
            'selected_elements': selected_elements
        }
        
        spectra_per_sample = self.spectra_per_sample_spin.value()
        
        # Start processing thread
        self.processing_thread = MultiElementProcessingThread(self.batch_file_paths, fitting_params, spectra_per_sample)
        self.processing_thread.progress.connect(self.progress_bar.setValue)
        self.processing_thread.error_occurred.connect(self.on_processing_error)
        self.processing_thread.finished.connect(self.on_batch_finished)
        
        # Disable buttons during processing
        self.fit_batch_btn.setEnabled(False)
        self.progress_bar.setValue(0)
        
        # Clear previous results and reset counters
        self.results_text.clear()
        self.files_processed_count = 0
        self.latest_processed_data = None
        
        # Initialize progress tracking
        self.total_files = len(self.batch_file_paths)
        
        # Start processing
        self.processing_thread.start()
    

    
    def on_processing_error(self, file_path, error_msg):
        """Handle processing error"""
        filename = os.path.basename(file_path)
        self.results_text.append(f"ERROR: {filename} - {error_msg}")
    
    def run_pb_as_deconvolution(self):
        """Run Pb-As deconvolution on the loaded spectrum"""
        if not hasattr(self, 'current_data') or self.current_data is None:
            QMessageBox.warning(self, "Error", "No file loaded. Please select a file first.")
            return
        
        try:
            x, y = self.current_data
            
            # Check if we have calibrations for both Pb and As
            has_pb_cal = self.calibration_manager.has_calibration('Pb')
            has_as_cal = self.calibration_manager.has_calibration('As')
            
            if not has_pb_cal or not has_as_cal:
                missing = []
                if not has_pb_cal:
                    missing.append('Pb')
                if not has_as_cal:
                    missing.append('As')
                
                response = QMessageBox.question(
                    self, "Missing Calibrations",
                    f"No calibration found for: {', '.join(missing)}\n\n"
                    f"Deconvolution will proceed, but concentrations cannot be calculated.\n"
                    f"Continue anyway?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )
                
                if response == QMessageBox.StandardButton.No:
                    return
            
            # Run deconvolution
            pb_results, as_results, x_fit, combined_fit, r_squared = self.peak_fitter.fit_pb_as_deconvolution(x, y)
            
            # Plot results
            self.plot_canvas.figure.clear()
            ax = self.plot_canvas.figure.add_subplot(111)
            
            # Plot original spectrum
            ax.plot(x, y, 'b-', linewidth=0.8, alpha=0.5, label='Original Spectrum')
            
            # Plot fitted curve
            ax.plot(x_fit, combined_fit, 'r-', linewidth=2, label=f'Combined Fit (R²={r_squared:.4f})')
            
            # Mark characteristic lines
            pb_lines = [10.5515, 10.4495, 12.6137]  # Lα1, Lα2, Lβ1
            as_lines = [10.5437, 10.5078, 11.7262]  # Kα1, Kα2, Kβ1
            
            for energy in pb_lines:
                ax.axvline(energy, color='green', linestyle='--', alpha=0.5, linewidth=1)
            for energy in as_lines:
                ax.axvline(energy, color='orange', linestyle='--', alpha=0.5, linewidth=1)
            
            # Add legend entries for lines
            ax.plot([], [], 'g--', alpha=0.5, label='Pb L-lines')
            ax.plot([], [], color='orange', linestyle='--', alpha=0.5, label='As K-lines')
            
            ax.set_xlabel('Energy (keV)', fontsize=10)
            ax.set_ylabel('Counts', fontsize=10)
            ax.set_title('Pb-As Deconvolution', fontsize=12, fontweight='bold')
            ax.set_xlim(9.5, 13.5)
            ax.grid(True, alpha=0.3)
            ax.legend(loc='best')
            
            self.plot_canvas.figure.tight_layout()
            self.plot_canvas.draw()
            
            # Display results
            results_text = f"""
=== Pb-As DECONVOLUTION RESULTS ===
File: {os.path.basename(self.current_file_path)}
Fit Quality: R² = {r_squared:.6f}

LEAD (Pb):
  Lines Used: {', '.join(pb_results['lines_used'])}
  Amplitude: {pb_results['amplitude']:.2f} ± {pb_results['amplitude_error']:.2f}
  FWHM: {pb_results['fwhm']:.4f} ± {pb_results['fwhm_error']:.4f} keV
  Integrated Intensity: {pb_results['integrated_intensity']:.1f} cps
  Concentration: {pb_results['concentration']:.2f} ppm""" if pb_results['concentration'] else "  Concentration: N/A (no calibration)"
            
            results_text += f"""

ARSENIC (As):
  Lines Used: {', '.join(as_results['lines_used'])}
  Amplitude: {as_results['amplitude']:.2f} ± {as_results['amplitude_error']:.2f}
  FWHM: {as_results['fwhm']:.4f} ± {as_results['fwhm_error']:.4f} keV
  Integrated Intensity: {as_results['integrated_intensity']:.1f} cps
  Concentration: {as_results['concentration']:.2f} ppm""" if as_results['concentration'] else "  Concentration: N/A (no calibration)"
            
            results_text += f"""

NOTES:
  • Pb Lα1 (10.5515 keV) and As Kα1 (10.5437 keV) overlap
  • Deconvolution uses all characteristic lines with theoretical intensity ratios
  • Pb Lβ1 (12.6 keV) and As Kβ1 (11.7 keV) are well separated
  • Both elements quantified simultaneously
"""
            
            self.results_text.setText(results_text)
            
            # Show success message
            msg = f"Deconvolution successful!\n\n"
            if pb_results['concentration'] and as_results['concentration']:
                msg += f"Pb: {pb_results['concentration']:.2f} ppm\n"
                msg += f"As: {as_results['concentration']:.2f} ppm\n"
                msg += f"Pb/As ratio: {pb_results['concentration']/as_results['concentration']:.2f}"
            else:
                msg += "Integrated intensities calculated.\nAdd calibrations to get concentrations."
            
            QMessageBox.information(self, "Deconvolution Complete", msg)
            
        except Exception as e:
            QMessageBox.critical(self, "Deconvolution Failed", 
                               f"Failed to perform Pb-As deconvolution:\n{str(e)}\n\n"
                               f"Make sure the spectrum covers the 10-13 keV range.")
    
    def on_batch_finished(self, results, sample_groups):
        """Handle batch processing completion"""
        # Ensure progress bar reaches 100%
        self.progress_bar.setValue(100)
        
        self.batch_results = results
        self.sample_groups = sample_groups
        self.fit_batch_btn.setEnabled(True)
        self.export_individual_btn.setEnabled(True)
        self.export_samples_btn.setEnabled(True)
        self.generate_all_reports_btn.setEnabled(True)
        self.generate_single_report_btn.setEnabled(True)
        
        # Initialize spectrum browser
        self.initialize_spectrum_browser()
        
        # Check if this is multi-element results
        is_multi_element = (results and 'element_results' in results[0])
        
        if is_multi_element:
            # Display multi-element sample statistics
            self.display_multi_element_sample_statistics(sample_groups)
            
            # Display multi-element summary
            self.display_multi_element_summary(results, sample_groups)
            
            # Plot concentration evolution for all elements
            self.plot_multi_element_concentration_evolution(results, sample_groups)
        else:
            # Display single-element sample statistics
            self.display_sample_statistics(sample_groups)
            
            # Plot concentration evolution for single element
            self.plot_concentration_evolution(results, sample_groups)
            
            # Display single-element summary
            successful = len([r for r in results if 'fit_params' in r])
            total = len(self.batch_file_paths)
            
            self.results_text.append(f"\n=== BATCH PROCESSING COMPLETE ===")
            self.results_text.append(f"Total files processed: {successful}/{total}")
            self.results_text.append(f"Samples analyzed: {len(sample_groups)}")
            
            QMessageBox.information(self, "Batch Processing Complete", 
                                  f"Processing complete!\n{successful}/{total} files processed successfully\n{len(sample_groups)} samples analyzed\n\nUse the Spectrum Browser to examine individual fits.")
    
    def display_fit_results(self, fit_params, r_squared, integrated_intensity, concentration):
        """Display fitting results for single file"""
        results_text = f"""
Gaussian-A Fit Results:
========================
Peak Center: {fit_params['center']:.4f} ± {fit_params['center_error']:.4f} keV
Amplitude: {fit_params['amplitude']:.2f} ± {fit_params['amplitude_error']:.2f}
FWHM: {fit_params['fwhm']:.4f} ± {fit_params['fwhm_error']:.4f} keV
Actual Peak Area: {fit_params['actual_peak_area']:.2f}
Background Slope: {fit_params['background_slope']:.4f}
Background Intercept: {fit_params['background_intercept']:.2f}
R-squared: {r_squared:.6f}

Quantitative Analysis:
=====================
Integrated Intensity: {integrated_intensity:.2f}
Calibrated Concentration: {concentration:.4f}
"""
        self.results_text.setText(results_text)
    
    def display_sample_statistics(self, sample_groups):
        """Display sample statistics in results text area"""
        self.results_text.append(f"\n=== SAMPLE STATISTICS ===")
        
        for group in sample_groups:
            self.results_text.append(f"\n{group.sample_name} (n={group.n_spectra}):")
            self.results_text.append(f"  Mean Intensity: {group.mean_integrated_intensity:.2f}")
            self.results_text.append(f"  Intensity SD: {group.std_integrated_intensity:.2f}")
            self.results_text.append(f"  Intensity SEM: {group.sem_integrated_intensity:.2f}")
            self.results_text.append(f"  Intensity RSD (%): {group.rsd_integrated_intensity:.2f}")
            self.results_text.append(f"  Mean Concentration: {group.mean_concentration:.4f}")
            self.results_text.append(f"  Concentration SD: {group.std_concentration:.4f}")
            self.results_text.append(f"  Concentration SEM: {group.sem_concentration:.4f}")
            self.results_text.append(f"  Concentration RSD (%): {group.rsd_concentration:.2f}")
    
    def plot_concentration_evolution(self, results, sample_groups):
        """Plot concentration evolution across all spectra for single element"""
        try:
            # Clear and setup figure
            self.plot_canvas.figure.clear()
            
            # Create two subplots: concentration evolution and sample statistics
            ax1 = self.plot_canvas.figure.add_subplot(2, 1, 1)
            ax2 = self.plot_canvas.figure.add_subplot(2, 1, 2)
            
            # Extract concentrations and file names
            concentrations = []
            file_names = []
            for result in results:
                if 'concentration' in result:
                    concentrations.append(result['concentration'])
                    file_names.append(os.path.basename(result['file_path']))
            
            if not concentrations:
                return
            
            # Plot 1: Concentration evolution
            x_indices = np.arange(len(concentrations))
            ax1.plot(x_indices, concentrations, 'bo-', linewidth=1.5, markersize=6, alpha=0.7)
            ax1.set_xlabel('Spectrum Index', fontsize=10)
            ax1.set_ylabel('Concentration (ppm)', fontsize=10)
            ax1.set_title('Concentration Evolution Across All Spectra', fontsize=11, fontweight='bold')
            ax1.grid(True, alpha=0.3)
            
            # Add mean line
            mean_conc = np.mean(concentrations)
            ax1.axhline(mean_conc, color='red', linestyle='--', linewidth=2, alpha=0.7, label=f'Mean: {mean_conc:.2f} ppm')
            ax1.legend()
            
            # Plot 2: Sample statistics (bar plot with error bars)
            if sample_groups:
                sample_names = [g.sample_name for g in sample_groups]
                means = [g.mean_concentration for g in sample_groups]
                sems = [g.sem_concentration for g in sample_groups]
                
                x_pos = np.arange(len(sample_names))
                bars = ax2.bar(x_pos, means, yerr=sems, capsize=5, alpha=0.7, edgecolor='black', linewidth=1.5)
                
                # Color bars by value
                colors = plt.cm.viridis(np.linspace(0.3, 0.9, len(means)))
                for bar, color in zip(bars, colors):
                    bar.set_color(color)
                
                ax2.set_xlabel('Sample', fontsize=10)
                ax2.set_ylabel('Concentration (ppm)', fontsize=10)
                ax2.set_title('Sample Statistics (Mean ± SEM)', fontsize=11, fontweight='bold')
                ax2.set_xticks(x_pos)
                ax2.set_xticklabels(sample_names, rotation=45, ha='right')
                ax2.grid(True, alpha=0.3, axis='y')
            
            self.plot_canvas.figure.tight_layout()
            self.plot_canvas.draw()
            
        except Exception as e:
            print(f"Error plotting concentration evolution: {e}")
    
    def plot_multi_element_concentration_evolution(self, results, sample_groups):
        """Plot concentration evolution for multiple elements"""
        try:
            # Get all elements
            if not results or 'element_results' not in results[0]:
                return
            
            elements = list(results[0]['element_results'].keys())
            n_elements = len(elements)
            
            if n_elements == 0:
                return
            
            # Clear and setup figure
            self.plot_canvas.figure.clear()
            
            # Determine grid layout
            if n_elements == 1:
                rows, cols = 1, 1
            elif n_elements == 2:
                rows, cols = 1, 2
            elif n_elements <= 4:
                rows, cols = 2, 2
            elif n_elements <= 6:
                rows, cols = 2, 3
            else:
                rows, cols = 3, 3
            
            # Plot each element
            for idx, element in enumerate(elements[:9]):  # Limit to 9 elements
                ax = self.plot_canvas.figure.add_subplot(rows, cols, idx + 1)
                
                # Extract concentrations for this element
                concentrations = []
                for result in results:
                    elem_result = result['element_results'].get(element, {})
                    if 'concentration' in elem_result:
                        concentrations.append(elem_result['concentration'])
                
                if not concentrations:
                    continue
                
                # Plot concentration evolution
                x_indices = np.arange(len(concentrations))
                ax.plot(x_indices, concentrations, 'o-', linewidth=1.5, markersize=5, alpha=0.7)
                
                # Add mean line
                mean_conc = np.mean(concentrations)
                ax.axhline(mean_conc, color='red', linestyle='--', linewidth=1.5, alpha=0.6)
                
                ax.set_xlabel('Spectrum #', fontsize=8)
                ax.set_ylabel('Conc. (ppm)', fontsize=8)
                ax.set_title(f'{element} - Mean: {mean_conc:.1f} ppm', fontsize=9, fontweight='bold')
                ax.grid(True, alpha=0.3)
                ax.tick_params(labelsize=7)
            
            self.plot_canvas.figure.suptitle('Multi-Element Concentration Evolution', fontsize=12, fontweight='bold')
            self.plot_canvas.figure.tight_layout()
            self.plot_canvas.draw()
            
        except Exception as e:
            print(f"Error plotting multi-element concentration evolution: {e}")
    
    def export_individual_results(self):
        """Export individual spectrum results to CSV"""
        if not self.batch_results:
            QMessageBox.warning(self, "Error", "No batch results to export")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Individual Results", "xrf_individual_results.csv", "CSV Files (*.csv)"
        )
        
        if file_path:
            try:
                # Prepare data for export
                export_data = []
                for i, result in enumerate(self.batch_results):
                    if 'fit_params' in result:
                        fp = result['fit_params']
                        sample_number = (i // self.spectra_per_sample_spin.value()) + 1
                        spectrum_number = (i % self.spectra_per_sample_spin.value()) + 1
                        
                        export_data.append({
                            'filename': result['filename'],
                            'sample_number': sample_number,
                            'spectrum_number': spectrum_number,
                            'peak_center_keV': fp['center'],
                            'peak_center_error_keV': fp['center_error'],
                            'amplitude': fp['amplitude'],
                            'amplitude_error': fp['amplitude_error'],
                            'fwhm_keV': fp['fwhm'],
                            'fwhm_error_keV': fp['fwhm_error'],
                            'actual_peak_area': fp['actual_peak_area'],
                            'background_slope': fp['background_slope'],
                            'background_intercept': fp['background_intercept'],
                            'r_squared': result['r_squared'],
                            'integrated_intensity': result['integrated_intensity'],
                            'concentration': result['concentration']
                        })
                
                # Create DataFrame and export
                df = pd.DataFrame(export_data)
                df.to_csv(file_path, index=False)
                
                QMessageBox.information(self, "Export Complete", f"Individual results exported to {file_path}")
                
            except Exception as e:
                QMessageBox.warning(self, "Export Error", f"Error exporting results: {str(e)}")
    
    def export_sample_statistics(self):
        """Export sample statistics to CSV"""
        if not self.sample_groups:
            QMessageBox.warning(self, "Error", "No sample statistics to export")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Sample Statistics", "xrf_sample_statistics.csv", "CSV Files (*.csv)"
        )
        
        if file_path:
            try:
                # Prepare data for export
                export_data = []
                for group in self.sample_groups:
                    export_data.append({
                        'sample_name': group.sample_name,
                        'n_spectra': group.n_spectra,
                        'mean_integrated_intensity': group.mean_integrated_intensity,
                        'std_integrated_intensity': group.std_integrated_intensity,
                        'rsd_integrated_intensity_percent': group.rsd_integrated_intensity,
                        'sem_integrated_intensity': group.sem_integrated_intensity,
                        'mean_concentration': group.mean_concentration,
                        'std_concentration': group.std_concentration,
                        'rsd_concentration_percent': group.rsd_concentration,
                        'sem_concentration': group.sem_concentration
                    })
                
                # Create DataFrame and export
                df = pd.DataFrame(export_data)
                df.to_csv(file_path, index=False)
                
                QMessageBox.information(self, "Export Complete", f"Sample statistics exported to {file_path}")
                
            except Exception as e:
                QMessageBox.warning(self, "Export Error", f"Error exporting sample statistics: {str(e)}")

    def initialize_spectrum_browser(self):
        """Initialize the spectrum browser with processed results"""
        if not self.batch_results:
            return
        
        # Check if multi-element or single-element results
        is_multi_element = (self.batch_results and 'element_results' in self.batch_results[0])
        
        if is_multi_element:
            # For multi-element, filter results that have at least one successful element fit
            self.filtered_results = [r for r in self.batch_results if 'element_results' in r and r['element_results']]
        else:
            # For single-element, filter successful results
            self.filtered_results = [r for r in self.batch_results if 'fit_params' in r]
        
        if not self.filtered_results:
            return
        
        # Initialize browser state
        self.current_spectrum_index = 0
        self.total_spectra = len(self.filtered_results)
        
        # Update UI controls
        self.total_spectra_label.setText(str(self.total_spectra))
        self.spectrum_spinbox.setMaximum(self.total_spectra)
        self.spectrum_spinbox.setValue(1)
        
        # Enable browser controls
        self.prev_spectrum_btn.setEnabled(True)
        self.next_spectrum_btn.setEnabled(True)
        
        # Show first spectrum
        self.show_current_spectrum()
    
    def show_current_spectrum(self):
        """Display the current spectrum in the browser"""
        if not hasattr(self, 'filtered_results') or not self.filtered_results:
            return
        
        if self.current_spectrum_index >= len(self.filtered_results):
            return
        
        result = self.filtered_results[self.current_spectrum_index]
        
        # Check if multi-element result
        is_multi_element = 'element_results' in result
        
        if is_multi_element:
            # For multi-element, show the first element's fit
            x = result['x_data']
            y = result['y_data']
            filename = result['filename']
            
            # Get first element's results
            element_results = result['element_results']
            if not element_results:
                return
            
            first_element = list(element_results.keys())[0]
            elem_result = element_results[first_element]
            
            fit_x = elem_result.get('fit_x')
            fit_y = elem_result.get('fit_y')
            fit_params = elem_result.get('fit_params')
            r_squared = elem_result.get('r_squared', 0)
            concentration = elem_result.get('concentration', 0)
        else:
            # Single element result
            x = result['x_data']
            y = result['y_data']
            fit_x = result['fit_x']
            fit_y = result['fit_y']
            filename = result['filename']
            fit_params = result['fit_params']
            r_squared = result['r_squared']
            concentration = result['concentration']
        
        # Store current data for zoom updates
        self.current_data = (x, y)
        
        # Store fit results for zoom updates
        self.current_fit_results = {
            'fit_params': fit_params,
            'fit_curve': fit_y,
            'r_squared': r_squared,
            'x_fit': fit_x,
            'integrated_intensity': result.get('integrated_intensity', 0),
            'concentration': concentration
        }
        
        # Calculate background curve for display
        background_x = None
        background_y = None
        if fit_params and 'background_slope' in fit_params and 'background_intercept' in fit_params:
            # Use the same x range as the fit for background display
            if fit_x is not None:
                background_x = fit_x
                background_y = self.fitter.linear_background(fit_x, 
                                                           fit_params['background_slope'], 
                                                           fit_params['background_intercept'])
        
        # Store background for zoom updates
        self.current_background = {'x': background_x, 'y': background_y}
        
        # Calculate sample and spectrum numbers for title
        # Get spectra per sample from the main GUI
        spectra_per_sample = getattr(self, 'spectra_per_sample_spin', None)
        if spectra_per_sample:
            spectra_per_sample_value = spectra_per_sample.value()
        else:
            spectra_per_sample_value = 6  # Default value
        
        # Calculate which sample this spectrum belongs to
        sample_number = (self.current_spectrum_index // spectra_per_sample_value) + 1
        spectrum_in_sample = (self.current_spectrum_index % spectra_per_sample_value) + 1
        
        # Create enhanced title with sample and spectrum information
        title = f"Sample {sample_number}, Spectrum {spectrum_in_sample}/{spectra_per_sample_value} (Overall: {self.current_spectrum_index + 1}/{self.total_spectra}): {filename}"
        
        # Update plot with background curve, R², and concentration
        # Ensure subplots are set up first if not already done
        if not hasattr(self.plot_canvas, 'ax1'):
            self.plot_canvas.setup_subplots()
        
        self.plot_canvas.plot_spectrum(
            x, y, 
            fit_x=fit_x, 
            fit_y=fit_y,
            background_x=background_x,
            background_y=background_y,
            r_squared=r_squared,
            concentration=concentration,
            title=title
        )
        
        # Update info label
        info_text = f"R² = {r_squared:.4f} | Conc = {concentration:.2f} ppm | Center = {fit_params['center']:.3f} keV"
        self.spectrum_info_label.setText(info_text)
        
        # Update spinbox
        self.spectrum_spinbox.blockSignals(True)
        self.spectrum_spinbox.setValue(self.current_spectrum_index + 1)
        self.spectrum_spinbox.blockSignals(False)
        
        # Update navigation buttons
        self.prev_spectrum_btn.setEnabled(self.current_spectrum_index > 0)
        self.next_spectrum_btn.setEnabled(self.current_spectrum_index < self.total_spectra - 1)
    
    def show_previous_spectrum(self):
        """Show the previous spectrum"""
        if hasattr(self, 'current_spectrum_index') and self.current_spectrum_index > 0:
            self.current_spectrum_index -= 1
            self.show_current_spectrum()
    
    def show_next_spectrum(self):
        """Show the next spectrum"""
        if hasattr(self, 'current_spectrum_index') and self.current_spectrum_index < self.total_spectra - 1:
            self.current_spectrum_index += 1
            self.show_current_spectrum()
    
    def go_to_spectrum(self, spectrum_number):
        """Go to a specific spectrum number"""
        if hasattr(self, 'filtered_results') and self.filtered_results:
            index = spectrum_number - 1
            if 0 <= index < len(self.filtered_results):
                self.current_spectrum_index = index
                self.show_current_spectrum()
    
    def apply_r2_filter(self, min_r2):
        """Filter spectra by minimum R² value"""
        if not hasattr(self, 'batch_results') or not self.batch_results:
            return
        
        # Filter results by R²
        self.filtered_results = [r for r in self.batch_results if 'fit_params' in r and r['r_squared'] >= min_r2]
        
        if not self.filtered_results:
            self.spectrum_info_label.setText("No spectra meet R² filter criteria")
            self.prev_spectrum_btn.setEnabled(False)
            self.next_spectrum_btn.setEnabled(False)
            self.total_spectra_label.setText("0")
            self.spectrum_spinbox.setMaximum(1)
            return
        
        # Reset to first spectrum
        self.current_spectrum_index = 0
        self.total_spectra = len(self.filtered_results)
        
        # Update UI
        self.total_spectra_label.setText(str(self.total_spectra))
        self.spectrum_spinbox.setMaximum(self.total_spectra)
        self.spectrum_spinbox.setValue(1)
        
        # Enable controls
        self.prev_spectrum_btn.setEnabled(True)
        self.next_spectrum_btn.setEnabled(True)
        
        # Show first filtered spectrum
        self.show_current_spectrum()
    
    def show_sorting_dialog(self, file_paths, folder_path):
        """Show dialog for file sorting options"""
        dialog = FileSortingDialog(file_paths, folder_path, self)
        if dialog.exec() == QDialog.Accepted:
            self.sorted_file_paths = dialog.get_sorted_files()
            self.batch_file_paths = self.sorted_file_paths  # Set the batch file paths for processing
            self.batch_folder_label.setText(f"{len(self.sorted_file_paths)} files sorted and ready")
            self.fit_batch_btn.setEnabled(True)
    
    def natural_sort_key(self, filename):
        """Generate a key for natural sorting of filenames"""
        import re
        
        # Extract the base filename without path
        basename = os.path.basename(filename)
        
        # Split the filename into text and number parts
        def convert(text):
            return int(text) if text.isdigit() else text.lower()
        
        # Use regex to split on numbers
        alphanum_key = [convert(c) for c in re.split('([0-9]+)', basename)]
        return alphanum_key
    
    def smart_sort_files(self, file_paths):
        """Sort files using natural sorting (handles numbers correctly)"""
        return sorted(file_paths, key=self.natural_sort_key)
    
    def detect_sorting_pattern(self, file_paths):
        """Detect the naming pattern in the files"""
        if not file_paths:
            return "unknown"
        
        # Get just the filenames
        filenames = [os.path.basename(f) for f in file_paths]
        
        # Check for common patterns
        patterns = {
            "sample_number": r"sample_(\d+)",
            "sample_number_extended": r"sample_(\d+)_",
            "number_only": r"^(\d+)",
            "letter_number": r"([A-Za-z]+)(\d+)",
            "date_pattern": r"(\d{4})[-_](\d{2})[-_](\d{2})",
            "time_pattern": r"(\d{2})[-:](\d{2})[-:](\d{2})"
        }
        
        for pattern_name, pattern in patterns.items():
            matches = [re.search(pattern, f) for f in filenames]
            if all(matches):
                return pattern_name
        
        return "unknown"
    

    
    def show_statistics_plot(self):
        """Show the sample statistics plot"""
        if hasattr(self, 'sample_groups') and self.sample_groups:
            self.plot_canvas.plot_sample_statistics(self.sample_groups)
    
    def update_top_plot_zoom(self):
        """Update the plot zoom range and redraw if data is available"""
        self.plot_canvas.display_min = self.display_min_spin.value()
        self.plot_canvas.display_max = self.display_max_spin.value()
        
        # Redraw current spectrum if available
        if hasattr(self, 'current_data') and self.current_data is not None:
            x, y = self.current_data
            
            # Check if we have fit results to redraw
            if hasattr(self, 'current_fit_results') and self.current_fit_results:
                fit_results = self.current_fit_results
                background = getattr(self, 'current_background', {'x': None, 'y': None})
                
                # Generate enhanced title if this is part of batch processing
                if hasattr(self, 'current_spectrum_index') and hasattr(self, 'total_spectra'):
                    # Get spectra per sample value
                    spectra_per_sample_value = getattr(self.spectra_per_sample_spin, 'value', lambda: 6)()
                    
                    # Calculate which sample this spectrum belongs to
                    sample_number = (self.current_spectrum_index // spectra_per_sample_value) + 1
                    spectrum_in_sample = (self.current_spectrum_index % spectra_per_sample_value) + 1
                    
                    title = f"Sample {sample_number}, Spectrum {spectrum_in_sample}/{spectra_per_sample_value} (Overall: {self.current_spectrum_index + 1}/{self.total_spectra}): {os.path.basename(self.current_file_path)}"
                else:
                    title = f"XRF Spectrum with Gaussian-A Fit - {os.path.basename(self.current_file_path)}"
                
                self.plot_canvas.plot_spectrum(
                    x, y,
                    fit_x=fit_results['x_fit'],
                    fit_y=fit_results['fit_curve'],
                    background_x=background['x'],
                    background_y=background['y'],
                    r_squared=fit_results['r_squared'],
                    concentration=fit_results['concentration'],
                    title=title
                )
            else:
                # Just raw data
                self.plot_canvas.plot_spectrum(x, y, title="XRF Spectrum")
        
        # Also update real-time plot if available
        if hasattr(self, 'latest_processed_data'):
            self.update_real_time_plot()

    def show_protocol_dialog(self):
        """Show the XRF SOP in a popup dialog with markdown formatting"""
        dialog = ProtocolDialog(self)
        dialog.exec()
    
    def export_protocol_text(self):
        """Export XRF SOP text to file"""
        file_path, _ = QFileDialog.getSaveFileName(self, "Export XRF SOP", "xrf_sop.txt", "Text Files (*.txt)")
        if file_path:
            try:
                with open('xrf_sop_markdown.md', 'r') as f:
                    protocol_text = f.read()
                with open(file_path, 'w') as f:
                    f.write(protocol_text)
                QMessageBox.information(self, "Export Complete", f"XRF SOP exported to {file_path}")
            except Exception as e:
                QMessageBox.warning(self, "Export Error", f"Error exporting XRF SOP: {e}")
    
    def generate_all_sample_reports(self):
        """Generate reports for all samples"""
        if not hasattr(self, 'sample_groups') or not self.sample_groups:
            QMessageBox.warning(self, "Error", "No sample data available for report generation")
            return
        
        # Get output directory
        output_dir = QFileDialog.getExistingDirectory(self, "Select Output Directory for Reports")
        if not output_dir:
            return
        
        try:
            # Generate reports for each sample
            for i, sample_group in enumerate(self.sample_groups):
                self.generate_sample_report(sample_group, output_dir, f"Sample_{i+1}")
            
            QMessageBox.information(self, "Report Generation Complete", 
                                  f"Generated {len(self.sample_groups)} sample reports in {output_dir}")
            
        except Exception as e:
            QMessageBox.warning(self, "Report Generation Error", f"Error generating reports: {str(e)}")
    
    def generate_single_sample_report(self):
        """Generate a report for the currently selected sample"""
        if not hasattr(self, 'sample_groups') or not self.sample_groups:
            QMessageBox.warning(self, "Error", "No sample data available for report generation")
            return
        
        # Get current sample index
        if not hasattr(self, 'current_spectrum_index') or not hasattr(self, 'filtered_results'):
            QMessageBox.warning(self, "Error", "Please select a sample in the spectrum browser first")
            return
        
        # Find which sample the current spectrum belongs to
        sample_index = self.current_spectrum_index // self.spectra_per_sample_spin.value()
        if sample_index >= len(self.sample_groups):
            QMessageBox.warning(self, "Error", "Invalid sample index")
            return
        
        # Get output directory
        output_dir = QFileDialog.getExistingDirectory(self, "Select Output Directory for Report")
        if not output_dir:
            return
        
        try:
            sample_group = self.sample_groups[sample_index]
            self.generate_sample_report(sample_group, output_dir, f"Sample_{sample_index+1}")
            
            QMessageBox.information(self, "Report Generation Complete", 
                                  f"Generated report for {sample_group.sample_name} in {output_dir}")
            
        except Exception as e:
            QMessageBox.warning(self, "Report Generation Error", f"Error generating report: {str(e)}")
    
    def generate_sample_report(self, sample_group, output_dir, filename_prefix):
        """Generate a comprehensive report for a single sample"""
        report_format = self.report_format_combo.currentText()
        
        if report_format == "HTML":
            self.generate_html_report(sample_group, output_dir, filename_prefix)
        elif report_format == "PDF":
            self.generate_pdf_report(sample_group, output_dir, filename_prefix)
        # Word report option removed

    def generate_pdf_report(self, sample_group, output_dir, filename_prefix):
        """Generate PDF report for a sample (now matches HTML content)"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            import pandas as pd

            # Create PDF file
            pdf_file = os.path.join(output_dir, f"{filename_prefix}_XRF_Report.pdf")
            doc = SimpleDocTemplate(pdf_file, pagesize=letter)
            styles = getSampleStyleSheet()

            # Build PDF content
            story = []

            # Header with logos
            header_data = []
            # Left logo (Pb logo)
            if os.path.exists('Pb_logo.png'):
                pb_logo = Image('Pb_logo.png', width=1.5*inch, height=1*inch, kind='proportional')
                header_data.append(pb_logo)
            else:
                header_data.append(Paragraph("", styles['Normal']))
            # Center title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Center
            )
            header_data.append(Paragraph(f"XRF Analysis Report<br/>{sample_group.sample_name}", title_style))
            # Right logo (NHM logo)
            if os.path.exists('NHM_logo_black2.jpg'):
                nhm_logo = Image('NHM_logo_black2.jpg', width=1.5*inch, height=1*inch, kind='proportional')
                header_data.append(nhm_logo)
            else:
                header_data.append(Paragraph("", styles['Normal']))
            # Create header table
            header_table = Table([header_data], colWidths=[2*inch, 4*inch, 2*inch])
            header_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(header_table)
            story.append(Spacer(1, 20))

            # Protocol Summary
            story.append(Paragraph("<b>Method:</b> XRF Analysis of Lead (Pb) in Pressed Pellets", styles['Normal']))
            story.append(Paragraph("<b>Peak:</b> Pb L-alpha at 10.5 keV", styles['Normal']))
            story.append(Paragraph("<b>Fitting:</b> Gaussian-A function with linear background", styles['Normal']))
            story.append(Paragraph("<b>Calibration:</b> NIST calibration curve (Concentration = 13.8913 × Intensity + 0)", styles['Normal']))
            story.append(Spacer(1, 12))

            # Sample Information
            story.append(Paragraph("<b>Sample Information</b>", styles['Heading2']))
            story.append(Paragraph(f"<b>Sample Name:</b> {sample_group.sample_name}", styles['Normal']))
            story.append(Paragraph(f"<b>Number of Spectra:</b> {sample_group.n_spectra}", styles['Normal']))
            story.append(Paragraph(f"<b>Analysis Date:</b> {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 12))

            # Statistical Summary Table
            story.append(Paragraph("<b>Statistical Summary</b>", styles['Heading2']))
            data = [
                ["Parameter", "Value", "Standard Deviation", "RSD (%)"],
                ["Mean Integrated Intensity", f"{sample_group.mean_integrated_intensity:.2f}", f"{sample_group.std_integrated_intensity:.2f}", f"{sample_group.rsd_integrated_intensity:.2f}"],
                ["Mean Concentration (ppm)", f"{sample_group.mean_concentration:.4f}", f"{sample_group.std_concentration:.4f}", f"{sample_group.rsd_concentration:.2f}"]
            ]
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(table)
            story.append(Spacer(1, 12))

            # Calibration Information
            story.append(Paragraph("<b>Calibration Information</b>", styles['Heading2']))
            story.append(Paragraph("<b>Calibration Equation:</b> Concentration = 13.8913 × Integrated Intensity + 0", styles['Normal']))
            story.append(Paragraph("<b>Calibration Source:</b> NIST Standard Reference Materials", styles['Normal']))
            story.append(Paragraph("<b>Dilution Factor:</b> 0.833 (2.0g sample + 0.4g binder)", styles['Normal']))
            story.append(Spacer(1, 12))

            # Quality Control
            story.append(Paragraph("<b>Quality Control</b>", styles['Heading2']))
            if sample_group.rsd_concentration <= 5.0:
                qc_status = "PASS"
                qc_color = colors.green
            else:
                qc_status = "FAIL"
                qc_color = colors.red
            qc_text = f"<b>Precision (RSD):</b> <font color='{qc_color.hexval()}'>{sample_group.rsd_concentration:.2f}% - {qc_status}</font>"
            story.append(Paragraph(qc_text, styles['Normal']))
            story.append(Paragraph("<b>Acceptance Criteria:</b> RSD ≤ 5.0%", styles['Normal']))

            # Build PDF
            doc.build(story)

        except ImportError:
            QMessageBox.warning(self, "PDF Generation Error", 
                              "PDF generation requires reportlab. Please install with: pip install reportlab")
        except Exception as e:
            QMessageBox.warning(self, "PDF Generation Error", f"Error generating PDF: {str(e)}")

    def generate_html_report(self, sample_group, output_dir, filename_prefix):
        """Generate HTML report for a sample"""
        html_content = self.create_report_content(sample_group, "html")
        
        # Save HTML file
        html_file = os.path.join(output_dir, f"{filename_prefix}_XRF_Report.html")
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Copy logo files to output directory for HTML display
        try:
            if os.path.exists('Pb_logo.png'):
                import shutil
                shutil.copy2('Pb_logo.png', output_dir)
            
            if os.path.exists('NHM_logo_black2.jpg'):
                import shutil
                shutil.copy2('NHM_logo_black2.jpg', output_dir)
        except Exception as e:
            print(f"Warning: Could not copy logo files: {e}")
    
    def create_report_content(self, sample_group, format_type):
        """Create the content for a report"""
        content = []
        
        # Header with logos
        content.append("<div style='display: flex; justify-content: space-between; align-items: center; margin-bottom: 30px;'>")
        
        # Left logo (Pb logo)
        if os.path.exists('Pb_logo.png'):
            content.append("<div style='flex: 1; text-align: left;'>")
            content.append("<img src='Pb_logo.png' alt='Pb Logo' style='max-height: 80px; max-width: 200px; object-fit: contain;'>")
            content.append("</div>")
        else:
            content.append("<div style='flex: 1;'></div>")
        
        # Center title
        content.append("<div style='flex: 2; text-align: center;'>")
        content.append(f"<h1 style='margin: 0;'>XRF Analysis Report - {sample_group.sample_name}</h1>")
        content.append("</div>")
        
        # Right logo (NHM logo)
        if os.path.exists('NHM_logo_black2.jpg'):
            content.append("<div style='flex: 1; text-align: right;'>")
            content.append("<img src='NHM_logo_black2.jpg' alt='NHM Logo' style='max-height: 80px; max-width: 200px; object-fit: contain;'>")
            content.append("</div>")
        else:
            content.append("<div style='flex: 1;'></div>")
        
        content.append("</div>")
        
        # Protocol Summary
        if self.include_protocol_check.isChecked():
            content.append("<h2>Protocol Summary</h2>")
            content.append("<p><strong>Method:</strong> XRF Analysis of Lead (Pb) in Pressed Pellets</p>")
            content.append("<p><strong>Peak:</strong> Pb L-alpha at 10.5 keV</p>")
            content.append("<p><strong>Fitting:</strong> Gaussian-A function with linear background</p>")
            content.append("<p><strong>Calibration:</strong> NIST calibration curve (Concentration = 13.8913 × Intensity + 0)</p>")
        
        # Sample Information
        content.append("<h2>Sample Information</h2>")
        content.append(f"<p><strong>Sample Name:</strong> {sample_group.sample_name}</p>")
        content.append(f"<p><strong>Number of Spectra:</strong> {sample_group.n_spectra}</p>")
        content.append(f"<p><strong>Analysis Date:</strong> {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}</p>")
        
        # Statistics Table
        if self.include_statistics_check.isChecked():
            content.append("<h2>Statistical Summary</h2>")
            content.append("<table border='1' style='border-collapse: collapse; width: 100%;'>")
            content.append("<tr><th>Parameter</th><th>Value</th><th>Standard Deviation</th><th>RSD (%)</th></tr>")
            content.append(f"<tr><td>Mean Integrated Intensity</td><td>{sample_group.mean_integrated_intensity:.2f}</td><td>{sample_group.std_integrated_intensity:.2f}</td><td>{sample_group.rsd_integrated_intensity:.2f}</td></tr>")
            content.append(f"<tr><td>Mean Concentration (ppm)</td><td>{sample_group.mean_concentration:.4f}</td><td>{sample_group.std_concentration:.4f}</td><td>{sample_group.rsd_concentration:.2f}</td></tr>")
            content.append("</table>")
        
        # Calibration Information
        if self.include_calibration_check.isChecked():
            content.append("<h2>Calibration Information</h2>")
            content.append("<p><strong>Calibration Equation:</strong> Concentration = 13.8913 × Integrated Intensity + 0</p>")
            content.append("<p><strong>Calibration Source:</strong> NIST Standard Reference Materials</p>")
            content.append("<p><strong>Dilution Factor:</strong> 0.833 (2.0g sample + 0.4g binder)</p>")
        
        # Example Spectrum (if available)
        if self.include_spectra_check.isChecked() and hasattr(self, 'batch_results'):
            content.append("<h2>Example Spectrum</h2>")
            content.append("<p>Representative spectrum with Gaussian-A fit shown in the main application window.</p>")
            content.append("<p><strong>Peak Center:</strong> ~10.5 keV (Pb L-alpha)</p>")
            content.append("<p><strong>Integration Region:</strong> 9.8 - 11.2 keV</p>")
        
        # Quality Control
        content.append("<h2>Quality Control</h2>")
        if sample_group.rsd_concentration <= 5.0:
            qc_status = "PASS"
            qc_color = "green"
        else:
            qc_status = "FAIL"
            qc_color = "red"
        
        content.append(f"<p><strong>Precision (RSD):</strong> <span style='color: {qc_color};'>{sample_group.rsd_concentration:.2f}% - {qc_status}</span></p>")
        content.append("<p><strong>Acceptance Criteria:</strong> RSD ≤ 5.0%</p>")
        
        # HTML wrapper
        if format_type == "html":
            html_wrapper = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>XRF Report - {sample_group.sample_name}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                    h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                    h2 {{ color: #34495e; margin-top: 30px; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f8f9fa; font-weight: bold; }}
                    .qc-pass {{ color: green; font-weight: bold; }}
                    .qc-fail {{ color: red; font-weight: bold; }}
                </style>
            </head>
            <body>
                {''.join(content)}
            </body>
            </html>
            """
            return html_wrapper
        
        return content
    
    def create_pdf_content(self, sample_group, styles):
        """Create content for PDF report"""
        story = []
        
        # Add content sections here (simplified for now)
        story.append(Paragraph(f"Sample: {sample_group.sample_name}", styles['Heading2']))
        story.append(Paragraph(f"Mean Concentration: {sample_group.mean_concentration:.4f} ppm", styles['Normal']))
        story.append(Paragraph(f"RSD: {sample_group.rsd_concentration:.2f}%", styles['Normal']))
        
        return story
    
    def create_word_content(self, doc, sample_group):
        """Create content for Word document"""
        # Add content sections here (simplified for now)
        doc.add_heading("Sample Information", level=2)
        doc.add_paragraph(f"Sample: {sample_group.sample_name}")
        doc.add_paragraph(f"Mean Concentration: {sample_group.mean_concentration:.4f} ppm")
        doc.add_paragraph(f"RSD: {sample_group.rsd_concentration:.2f}%")

    def setup_reference_materials_table(self):
        """Setup the reference materials table with certified values"""
        elements = list(ELEMENT_DEFINITIONS.keys())
        materials = list(REFERENCE_MATERIALS.keys())
        
        self.ref_materials_table.setRowCount(len(elements))
        self.ref_materials_table.setColumnCount(len(materials) + 1)
        
        # Set headers
        headers = ['Element'] + materials
        self.ref_materials_table.setHorizontalHeaderLabels(headers)
        
        # Populate table
        for i, element in enumerate(elements):
            # Element name
            element_item = QTableWidgetItem(f"{element} ({ELEMENT_DEFINITIONS[element]['name']})")
            element_item.setFlags(element_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.ref_materials_table.setItem(i, 0, element_item)
            
            # Certified values for each material
            for j, material in enumerate(materials):
                value = REFERENCE_MATERIALS[material].get(element)
                if value is not None:
                    if isinstance(value, str):
                        display_value = value
                    else:
                        display_value = str(value)
                else:
                    display_value = "N/A"
                
                value_item = QTableWidgetItem(display_value)
                value_item.setFlags(value_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                
                # Color coding for availability
                if value is not None and value != "N/A":
                    value_item.setBackground(QColor(144, 238, 144))  # Light green
                else:
                    value_item.setBackground(QColor(211, 211, 211))  # Light gray
                
                self.ref_materials_table.setItem(i, j + 1, value_item)
        
        # Resize columns to content
        self.ref_materials_table.resizeColumnsToContents()
        
        # Set alternating row colors
        self.ref_materials_table.setAlternatingRowColors(True)
    
    def on_element_changed(self, element_symbol):
        """Handle element selection change"""
        if element_symbol in ELEMENT_DEFINITIONS:
            element_data = ELEMENT_DEFINITIONS[element_symbol]
            
            # Check if element has alternative peak
            has_alternative = 'alternative_peak' in element_data
            self.peak_selection_combo.setEnabled(has_alternative)
            
            if has_alternative:
                alt_peak = element_data['alternative_peak']
                self.peak_info_label.setText(f"Alternative: {alt_peak['name']} at {alt_peak['energy']} keV - {alt_peak['note']}")
            else:
                self.peak_info_label.setText("No alternative peak available")
                self.peak_selection_combo.setCurrentText('Primary (default)')
            
            # Update display based on current peak selection
            self.update_element_display(element_symbol)
            
            # Update peak fitter to current element
            self.peak_fitter.set_element(element_symbol)
            
            # Update main fitter as well
            self.fitter.set_element(element_symbol)
            
            # Update fitting parameters in the main UI
            self.update_fitting_parameters_for_element(element_symbol)
    
    def on_peak_selection_changed(self, peak_choice):
        """Handle peak selection change (primary vs alternative)"""
        current_element = self.element_combo.currentText()
        use_alternative = (peak_choice == 'Alternative')
        
        # Update both fitters
        self.peak_fitter.set_use_alternative_peak(current_element, use_alternative)
        self.fitter.set_use_alternative_peak(current_element, use_alternative)
        
        # Update display
        self.update_element_display(current_element)
    
    def update_element_display(self, element_symbol):
        """Update element properties display based on selected peak"""
        if element_symbol not in ELEMENT_DEFINITIONS:
            return
        
        element_data = ELEMENT_DEFINITIONS[element_symbol]
        use_alternative = self.peak_selection_combo.currentText() == 'Alternative'
        
        if use_alternative and 'alternative_peak' in element_data:
            # Use alternative peak
            alt_peak = element_data['alternative_peak']
            self.element_energy_label.setText(f"{alt_peak['energy']} ({alt_peak['name']})")
            self.element_peak_region_label.setText(f"{alt_peak['peak_region'][0]} - {alt_peak['peak_region'][1]}")
            self.element_integration_region_label.setText(f"{alt_peak['integration_region'][0]} - {alt_peak['integration_region'][1]}")
        else:
            # Use primary peak
            self.element_energy_label.setText(f"{element_data['primary_energy']}")
            self.element_peak_region_label.setText(f"{element_data['peak_region'][0]} - {element_data['peak_region'][1]}")
            self.element_integration_region_label.setText(f"{element_data['integration_region'][0]} - {element_data['integration_region'][1]}")
        
        # Update calibration parameters
        if hasattr(self.peak_fitter, 'element_calibrations') and element_symbol in self.peak_fitter.element_calibrations:
            cal_data = self.peak_fitter.element_calibrations[element_symbol]
            self.element_slope_edit.setText(str(cal_data['slope']))
            self.element_intercept_edit.setText(str(cal_data['intercept']))
        else:
            # Use default calibration
            default_cal = element_data['default_calibration']
            self.element_slope_edit.setText(str(default_cal['slope']))
            self.element_intercept_edit.setText(str(default_cal['intercept']))
            
            # Highlight current element in reference materials table
            self.highlight_element_in_table(element_symbol)
    
    def highlight_element_in_table(self, element_symbol):
        """Highlight the selected element row in the reference materials table"""
        elements = list(ELEMENT_DEFINITIONS.keys())
        if element_symbol in elements:
            row_index = elements.index(element_symbol)
            self.ref_materials_table.selectRow(row_index)
    
    def setup_search_tab(self, tab):
        """Setup the element search/identification tab"""
        layout = QVBoxLayout(tab)
        
        # Instructions
        info_label = QLabel(
            "🔍 <b>Element Search</b><br>"
            "Load a spectrum file to automatically identify possible elements based on detected peaks. "
            "The tool will match peak energies against the XRF lines database (Na to U)."
        )
        info_label.setWordWrap(True)
        info_label.setStyleSheet("""
            QLabel {
                background-color: #E8F5E9;
                padding: 10px;
                border-radius: 5px;
                font-size: 11px;
            }
        """)
        layout.addWidget(info_label)
        
        # File loading section
        file_group = QGroupBox("Load Spectrum")
        file_layout = QHBoxLayout(file_group)
        
        self.search_file_label = QLabel("No file loaded")
        self.search_file_label.setStyleSheet("color: gray;")
        file_layout.addWidget(self.search_file_label)
        
        load_file_btn = QPushButton("📁 Load Spectrum File")
        load_file_btn.clicked.connect(self.load_spectrum_for_search)
        file_layout.addWidget(load_file_btn)
        
        self.send_to_quant_btn = QPushButton("➡️ Send to Quant Tab")
        self.send_to_quant_btn.setEnabled(False)
        self.send_to_quant_btn.clicked.connect(self.send_spectrum_to_quant)
        self.send_to_quant_btn.setToolTip("Send loaded spectrum to Quant tab for analysis")
        self.send_to_quant_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; }")
        file_layout.addWidget(self.send_to_quant_btn)
        
        layout.addWidget(file_group)
        
        # Peak detection parameters
        params_group = QGroupBox("Peak Detection Parameters")
        params_layout = QGridLayout(params_group)
        
        params_layout.addWidget(QLabel("Energy Tolerance (keV):"), 0, 0)
        self.search_energy_tolerance = QDoubleSpinBox()
        self.search_energy_tolerance.setRange(0.01, 1.0)
        self.search_energy_tolerance.setValue(0.1)
        self.search_energy_tolerance.setSingleStep(0.01)
        self.search_energy_tolerance.setDecimals(2)
        params_layout.addWidget(self.search_energy_tolerance, 0, 1)
        
        params_layout.addWidget(QLabel("Min Peak Height (cps):"), 1, 0)
        self.search_min_height = QSpinBox()
        self.search_min_height.setRange(10, 10000)
        self.search_min_height.setValue(100)
        self.search_min_height.setSingleStep(10)
        params_layout.addWidget(self.search_min_height, 1, 1)
        
        params_layout.addWidget(QLabel("Min Relative Intensity:"), 2, 0)
        self.search_min_rel_intensity = QSpinBox()
        self.search_min_rel_intensity.setRange(0, 100)
        self.search_min_rel_intensity.setValue(50)
        self.search_min_rel_intensity.setSuffix("%")
        params_layout.addWidget(self.search_min_rel_intensity, 2, 1)
        
        # Baseline subtraction options
        params_layout.addWidget(QLabel("Baseline Subtraction:"), 3, 0)
        self.search_baseline_method = QComboBox()
        self.search_baseline_method.addItems(["None", "Linear", "Polynomial", "SNIP"])
        self.search_baseline_method.setCurrentText("None")
        params_layout.addWidget(self.search_baseline_method, 3, 1)
        
        params_layout.addWidget(QLabel("Baseline Iterations:"), 4, 0)
        self.search_baseline_iterations = QSpinBox()
        self.search_baseline_iterations.setRange(1, 50)
        self.search_baseline_iterations.setValue(10)
        self.search_baseline_iterations.setEnabled(False)
        self.search_baseline_iterations.setToolTip("Number of iterations for SNIP algorithm")
        params_layout.addWidget(self.search_baseline_iterations, 4, 1)
        
        # Enable/disable iterations based on method
        self.search_baseline_method.currentTextChanged.connect(self.on_baseline_method_changed)
        
        # X-ray tube filter
        params_layout.addWidget(QLabel("X-ray Tube Element:"), 6, 0)
        self.search_tube_element = QComboBox()
        self.search_tube_element.addItems(["None", "Rh (Rhodium)", "W (Tungsten)", "Mo (Molybdenum)", "Cr (Chromium)", "Ag (Silver)"])
        self.search_tube_element.setCurrentText("Rh (Rhodium)")
        self.search_tube_element.setToolTip("Exclude X-ray tube element lines from search results")
        params_layout.addWidget(self.search_tube_element, 6, 1)
        
        search_btn = QPushButton("🔎 Search for Elements")
        search_btn.clicked.connect(self.search_for_elements)
        params_layout.addWidget(search_btn, 7, 0, 1, 2)
        
        layout.addWidget(params_group)
        
        # Detected elements list
        elements_group = QGroupBox("Detected Elements")
        elements_layout = QVBoxLayout(elements_group)
        
        self.detected_elements_table = QTableWidget(0, 7)
        self.detected_elements_table.setHorizontalHeaderLabels(['Element', 'Z', 'Lines', 'Energies (keV)', 'Peak Heights', 'Confidence', '# Matches'])
        self.detected_elements_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.detected_elements_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.detected_elements_table.itemSelectionChanged.connect(self.on_detected_element_selected)
        elements_layout.addWidget(self.detected_elements_table)
        
        layout.addWidget(elements_group)
        
        # Peak details table
        details_group = QGroupBox("Peak Details for Selected Element")
        details_layout = QVBoxLayout(details_group)
        
        self.peak_details_table = QTableWidget(0, 5)
        self.peak_details_table.setHorizontalHeaderLabels(['Line', 'Expected (keV)', 'Detected (keV)', 'Δ (keV)', 'Rel. Intensity'])
        self.peak_details_table.setMaximumHeight(150)
        details_layout.addWidget(self.peak_details_table)
        
        layout.addWidget(details_group)
        
        # Store search results
        self.search_spectrum_data = None
        self.search_detected_peaks = []
        self.search_element_matches = {}
        self.search_baseline_subtracted = None
    
    def on_baseline_method_changed(self):
        """Enable/disable baseline iterations based on selected method"""
        method = self.search_baseline_method.currentText()
        self.search_baseline_iterations.setEnabled(method == "SNIP")
    
    def apply_baseline_subtraction(self, energy, counts, method):
        """Apply baseline subtraction to spectrum data"""
        if method == "Linear":
            # Simple linear baseline from first to last point
            baseline = np.linspace(counts[0], counts[-1], len(counts))
            return counts - baseline
        
        elif method == "Polynomial":
            # Polynomial baseline (degree 3)
            from scipy.signal import savgol_filter
            # Smooth the data first
            smoothed = savgol_filter(counts, window_length=51, polyorder=3)
            # Use minimum values as baseline estimate
            baseline = np.minimum.accumulate(smoothed)
            baseline = np.minimum.accumulate(baseline[::-1])[::-1]
            return np.maximum(counts - baseline, 0)
        
        elif method == "SNIP":
            # Statistics-sensitive Non-linear Iterative Peak-clipping
            iterations = self.search_baseline_iterations.value()
            return self.snip_baseline(counts, iterations)
        
        return counts
    
    def snip_baseline(self, spectrum, iterations):
        """
        SNIP (Statistics-sensitive Non-linear Iterative Peak-clipping) algorithm
        for baseline estimation in XRF spectra
        """
        import numpy as np
        
        # Work with log-transformed spectrum
        spectrum = np.array(spectrum, dtype=float)
        # Avoid log(0) by adding small value
        spectrum_log = np.log(np.log(np.sqrt(spectrum + 1) + 1) + 1)
        
        # Apply SNIP algorithm
        working = np.copy(spectrum_log)
        
        for p in range(1, iterations + 1):
            for i in range(p, len(working) - p):
                a = working[i]
                b = (working[i - p] + working[i + p]) / 2.0
                if b < a:
                    working[i] = b
        
        # Transform back
        baseline = (np.exp(np.exp(working) - 1) - 1) ** 2 - 1
        
        # Ensure baseline doesn't exceed original spectrum
        baseline = np.minimum(baseline, spectrum)
        
        # Subtract baseline
        corrected = spectrum - baseline
        
        return np.maximum(corrected, 0)
    
    def load_spectrum_for_search(self):
        """Load a spectrum file for element search"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Spectrum File", "", "Text Files (*.txt);;All Files (*)"
        )
        
        if file_path:
            try:
                # Parse the spectrum file using existing parser
                energy, counts, format_type = parse_xrf_file_smart(file_path)
                
                if energy is None or counts is None:
                    raise Exception("Could not parse spectrum file")
                
                self.search_spectrum_data = {'energy': energy, 'counts': counts, 'file_path': file_path}
                
                # Update UI
                file_name = os.path.basename(file_path)
                self.search_file_label.setText(f"✓ {file_name} ({format_type})")
                self.search_file_label.setStyleSheet("color: green; font-weight: bold;")
                self.send_to_quant_btn.setEnabled(True)
                
                # Plot the spectrum on the main plot canvas
                self.plot_search_spectrum()
                
                QMessageBox.information(self, "Success", f"Spectrum loaded successfully!\n{len(energy)} data points\nFormat: {format_type}")
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load spectrum:\n{str(e)}")
    
    def send_spectrum_to_quant(self):
        """Send the loaded spectrum from Search tab to Quant tab"""
        if self.search_spectrum_data is None:
            QMessageBox.warning(self, "No Data", "No spectrum loaded to send.")
            return
        
        # Set the data in Quant tab
        file_path = self.search_spectrum_data['file_path']
        self.current_file_path = file_path
        self.current_data = (self.search_spectrum_data['energy'], self.search_spectrum_data['counts'])
        
        # Update Quant tab UI
        self.single_file_label.setText(os.path.basename(file_path))
        self.fit_single_btn.setEnabled(True)
        # pb_as_deconv_btn removed - deconvolution is now automatic
        
        # Load and display in Quant tab
        self.load_and_display_file(file_path)
        
        # Switch to Quant tab
        self.tab_widget.setCurrentIndex(1)  # Quant is tab index 1
        
        QMessageBox.information(self, "Success", 
                              f"Spectrum sent to Quant tab!\n\n"
                              f"File: {os.path.basename(file_path)}\n"
                              f"You can now fit peaks for selected elements.")
    
    def plot_search_spectrum(self, highlight_peaks=None):
        """Plot the loaded spectrum with optional peak highlights"""
        if self.search_spectrum_data is None:
            return
        
        # Use the main plot canvas on the right
        self.plot_canvas.figure.clear()
        ax = self.plot_canvas.figure.add_subplot(111)
        
        energy = self.search_spectrum_data['energy']
        counts = self.search_spectrum_data['counts']
        
        # Plot original spectrum
        ax.plot(energy, counts, 'b-', linewidth=0.8, alpha=0.5, label='Original')
        
        # Plot baseline-subtracted spectrum if available
        if self.search_baseline_subtracted is not None:
            ax.plot(energy, self.search_baseline_subtracted, 'g-', linewidth=0.8, alpha=0.7, label='Baseline Corrected')
            plot_data = self.search_baseline_subtracted
        else:
            plot_data = counts
        
        # Highlight detected peaks if available
        if highlight_peaks:
            for peak in highlight_peaks:
                ax.axvline(peak['energy'], color='red', linestyle='--', alpha=0.5, linewidth=1)
                ax.plot(peak['energy'], peak['height'], 'ro', markersize=8)
        
        ax.set_xlabel('Energy (keV)', fontsize=10)
        ax.set_ylabel('Counts', fontsize=10)
        ax.set_title('XRF Spectrum - Search Mode', fontsize=12, fontweight='bold')
        ax.grid(True, alpha=0.3)
        ax.legend()
        
        self.plot_canvas.figure.tight_layout()
        self.plot_canvas.draw()
    
    def search_for_elements(self):
        """Search for elements based on detected peaks in the spectrum"""
        if self.search_spectrum_data is None:
            QMessageBox.warning(self, "No Data", "Please load a spectrum file first.")
            return
        
        energy = self.search_spectrum_data['energy']
        counts = self.search_spectrum_data['counts']
        
        # Apply baseline subtraction if selected
        baseline_method = self.search_baseline_method.currentText()
        if baseline_method != "None":
            counts_corrected = self.apply_baseline_subtraction(energy, counts, baseline_method)
            self.search_baseline_subtracted = counts_corrected
        else:
            counts_corrected = counts
            self.search_baseline_subtracted = None
        
        # Detect peaks
        from scipy.signal import find_peaks
        
        min_height = self.search_min_height.value()
        peaks, properties = find_peaks(counts_corrected, height=min_height, distance=10, prominence=min_height*0.3)
        
        if len(peaks) == 0:
            QMessageBox.warning(self, "No Peaks", "No peaks detected. Try lowering the minimum height threshold.")
            return
        
        # Store detected peaks
        self.search_detected_peaks = []
        for i, peak_idx in enumerate(peaks):
            self.search_detected_peaks.append({
                'energy': energy[peak_idx],
                'height': counts_corrected[peak_idx],
                'index': peak_idx
            })
        
        # Match peaks to elements
        tolerance = self.search_energy_tolerance.value()
        min_rel_intensity = self.search_min_rel_intensity.value()
        
        # Get tube element to exclude
        tube_element_text = self.search_tube_element.currentText()
        if tube_element_text != "None":
            tube_element = tube_element_text.split()[0]  # Extract symbol (e.g., "Rh" from "Rh (Rhodium)")
        else:
            tube_element = None
        
        element_matches = {}
        
        for element, lines in XRF_LINES_DB.items():
            # Skip X-ray tube element
            if tube_element and element == tube_element:
                continue
            
            matched_lines = []
            
            for line in lines:
                # Only consider lines with sufficient relative intensity
                if line['Relative_Intensity'] < min_rel_intensity:
                    continue
                
                expected_energy = line['Energy_keV']
                
                # Check if any detected peak matches this line
                for peak in self.search_detected_peaks:
                    if abs(peak['energy'] - expected_energy) <= tolerance:
                        matched_lines.append({
                            'line': line['Line'],
                            'expected_energy': expected_energy,
                            'detected_energy': peak['energy'],
                            'delta': peak['energy'] - expected_energy,
                            'rel_intensity': line['Relative_Intensity'],
                            'peak_height': peak['height']
                        })
                        break
            
            if matched_lines:
                # Require at least 2 matched lines for confident identification
                # Exception: if only 1 line but it's very strong (>90% relative intensity)
                num_lines = len(matched_lines)
                
                if num_lines >= 2:
                    # Multiple lines matched - high confidence
                    # Base confidence on number of lines and their relative intensities
                    base_confidence = min(70, num_lines * 25)
                    intensity_bonus = sum(m['rel_intensity'] for m in matched_lines) / len(matched_lines) * 0.3
                    confidence = min(100, base_confidence + intensity_bonus)
                elif num_lines == 1 and matched_lines[0]['rel_intensity'] >= 90:
                    # Single very strong line - moderate confidence
                    confidence = 60
                else:
                    # Single weak line - skip this element
                    continue
                
                element_matches[element] = {
                    'atomic_number': lines[0]['Atomic_Number'],
                    'matched_lines': matched_lines,
                    'confidence': confidence,
                    'all_lines': lines  # Store all lines for plotting
                }
        
        self.search_element_matches = element_matches
        
        # Update results table
        self.update_detected_elements_table()
        
        # Plot spectrum with highlighted peaks
        self.plot_search_spectrum(highlight_peaks=self.search_detected_peaks)
        
        # Prepare result message
        result_msg = f"Found {len(peaks)} peaks in spectrum.\n"
        result_msg += f"Identified {len(element_matches)} possible elements."
        
        if tube_element:
            result_msg += f"\n\n✓ Filtered out {tube_element} (X-ray tube) lines"
        
        QMessageBox.information(self, "Search Complete", result_msg)
    
    def update_detected_elements_table(self):
        """Update the table showing detected elements"""
        self.detected_elements_table.setRowCount(0)
        
        # Sort by confidence
        sorted_elements = sorted(self.search_element_matches.items(), 
                                key=lambda x: x[1]['confidence'], reverse=True)
        
        for element, data in sorted_elements:
            row = self.detected_elements_table.rowCount()
            self.detected_elements_table.insertRow(row)
            
            # Element symbol
            elem_item = QTableWidgetItem(element)
            elem_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.detected_elements_table.setItem(row, 0, elem_item)
            
            # Atomic number
            z_item = QTableWidgetItem(str(data['atomic_number']))
            z_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.detected_elements_table.setItem(row, 1, z_item)
            
            # Emission lines (e.g., "Kα1, Kβ1")
            matched_lines = data['matched_lines']
            lines_str = ', '.join([line['line'] for line in matched_lines])
            lines_item = QTableWidgetItem(lines_str)
            self.detected_elements_table.setItem(row, 2, lines_item)
            
            # Energies
            energies_str = ', '.join([f"{line['detected_energy']:.3f}" for line in matched_lines])
            energies_item = QTableWidgetItem(energies_str)
            self.detected_elements_table.setItem(row, 3, energies_item)
            
            # Peak heights
            heights_str = ', '.join([f"{line['peak_height']:.0f}" for line in matched_lines])
            heights_item = QTableWidgetItem(heights_str)
            self.detected_elements_table.setItem(row, 4, heights_item)
            
            # Confidence
            conf_item = QTableWidgetItem(f"{data['confidence']:.0f}%")
            conf_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            
            # Color code by confidence
            if data['confidence'] >= 80:
                conf_item.setBackground(QColor(144, 238, 144))  # Light green
            elif data['confidence'] >= 50:
                conf_item.setBackground(QColor(255, 255, 200))  # Light yellow
            else:
                conf_item.setBackground(QColor(255, 200, 200))  # Light red
            
            self.detected_elements_table.setItem(row, 5, conf_item)
            
            # Number of matched lines
            num_lines_item = QTableWidgetItem(str(len(matched_lines)))
            num_lines_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.detected_elements_table.setItem(row, 6, num_lines_item)
        
        self.detected_elements_table.resizeColumnsToContents()
    
    def on_detected_element_selected(self):
        """Handle selection of a detected element"""
        selected_rows = self.detected_elements_table.selectedItems()
        if not selected_rows:
            return
        
        row = selected_rows[0].row()
        element = self.detected_elements_table.item(row, 0).text()
        
        if element in self.search_element_matches:
            self.update_peak_details_table(element)
            self.plot_spectrum_with_element_lines(element)
    
    def update_peak_details_table(self, element):
        """Update the peak details table for selected element"""
        self.peak_details_table.setRowCount(0)
        
        if element not in self.search_element_matches:
            return
        
        matched_lines = self.search_element_matches[element]['matched_lines']
        
        for line_data in matched_lines:
            row = self.peak_details_table.rowCount()
            self.peak_details_table.insertRow(row)
            
            # Line type
            self.peak_details_table.setItem(row, 0, QTableWidgetItem(line_data['line']))
            
            # Expected energy
            exp_item = QTableWidgetItem(f"{line_data['expected_energy']:.4f}")
            exp_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.peak_details_table.setItem(row, 1, exp_item)
            
            # Detected energy
            det_item = QTableWidgetItem(f"{line_data['detected_energy']:.4f}")
            det_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.peak_details_table.setItem(row, 2, det_item)
            
            # Delta
            delta_item = QTableWidgetItem(f"{line_data['delta']:.4f}")
            delta_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.peak_details_table.setItem(row, 3, delta_item)
            
            # Relative intensity
            rel_item = QTableWidgetItem(f"{line_data['rel_intensity']}")
            rel_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.peak_details_table.setItem(row, 4, rel_item)
        
        self.peak_details_table.resizeColumnsToContents()
    
    def plot_spectrum_with_element_lines(self, element):
        """Plot spectrum with ALL characteristic lines for the selected element"""
        if self.search_spectrum_data is None or element not in self.search_element_matches:
            return
        
        # Use the main plot canvas on the right
        self.plot_canvas.figure.clear()
        ax = self.plot_canvas.figure.add_subplot(111)
        
        energy = self.search_spectrum_data['energy']
        counts = self.search_spectrum_data['counts']
        
        # Plot original spectrum
        ax.plot(energy, counts, 'b-', linewidth=0.8, alpha=0.5, label='Original')
        
        # Plot baseline-subtracted spectrum if available
        if self.search_baseline_subtracted is not None:
            ax.plot(energy, self.search_baseline_subtracted, 'g-', linewidth=0.8, alpha=0.7, label='Baseline Corrected')
            plot_data = self.search_baseline_subtracted
        else:
            plot_data = counts
        
        # Get all lines for this element
        all_lines = self.search_element_matches[element]['all_lines']
        matched_lines = self.search_element_matches[element]['matched_lines']
        matched_energies = {line['expected_energy'] for line in matched_lines}
        
        # Plot ALL characteristic lines for this element
        for line in all_lines:
            expected_e = line['Energy_keV']
            rel_intensity = line['Relative_Intensity']
            line_name = line['Line']
            
            # Check if this line was matched to a detected peak
            is_matched = expected_e in matched_energies
            
            # Find the y-position for the line marker
            energy_idx = np.argmin(np.abs(energy - expected_e))
            if energy_idx < len(plot_data):
                y_pos = plot_data[energy_idx]
            else:
                y_pos = np.max(plot_data) * 0.9
            
            if is_matched:
                # Matched line - solid red with marker
                ax.axvline(expected_e, color='red', linestyle='-', alpha=0.7, linewidth=2)
                ax.plot(expected_e, y_pos, 'ro', markersize=10)
                ax.annotate(f"{element} {line_name}\n{expected_e:.3f} keV\n✓ Detected", 
                           xy=(expected_e, y_pos), xytext=(10, 10),
                           textcoords='offset points', fontsize=7,
                           bbox=dict(boxstyle='round,pad=0.3', facecolor='lightgreen', alpha=0.8),
                           arrowprops=dict(arrowstyle='->', connectionstyle='arc3,rad=0'))
            else:
                # Expected but not detected - dashed gray
                ax.axvline(expected_e, color='gray', linestyle='--', alpha=0.4, linewidth=1)
                # Only annotate strong lines that weren't detected
                if rel_intensity >= 50:
                    ax.annotate(f"{element} {line_name}\n{expected_e:.3f} keV\n(expected)", 
                               xy=(expected_e, y_pos * 0.5), xytext=(5, -20),
                               textcoords='offset points', fontsize=6,
                               bbox=dict(boxstyle='round,pad=0.2', facecolor='lightgray', alpha=0.6),
                               alpha=0.7)
        
        # Add legend entries
        from matplotlib.lines import Line2D
        legend_elements = [
            Line2D([0], [0], color='red', linewidth=2, label='Detected Lines'),
            Line2D([0], [0], color='gray', linestyle='--', linewidth=1, label='Expected Lines (not detected)')
        ]
        
        ax.set_xlabel('Energy (keV)', fontsize=10)
        ax.set_ylabel('Counts', fontsize=10)
        ax.set_title(f'XRF Spectrum - All {element} Characteristic Lines', fontsize=12, fontweight='bold')
        ax.grid(True, alpha=0.3)
        ax.legend(handles=legend_elements, loc='best', fontsize=8)
        
        self.plot_canvas.figure.tight_layout()
        self.plot_canvas.draw()
    
    def setup_calibration_plots_tab(self, tab):
        """Setup the calibration plots visualization tab - controls only, plots go to right pane"""
        layout = QVBoxLayout(tab)
        
        # Info label
        info_label = QLabel(
            "📊 <b>Calibration Plots Viewer</b><br>"
            "View calibration curves by element or standard. Plots will appear in the right pane."
        )
        info_label.setWordWrap(True)
        info_label.setStyleSheet("""
            QLabel {
                background-color: #E3F2FD;
                padding: 10px;
                border-radius: 5px;
                font-size: 11px;
            }
        """)
        layout.addWidget(info_label)
        
        # View mode selector
        mode_group = QGroupBox("View Mode")
        mode_layout = QVBoxLayout(mode_group)
        
        self.cal_view_by_element_radio = QRadioButton("View by Element")
        self.cal_view_by_standard_radio = QRadioButton("View by Standard")
        self.cal_view_by_element_radio.setChecked(True)
        self.cal_view_by_element_radio.toggled.connect(self.update_calibration_view_mode)
        
        mode_layout.addWidget(self.cal_view_by_element_radio)
        mode_layout.addWidget(self.cal_view_by_standard_radio)
        layout.addWidget(mode_group)
        
        # Elements list (for "View by Element" mode)
        self.cal_elements_group = QGroupBox("Select Element")
        elements_layout = QVBoxLayout(self.cal_elements_group)
        
        self.cal_element_list = QListWidget()
        for symbol, data in ELEMENT_DEFINITIONS.items():
            self.cal_element_list.addItem(f"{symbol} ({data['name']})")
        self.cal_element_list.currentItemChanged.connect(self.on_calibration_element_selected)
        elements_layout.addWidget(self.cal_element_list)
        
        layout.addWidget(self.cal_elements_group)
        
        # Standards list (for "View by Standard" mode)
        self.cal_standards_group = QGroupBox("Select Standard")
        standards_layout = QVBoxLayout(self.cal_standards_group)
        
        self.cal_standard_list = QListWidget()
        for std_name in REFERENCE_MATERIALS.keys():
            self.cal_standard_list.addItem(std_name)
        self.cal_standard_list.currentItemChanged.connect(self.on_calibration_standard_selected)
        standards_layout.addWidget(self.cal_standard_list)
        
        layout.addWidget(self.cal_standards_group)
        self.cal_standards_group.hide()  # Hidden by default
        
        # Refresh button
        refresh_btn = QPushButton("🔄 Refresh Plots")
        refresh_btn.clicked.connect(self.refresh_calibration_plots)
        layout.addWidget(refresh_btn)
        
        # Info display for selected calibration
        self.cal_info_label = QLabel("Select an element or standard to view calibration details")
        self.cal_info_label.setWordWrap(True)
        self.cal_info_label.setStyleSheet("""
            QLabel {
                background-color: #E8F5E9;
                padding: 8px;
                border-radius: 5px;
                font-size: 10px;
            }
        """)
        layout.addWidget(self.cal_info_label)
        
        # Statistics table
        stats_group = QGroupBox("Calibration Statistics")
        stats_layout = QVBoxLayout(stats_group)
        
        self.cal_stats_table = QTableWidget(0, 4)
        self.cal_stats_table.setHorizontalHeaderLabels(['Standard', 'Certified (ppm)', 'Measured Intensity (cps)', 'Predicted (ppm)'])
        self.cal_stats_table.setMaximumHeight(200)
        stats_layout.addWidget(self.cal_stats_table)
        
        layout.addWidget(stats_group)
        layout.addStretch()
    
    def update_calibration_plot(self):
        """Legacy method - redirects to new refresh method"""
        self.refresh_calibration_plots()
    
    def update_calibration_view_mode(self):
        """Switch between viewing by element or by standard"""
        if self.cal_view_by_element_radio.isChecked():
            self.cal_elements_group.show()
            self.cal_standards_group.hide()
            # Trigger element view if one is selected
            if self.cal_element_list.currentItem():
                self.on_calibration_element_selected(self.cal_element_list.currentItem(), None)
        else:
            self.cal_elements_group.hide()
            self.cal_standards_group.show()
            # Trigger standard view if one is selected
            if self.cal_standard_list.currentItem():
                self.on_calibration_standard_selected(self.cal_standard_list.currentItem(), None)
    
    def on_calibration_element_selected(self, current, previous):
        """Handle element selection in calibration plots tab"""
        if current:
            element_text = current.text()
            element_symbol = element_text.split()[0]
            self.display_single_element_calibration(element_symbol)
    
    def on_calibration_standard_selected(self, current, previous):
        """Handle standard selection in calibration plots tab"""
        if current:
            standard_name = current.text()
            self.display_standard_all_elements(standard_name)
    
    def refresh_calibration_plots(self):
        """Refresh the current calibration plot view"""
        if self.cal_view_by_element_radio.isChecked():
            if self.cal_element_list.currentItem():
                self.on_calibration_element_selected(self.cal_element_list.currentItem(), None)
        else:
            if self.cal_standard_list.currentItem():
                self.on_calibration_standard_selected(self.cal_standard_list.currentItem(), None)
    
    def display_single_element_calibration(self, element_symbol):
        """Display calibration plot for a single element (original behavior)"""
        # Check if calibration exists
        if not hasattr(self.calibration_manager, 'calibrations'):
            self.plot_canvas.figure.clear()
            self.plot_canvas.draw()
            self.cal_stats_table.setRowCount(0)
            return
        
        if element_symbol not in self.calibration_manager.calibrations:
            self.plot_canvas.figure.clear()
            self.plot_canvas.draw()
            self.cal_stats_table.setRowCount(0)
            return
        
        # Get calibration data
        cal_data = self.calibration_manager.calibrations[element_symbol]
        slope = cal_data.get('slope', 0)
        intercept = cal_data.get('intercept', 0)
        r_squared = cal_data.get('r_squared', 0)
        standards_used = cal_data.get('standards_used', [])
        timestamp = cal_data.get('timestamp', 'Unknown')
        
        # Update info label
        element_name = ELEMENT_DEFINITIONS[element_symbol]['name']
        r2_warning = ""
        if r_squared < 0.5:
            r2_warning = " ❌ POOR CALIBRATION - R² too low!"
        elif r_squared < 0.90:
            r2_warning = " ⚠️ Warning: Low R² - check data quality"
        elif r_squared < 0.95:
            r2_warning = " ⚠️ Acceptable but could be improved"
        
        intercept_note = ""
        if intercept < 0:
            intercept_note = " (Negative intercept: background slightly overcorrected - normal)"
        elif intercept > 100:
            intercept_note = f" ⚠️ (Large positive intercept: check for contamination or systematic bias)"
        
        self.cal_info_label.setText(
            f"📊 {element_name} ({element_symbol}) Calibration{r2_warning}\n"
            f"Equation: Concentration = {slope:.6f} × Intensity + {intercept:.2f}{intercept_note}\n"
            f"R² = {r_squared:.6f} | Standards: {', '.join(standards_used) if standards_used else 'N/A'} | "
            f"Created: {timestamp}"
        )
        
        # Set background color based on R² quality
        if r_squared < 0.5:
            bg_color = "#FFCCCC"
            border = "2px solid red"
        elif r_squared < 0.90:
            bg_color = "#FFF4CC"
            border = "2px solid orange"
        else:
            bg_color = "#E3F2FD"
            border = "none"
        
        self.cal_info_label.setStyleSheet(f"""
            QLabel {{
                background-color: {bg_color};
                padding: 10px;
                border-radius: 5px;
                font-size: 11px;
                border: {border};
            }}
        """)
        
        # Plot the calibration curve (reuse existing plotting logic)
        self._plot_element_calibration_curve(element_symbol, cal_data)
    
    def display_standard_all_elements(self, standard_name):
        """Display all calibrated element plots for a selected standard"""
        # Check if calibrations exist
        if not hasattr(self.calibration_manager, 'calibrations'):
            self.plot_canvas.figure.clear()
            self.plot_canvas.draw()
            self.cal_stats_table.setRowCount(0)
            return
        
        # Find which elements have calibrations that use this standard
        elements_with_standard = []
        for element_symbol, cal_data in self.calibration_manager.calibrations.items():
            standards_used = cal_data.get('standards_used', [])
            if standard_name in standards_used:
                elements_with_standard.append(element_symbol)
        
        if not elements_with_standard:
            self.plot_canvas.figure.clear()
            self.plot_canvas.draw()
            self.cal_stats_table.setRowCount(0)
            return
        
        # Update info label
        self.cal_info_label.setText(
            f"📊 Standard: {standard_name}\n"
            f"Showing calibration curves for {len(elements_with_standard)} element(s): {', '.join(elements_with_standard)}"
        )
        self.cal_info_label.setStyleSheet("""
            QLabel {
                background-color: #E3F2FD;
                padding: 10px;
                border-radius: 5px;
                font-size: 11px;
            }
        """)
        
        # Create subplot grid
        self.plot_canvas.figure.clear()
        n_plots = len(elements_with_standard)
        
        # Determine grid layout
        if n_plots == 1:
            rows, cols = 1, 1
        elif n_plots == 2:
            rows, cols = 1, 2
        elif n_plots <= 4:
            rows, cols = 2, 2
        elif n_plots <= 6:
            rows, cols = 2, 3
        elif n_plots <= 9:
            rows, cols = 3, 3
        else:
            rows, cols = 4, 3
        
        # Plot each element
        stats_data = []
        for idx, element_symbol in enumerate(sorted(elements_with_standard)):
            ax = self.plot_canvas.figure.add_subplot(rows, cols, idx + 1)
            cal_data = self.calibration_manager.calibrations[element_symbol]
            
            # Plot this element's calibration with the standard highlighted
            self._plot_element_in_subplot(ax, element_symbol, cal_data, standard_name, stats_data)
        
        self.plot_canvas.figure.tight_layout()
        self.plot_canvas.draw()
        
        # Update statistics table with data from all elements
        self._update_stats_table_multi_element(stats_data)
    
    def _plot_element_calibration_curve(self, element_symbol, cal_data):
        """Helper method to plot a single element's calibration curve"""
        slope = cal_data.get('slope', 0)
        intercept = cal_data.get('intercept', 0)
        r_squared = cal_data.get('r_squared', 0)
        standards_used = cal_data.get('standards_used', [])
        
        # Collect standard data points
        standard_intensities = []
        standard_concentrations = []
        standard_names = []
        
        for std_name in standards_used:
            cert_conc = self._get_certified_concentration(std_name, element_symbol)
            if cert_conc is not None and slope != 0:
                intensity = (cert_conc - intercept) / slope
                standard_intensities.append(intensity)
                standard_concentrations.append(cert_conc)
                standard_names.append(std_name)
        
        # Create the plot
        self.plot_canvas.figure.clear()
        ax = self.plot_canvas.figure.add_subplot(111)
        
        if standard_intensities:
            raw_intensities_dict = cal_data.get('raw_intensities', {})
            
            # Color palette
            standard_colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', 
                             '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf']
            
            # Plot individual measurements
            if raw_intensities_dict:
                for idx, std_name in enumerate(standards_used):
                    if std_name in raw_intensities_dict and raw_intensities_dict[std_name]:
                        cert_conc = self._get_certified_concentration(std_name, element_symbol)
                        if cert_conc is not None:
                            raw_ints = raw_intensities_dict[std_name]
                            raw_concs = [cert_conc] * len(raw_ints)
                            color = standard_colors[idx % len(standard_colors)]
                            ax.scatter(raw_ints, raw_concs, color=color, s=40, alpha=0.6, 
                                     zorder=3, label=f'{std_name} (n={len(raw_ints)})')
            
            # Plot calibration line
            max_intensity = max(standard_intensities) * 1.2
            if raw_intensities_dict:
                all_raw = [val for vals in raw_intensities_dict.values() for val in vals]
                if all_raw:
                    max_intensity = max(max_intensity, max(all_raw) * 1.2)
            
            intensity_range = np.linspace(0, max_intensity, 100)
            concentration_range = slope * intensity_range + intercept
            
            ax.plot(intensity_range, concentration_range, 'k-', linewidth=2.5, 
                   label=f'Calibration: y = {slope:.4f}x + {intercept:.2f}', zorder=4)
            
            # Plot averaged standard points
            for idx, (intensity, conc, name) in enumerate(zip(standard_intensities, standard_concentrations, standard_names)):
                color = standard_colors[idx % len(standard_colors)]
                ax.scatter([intensity], [conc], color=color, s=150, alpha=0.9, zorder=5, 
                          edgecolors='black', linewidths=2.5, marker='o')
            
            # Add labels
            for i, name in enumerate(standard_names):
                ax.annotate(name, (standard_intensities[i], standard_concentrations[i]),
                           xytext=(5, 5), textcoords='offset points', fontsize=9, fontweight='bold',
                           bbox=dict(boxstyle='round,pad=0.3', facecolor='white', alpha=0.7, edgecolor='gray'))
            
            # Formatting
            element_name = ELEMENT_DEFINITIONS[element_symbol]['name']
            ax.set_xlabel('Integrated Intensity (cps)', fontsize=10)
            ax.set_ylabel(f'{element_symbol} Concentration (ppm)', fontsize=10)
            ax.set_title(f'{element_name} ({element_symbol}) Calibration Curve\nR² = {r_squared:.6f}', 
                        fontsize=12, fontweight='bold')
            ax.legend(loc='best')
            ax.grid(True, alpha=0.3)
            ax.set_xlim(0, max_intensity)
            ax.set_ylim(0, max(standard_concentrations) * 1.2)
            
            # Update statistics table
            self.cal_stats_table.setRowCount(len(standard_names))
            for i, name in enumerate(standard_names):
                self.cal_stats_table.setItem(i, 0, QTableWidgetItem(name))
                
                cert_item = QTableWidgetItem(f"{standard_concentrations[i]:.1f}")
                cert_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                self.cal_stats_table.setItem(i, 1, cert_item)
                
                intensity_item = QTableWidgetItem(f"{standard_intensities[i]:.1f}")
                intensity_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                self.cal_stats_table.setItem(i, 2, intensity_item)
                
                predicted = slope * standard_intensities[i] + intercept
                pred_item = QTableWidgetItem(f"{predicted:.1f}")
                pred_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                
                error_pct = abs((predicted - standard_concentrations[i]) / standard_concentrations[i] * 100)
                if error_pct < 2:
                    pred_item.setBackground(QColor(144, 238, 144))
                elif error_pct < 5:
                    pred_item.setBackground(QColor(255, 255, 200))
                else:
                    pred_item.setBackground(QColor(255, 200, 200))
                
                self.cal_stats_table.setItem(i, 3, pred_item)
            
            self.cal_stats_table.resizeColumnsToContents()
        else:
            ax.text(0.5, 0.5, 'No standard data available', 
                   ha='center', va='center', transform=ax.transAxes, fontsize=14)
        
        self.plot_canvas.figure.tight_layout()
        self.plot_canvas.draw()
    
    def _plot_element_in_subplot(self, ax, element_symbol, cal_data, highlight_standard, stats_data):
        """Helper method to plot an element's calibration in a subplot"""
        slope = cal_data.get('slope', 0)
        intercept = cal_data.get('intercept', 0)
        r_squared = cal_data.get('r_squared', 0)
        standards_used = cal_data.get('standards_used', [])
        raw_intensities_dict = cal_data.get('raw_intensities', {})
        
        # Collect standard data
        standard_intensities = []
        standard_concentrations = []
        standard_names = []
        
        for std_name in standards_used:
            cert_conc = self._get_certified_concentration(std_name, element_symbol)
            if cert_conc is not None and slope != 0:
                intensity = (cert_conc - intercept) / slope
                standard_intensities.append(intensity)
                standard_concentrations.append(cert_conc)
                standard_names.append(std_name)
        
        if not standard_intensities:
            ax.text(0.5, 0.5, 'No data', ha='center', va='center', transform=ax.transAxes)
            ax.set_title(f'{element_symbol}', fontsize=10)
            return
        
        # Plot individual measurements (smaller, semi-transparent)
        if raw_intensities_dict:
            for std_name in standards_used:
                if std_name in raw_intensities_dict and raw_intensities_dict[std_name]:
                    cert_conc = self._get_certified_concentration(std_name, element_symbol)
                    if cert_conc is not None:
                        raw_ints = raw_intensities_dict[std_name]
                        raw_concs = [cert_conc] * len(raw_ints)
                        # Highlight the selected standard
                        if std_name == highlight_standard:
                            ax.scatter(raw_ints, raw_concs, color='red', s=60, alpha=0.8, 
                                     zorder=4, marker='o', edgecolors='darkred', linewidths=1.5)
                        else:
                            ax.scatter(raw_ints, raw_concs, color='gray', s=30, alpha=0.4, zorder=2)
        
        # Plot calibration line
        max_intensity = max(standard_intensities) * 1.2
        if raw_intensities_dict:
            all_raw = [val for vals in raw_intensities_dict.values() for val in vals]
            if all_raw:
                max_intensity = max(max_intensity, max(all_raw) * 1.2)
        
        intensity_range = np.linspace(0, max_intensity, 100)
        concentration_range = slope * intensity_range + intercept
        ax.plot(intensity_range, concentration_range, 'k-', linewidth=1.5, alpha=0.7, zorder=3)
        
        # Plot averaged standard points
        for intensity, conc, name in zip(standard_intensities, standard_concentrations, standard_names):
            if name == highlight_standard:
                ax.scatter([intensity], [conc], color='red', s=120, alpha=1.0, zorder=5, 
                          edgecolors='darkred', linewidths=2, marker='o')
                # Add label for highlighted standard
                ax.annotate(name, (intensity, conc), xytext=(5, 5), textcoords='offset points',
                           fontsize=8, fontweight='bold', color='darkred')
                
                # Add to stats data
                predicted = slope * intensity + intercept
                stats_data.append({
                    'element': element_symbol,
                    'standard': name,
                    'certified': conc,
                    'intensity': intensity,
                    'predicted': predicted
                })
            else:
                ax.scatter([intensity], [conc], color='gray', s=60, alpha=0.5, zorder=3, 
                          edgecolors='black', linewidths=1)
        
        # Formatting
        element_name = ELEMENT_DEFINITIONS[element_symbol]['name']
        ax.set_xlabel('Intensity (cps)', fontsize=8)
        ax.set_ylabel('Conc. (ppm)', fontsize=8)
        ax.set_title(f'{element_symbol} - R²={r_squared:.3f}', fontsize=9, fontweight='bold')
        ax.grid(True, alpha=0.2)
        ax.set_xlim(0, max_intensity)
        ax.set_ylim(0, max(standard_concentrations) * 1.2)
        ax.tick_params(labelsize=7)
    
    def _get_certified_concentration(self, std_name, element_symbol):
        """Helper method to get certified concentration for a standard and element"""
        if std_name in REFERENCE_MATERIALS:
            cert_value = REFERENCE_MATERIALS[std_name].get(element_symbol)
            if cert_value and cert_value != "N/A":
                try:
                    if isinstance(cert_value, str):
                        if '%' in cert_value:
                            if '<' in cert_value:
                                return None
                            return float(cert_value.replace('%', '')) * 10000
                        elif '<' in cert_value:
                            return None
                        else:
                            return float(cert_value)
                    else:
                        return float(cert_value)
                except (ValueError, TypeError):
                    return None
        elif hasattr(self, 'custom_standards_data') and std_name in self.custom_standards_data:
            if element_symbol in self.custom_standards_data[std_name]:
                return self.custom_standards_data[std_name][element_symbol]
        return None
    
    def _update_stats_table_multi_element(self, stats_data):
        """Update statistics table for multi-element view"""
        self.cal_stats_table.setRowCount(len(stats_data))
        
        for i, data in enumerate(stats_data):
            # Element + Standard name
            self.cal_stats_table.setItem(i, 0, QTableWidgetItem(f"{data['element']} - {data['standard']}"))
            
            # Certified concentration
            cert_item = QTableWidgetItem(f"{data['certified']:.1f}")
            cert_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.cal_stats_table.setItem(i, 1, cert_item)
            
            # Measured intensity
            intensity_item = QTableWidgetItem(f"{data['intensity']:.1f}")
            intensity_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.cal_stats_table.setItem(i, 2, intensity_item)
            
            # Predicted concentration
            pred_item = QTableWidgetItem(f"{data['predicted']:.1f}")
            pred_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            
            # Color code based on accuracy
            error_pct = abs((data['predicted'] - data['certified']) / data['certified'] * 100)
            if error_pct < 2:
                pred_item.setBackground(QColor(144, 238, 144))
            elif error_pct < 5:
                pred_item.setBackground(QColor(255, 255, 200))
            else:
                pred_item.setBackground(QColor(255, 200, 200))
            
            self.cal_stats_table.setItem(i, 3, pred_item)
        
        self.cal_stats_table.resizeColumnsToContents()
    
    def update_fitting_parameters_for_element(self, element_symbol):
        """Update the fitting parameters in the main UI for the selected element"""
        if element_symbol in ELEMENT_DEFINITIONS:
            element_data = ELEMENT_DEFINITIONS[element_symbol]
            
            # Note: Fitting parameters UI removed - now managed via Multi-Element Calibrations tab
            # Peak and integration regions are still used internally by the fitter
            
            # Update window title to reflect current element
            element_name = element_data['name']
            energy = element_data['primary_energy']
            self.setWindowTitle(f"XRF Peak Fitting - {element_name} ({element_symbol}) at {energy} keV")
    
    def update_element_calibration(self):
        """Update calibration parameters for the current element"""
        try:
            current_element = self.element_combo.currentText()
            slope = float(self.element_slope_edit.text())
            intercept = float(self.element_intercept_edit.text())
            
            # Update the peak fitter
            self.peak_fitter.update_element_calibration(current_element, slope, intercept)
            self.fitter.update_element_calibration(current_element, slope, intercept)
            
            # Save calibration persistently
            self.calibration_manager.update_calibration(
                current_element, slope, intercept, None, ["Manual Entry"]
            )
            
            # Refresh status display
            self.refresh_calibration_status()
            
            # Show confirmation
            QMessageBox.information(self, "Calibration Updated", 
                                  f"Calibration for {current_element} has been updated and saved:\n"
                                  f"Slope: {slope:.4f}\n"
                                  f"Intercept: {intercept:.4f}")
            
        except ValueError:
            QMessageBox.warning(self, "Invalid Input", 
                              "Please enter valid numeric values for slope and intercept.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to update calibration: {str(e)}")
    
    def create_calibration_from_standards(self):
        """Create calibration curve from reference material standards"""
        current_element = self.element_combo.currentText()
        
        # Collect available standards for this element
        available_standards = []
        concentrations = []
        
        for material_name, material_data in REFERENCE_MATERIALS.items():
            value = material_data.get(current_element)
            if value is not None and value != "N/A":
                # Parse concentration value
                try:
                    if isinstance(value, str):
                        # Handle percentage values
                        if '%' in value:
                            if '<' in value:
                                continue  # Skip below detection limit values
                            conc = float(value.replace('%', '')) * 10000  # Convert % to ppm
                        else:
                            conc = float(value)
                    else:
                        conc = float(value)
                    
                    available_standards.append(material_name)
                    concentrations.append(conc)
                except (ValueError, TypeError):
                    continue
        
        if len(available_standards) < 2:
            QMessageBox.warning(self, "Insufficient Standards", 
                              f"At least 2 reference materials with certified values are needed for {current_element}.\n"
                              f"Only {len(available_standards)} standards available.")
            return
        
        # Show dialog to create calibration
        self.show_calibration_creation_dialog(current_element, available_standards, concentrations)
    
    def show_calibration_creation_dialog(self, element, standards, concentrations):
        """Show dialog for creating calibration from standards"""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Create {element} Calibration from Standards")
        dialog.setGeometry(300, 300, 600, 400)
        
        layout = QVBoxLayout(dialog)
        
        # Instructions
        info_label = QLabel(f"To create a calibration for {element}, you need to:\n"
                           f"1. Measure XRF spectra of the reference materials\n"
                           f"2. Enter the integrated intensities for each standard\n"
                           f"3. The system will calculate the calibration curve")
        info_label.setWordWrap(True)
        layout.addWidget(info_label)
        
        # Standards table
        table = QTableWidget(len(standards), 3)
        table.setHorizontalHeaderLabels(['Standard', 'Certified Conc. (ppm)', 'Measured Intensity'])
        
        intensity_edits = []
        for i, (standard, conc) in enumerate(zip(standards, concentrations)):
            # Standard name
            table.setItem(i, 0, QTableWidgetItem(standard))
            table.item(i, 0).setFlags(table.item(i, 0).flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            # Certified concentration
            table.setItem(i, 1, QTableWidgetItem(f"{conc:.2f}"))
            table.item(i, 1).setFlags(table.item(i, 1).flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            # Intensity input
            intensity_edit = QLineEdit()
            intensity_edit.setPlaceholderText("Enter measured intensity...")
            table.setCellWidget(i, 2, intensity_edit)
            intensity_edits.append(intensity_edit)
        
        table.resizeColumnsToContents()
        layout.addWidget(table)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        calculate_btn = QPushButton("Calculate Calibration")
        cancel_btn = QPushButton("Cancel")
        
        def calculate_calibration():
            try:
                # Get intensities
                intensities = []
                for edit in intensity_edits:
                    if edit.text().strip():
                        intensities.append(float(edit.text()))
                    else:
                        QMessageBox.warning(dialog, "Missing Data", "Please enter all intensity values.")
                        return
                
                if len(intensities) != len(concentrations):
                    QMessageBox.warning(dialog, "Data Mismatch", "Number of intensities doesn't match standards.")
                    return
                
                # Calculate linear regression
                slope, intercept, r_value, p_value, std_err = stats.linregress(intensities, concentrations)
                
                # Show results
                result_msg = (f"Calibration Results for {element}:\n\n"
                             f"Equation: Concentration = {slope:.4f} × Intensity + {intercept:.4f}\n"
                             f"R² = {r_value**2:.4f}\n"
                             f"Standard Error = {std_err:.4f}\n\n"
                             f"Apply this calibration?")
                
                reply = QMessageBox.question(dialog, "Calibration Results", result_msg,
                                           QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                
                if reply == QMessageBox.StandardButton.Yes:
                    # Update calibration
                    self.peak_fitter.update_element_calibration(element, slope, intercept)
                    self.element_slope_edit.setText(f"{slope:.4f}")
                    self.element_intercept_edit.setText(f"{intercept:.4f}")
                    
                    QMessageBox.information(dialog, "Success", f"Calibration for {element} has been updated!")
                    dialog.accept()
                
            except ValueError:
                QMessageBox.warning(dialog, "Invalid Input", "Please enter valid numeric values for intensities.")
            except Exception as e:
                QMessageBox.critical(dialog, "Error", f"Failed to calculate calibration: {str(e)}")
        
        calculate_btn.clicked.connect(calculate_calibration)
        cancel_btn.clicked.connect(dialog.reject)
        
        button_layout.addWidget(calculate_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        dialog.exec()
    
    def view_standards_plot(self):
        """Show a plot of available standards for the current element"""
        current_element = self.element_combo.currentText()
        
        # Collect data for plotting
        materials = []
        concentrations = []
        
        for material_name, material_data in REFERENCE_MATERIALS.items():
            value = material_data.get(current_element)
            if value is not None and value != "N/A":
                try:
                    if isinstance(value, str):
                        if '%' in value:
                            if '<' in value:
                                continue
                            conc = float(value.replace('%', '')) * 10000  # Convert % to ppm
                        else:
                            conc = float(value)
                    else:
                        conc = float(value)
                    
                    materials.append(material_name)
                    concentrations.append(conc)
                except (ValueError, TypeError):
                    continue
        
        if not materials:
            QMessageBox.information(self, "No Data", f"No certified values available for {current_element}.")
            return
        
        # Create plot
        try:
            import matplotlib.pyplot as plt
            
            fig, ax = plt.subplots(figsize=(10, 6))
            
            # Bar plot of concentrations
            bars = ax.bar(range(len(materials)), concentrations, color='steelblue', alpha=0.7)
            
            # Customize plot
            ax.set_xlabel('Reference Material')
            ax.set_ylabel(f'{current_element} Concentration (ppm)')
            ax.set_title(f'Certified {current_element} Concentrations in Reference Materials')
            ax.set_xticks(range(len(materials)))
            ax.set_xticklabels(materials, rotation=45, ha='right')
            
            # Add value labels on bars
            for bar, conc in zip(bars, concentrations):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + max(concentrations)*0.01,
                       f'{conc:.1f}', ha='center', va='bottom')
            
            ax.grid(True, alpha=0.3)
            plt.tight_layout()
            plt.show()
            
        except ImportError:
            QMessageBox.warning(self, "Plot Error", "Matplotlib is required for plotting.")
        except Exception as e:
            QMessageBox.critical(self, "Plot Error", f"Failed to create plot: {str(e)}")

    def auto_calibrate_from_spectra(self):
        """Automatically create calibration by analyzing XRF spectra files of reference materials"""
        current_element = self.element_combo.currentText()
        
        # Collect available standards for this element
        available_standards = []
        concentrations = []
        
        for material_name, material_data in REFERENCE_MATERIALS.items():
            value = material_data.get(current_element)
            if value is not None and value != "N/A":
                try:
                    if isinstance(value, str):
                        if '%' in value:
                            if '<' in value:
                                continue  # Skip below detection limit values
                            conc = float(value.replace('%', '')) * 10000  # Convert % to ppm
                        else:
                            conc = float(value)
                    else:
                        conc = float(value)
                    
                    available_standards.append(material_name)
                    concentrations.append(conc)
                except (ValueError, TypeError):
                    continue
        
        if len(available_standards) < 2:
            QMessageBox.warning(self, "Insufficient Standards", 
                              f"At least 2 reference materials with certified values are needed for {current_element}.\n"
                              f"Only {len(available_standards)} standards available.")
            return
        
        # Show dialog for file selection and automatic analysis
        self.show_auto_calibration_dialog(current_element, available_standards, concentrations)
    
    def show_auto_calibration_dialog(self, element, standards, concentrations):
        """Show dialog for automatic calibration from spectra files"""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Auto-Calibrate {element} from Spectra Files")
        dialog.setGeometry(300, 200, 900, 650)
        
        layout = QVBoxLayout(dialog)
        
        # Instructions
        info_label = QLabel(f"Automatic {element} Calibration from XRF Spectra Files\n\n"
                           f"This tool will:\n"
                           f"1. Load XRF spectra files for each reference material\n"
                           f"2. Automatically fit peaks and calculate integrated intensities\n"
                           f"3. Create a calibration curve using the certified concentrations\n\n"
                           f"Check the standards you want to use and select their spectra files:")
        info_label.setWordWrap(True)
        info_label.setStyleSheet("QLabel { background-color: #E3F2FD; padding: 10px; border-radius: 5px; }")
        layout.addWidget(info_label)
        
        # File selection table with checkboxes
        file_table = QTableWidget(len(standards), 5)
        file_table.setHorizontalHeaderLabels(['Use', 'Standard', 'Certified Conc. (ppm)', 'Spectra File', 'Status'])
        
        file_buttons = []
        file_paths = {}
        status_labels = []
        standard_checkboxes = []
        
        for i, (standard, conc) in enumerate(zip(standards, concentrations)):
            # Checkbox for selection
            checkbox = QCheckBox()
            checkbox.setChecked(True)  # Default to checked
            checkbox_widget = QWidget()
            checkbox_layout = QHBoxLayout(checkbox_widget)
            checkbox_layout.addWidget(checkbox)
            checkbox_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            checkbox_layout.setContentsMargins(0, 0, 0, 0)
            file_table.setCellWidget(i, 0, checkbox_widget)
            standard_checkboxes.append(checkbox)
            
            # Standard name
            file_table.setItem(i, 1, QTableWidgetItem(standard))
            file_table.item(i, 1).setFlags(file_table.item(i, 1).flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            # Certified concentration
            file_table.setItem(i, 2, QTableWidgetItem(f"{conc:.2f}"))
            file_table.item(i, 2).setFlags(file_table.item(i, 2).flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            # File selection button
            file_btn = QPushButton(f"Select {standard} File...")
            file_btn.clicked.connect(lambda checked, std=standard: self.select_standard_file(std, file_paths, status_labels, standards))
            file_table.setCellWidget(i, 3, file_btn)
            file_buttons.append(file_btn)
            
            # Status
            status_item = QTableWidgetItem("No file selected")
            status_item.setFlags(status_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            file_table.setItem(i, 4, status_item)
            status_labels.append(status_item)
        
        file_table.resizeColumnsToContents()
        layout.addWidget(file_table)
        
        # Progress area
        progress_group = QGroupBox("Analysis Progress")
        progress_layout = QVBoxLayout(progress_group)
        
        self.calibration_progress = QTextEdit()
        self.calibration_progress.setMaximumHeight(150)
        self.calibration_progress.setReadOnly(True)
        progress_layout.addWidget(self.calibration_progress)
        
        layout.addWidget(progress_group)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        analyze_btn = QPushButton("🔬 Analyze Selected Spectra & Create Calibration")
        analyze_btn.clicked.connect(lambda: self.analyze_all_standards(element, standards, concentrations, file_paths, standard_checkboxes, dialog))
        analyze_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                padding: 8px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        
        button_layout.addWidget(analyze_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        # Store references for use in other methods
        dialog.file_paths = file_paths
        dialog.status_labels = status_labels
        dialog.standards = standards
        dialog.standard_checkboxes = standard_checkboxes
        
        dialog.exec()
    
    def select_standard_file(self, standard_name, file_paths, status_labels, standards):
        """Select a spectra file for a specific standard"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, f"Select {standard_name} Spectra File", "", 
            "All Files (*);;CSV Files (*.csv);;Text Files (*.txt);;Excel Files (*.xlsx)"
        )
        
        if file_path:
            file_paths[standard_name] = file_path
            # Update status in table
            standard_index = standards.index(standard_name)
            status_labels[standard_index].setText(f"✓ {os.path.basename(file_path)}")
            status_labels[standard_index].setBackground(QColor(144, 238, 144))  # Light green
    
    def analyze_all_standards(self, element, standards, concentrations, file_paths, standard_checkboxes, dialog):
        """Analyze all selected spectra files and create calibration"""
        self.calibration_progress.clear()
        self.calibration_progress.append(f"Starting automatic calibration for {element}...\n")
        
        # Filter to only checked standards
        selected_standards = []
        selected_concentrations = []
        for i, (standard, conc) in enumerate(zip(standards, concentrations)):
            if standard_checkboxes[i].isChecked():
                selected_standards.append(standard)
                selected_concentrations.append(conc)
        
        if len(selected_standards) < 2:
            QMessageBox.warning(dialog, "Insufficient Standards", 
                              "Please select at least 2 standards for calibration.")
            return
        
        # Check that all selected standards have files
        missing_files = [std for std in selected_standards if std not in file_paths]
        if missing_files:
            QMessageBox.warning(dialog, "Missing Files", 
                              f"Please select spectra files for: {', '.join(missing_files)}")
            return
        
        self.calibration_progress.append(f"Using {len(selected_standards)} selected standards\n")
        
        # Create a temporary peak fitter for this element
        temp_fitter = XRFPeakFitter(element=element)
        
        measured_intensities = []
        valid_concentrations = []
        valid_standards = []
        
        for i, (standard, conc) in enumerate(zip(selected_standards, selected_concentrations)):
            if standard not in file_paths:
                continue
                
            file_path = file_paths[standard]
            self.calibration_progress.append(f"Analyzing {standard} ({os.path.basename(file_path)})...")
            QApplication.processEvents()  # Update UI
            
            try:
                # Load and analyze the file
                data = self.read_xrf_file(file_path)
                if data is None:
                    self.calibration_progress.append(f"  ✗ Failed to load file")
                    continue
                
                x, y = data
                
                # Fit the peak for this element
                fit_params, fit_curve, r_squared, x_fit, integrated_intensity, concentration = temp_fitter.fit_peak(
                    x, y, 
                    peak_region=None,  # Use element defaults
                    background_subtract=True,
                    integration_region=None  # Use element defaults
                )
                
                measured_intensities.append(integrated_intensity)
                valid_concentrations.append(conc)
                valid_standards.append(standard)
                
                self.calibration_progress.append(f"  ✓ Integrated Intensity: {integrated_intensity:.2f}")
                self.calibration_progress.append(f"  ✓ R² = {r_squared:.4f}")
                
            except Exception as e:
                self.calibration_progress.append(f"  ✗ Error: {str(e)}")
                continue
        
        if len(measured_intensities) < 2:
            QMessageBox.warning(dialog, "Insufficient Data", 
                              "Need at least 2 successful analyses to create calibration.")
            return
        
        # Calculate calibration curve
        try:
            slope, intercept, r_value, p_value, std_err = stats.linregress(measured_intensities, valid_concentrations)
            
            self.calibration_progress.append(f"\n🎯 Calibration Results:")
            self.calibration_progress.append(f"Equation: Concentration = {slope:.4f} × Intensity + {intercept:.4f}")
            self.calibration_progress.append(f"R² = {r_value**2:.4f}")
            self.calibration_progress.append(f"Standard Error = {std_err:.4f}")
            self.calibration_progress.append(f"Standards used: {', '.join(valid_standards)}")
            
            # Show results and ask for confirmation
            result_msg = (f"Automatic Calibration Results for {element}:\n\n"
                         f"Equation: Concentration = {slope:.4f} × Intensity + {intercept:.4f}\n"
                         f"R² = {r_value**2:.4f}\n"
                         f"Standard Error = {std_err:.4f}\n"
                         f"Standards selected: {len(selected_standards)}\n"
                         f"Successfully analyzed: {len(valid_standards)}\n"
                         f"Standards used: {', '.join(valid_standards)}\n\n"
                         f"Apply this calibration?")
            
            reply = QMessageBox.question(dialog, "Calibration Results", result_msg,
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                # Update calibration in fitters
                self.peak_fitter.update_element_calibration(element, slope, intercept)
                self.fitter.update_element_calibration(element, slope, intercept)
                
                # Save calibration persistently with the standards that were actually used
                self.calibration_manager.update_calibration(
                    element, slope, intercept, r_value**2, valid_standards
                )
                
                # Update UI
                self.element_slope_edit.setText(f"{slope:.4f}")
                self.element_intercept_edit.setText(f"{intercept:.4f}")
                self.refresh_calibration_status()  # Update status display
                
                QMessageBox.information(dialog, "Success", 
                                      f"Automatic calibration for {element} has been created and applied!\n\n"
                                      f"The calibration is now active and will be used for all {element} analyses.\n"
                                      f"Calibration saved permanently.")
                dialog.accept()
            
        except Exception as e:
            QMessageBox.critical(dialog, "Calibration Error", f"Failed to calculate calibration: {str(e)}")

    def select_all_elements(self):
        """Select all elements for analysis"""
        for checkbox in self.element_checkboxes.values():
            checkbox.setChecked(True)
    
    def select_none_elements(self):
        """Deselect all elements"""
        for checkbox in self.element_checkboxes.values():
            checkbox.setChecked(False)
    
    def select_common_elements(self):
        """Select commonly analyzed elements"""
        common_elements = ['Pb', 'Zn', 'Cu', 'Cr']
        for element, checkbox in self.element_checkboxes.items():
            checkbox.setChecked(element in common_elements)
    
    def get_selected_elements(self):
        """Get list of selected elements for analysis"""
        selected = []
        for element, checkbox in self.element_checkboxes.items():
            if checkbox.isChecked():
                selected.append(element)
        return selected

    def setup_calibration_status_table(self):
        """Setup the calibration status table"""
        elements = list(ELEMENT_DEFINITIONS.keys())
        self.calibration_status_table.setRowCount(len(elements))
        self.calibration_status_table.setColumnCount(6)
        self.calibration_status_table.setHorizontalHeaderLabels([
            'Element', 'Status', 'Equation', 'R²', 'Created', 'Standards Used'
        ])
        
        for i, element in enumerate(elements):
            # Element name
            element_data = ELEMENT_DEFINITIONS[element]
            element_item = QTableWidgetItem(f"{element} ({element_data['name']})")
            element_item.setFlags(element_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.calibration_status_table.setItem(i, 0, element_item)
        
        self.refresh_calibration_status()
        self.calibration_status_table.resizeColumnsToContents()
    
    def refresh_calibration_status(self):
        """Refresh the calibration status display"""
        elements = list(ELEMENT_DEFINITIONS.keys())
        
        for i, element in enumerate(elements):
            calibration = self.calibration_manager.get_calibration(element)
            
            if calibration:
                # Has custom calibration
                status_item = QTableWidgetItem("✅ Calibrated")
                status_item.setBackground(QColor(144, 238, 144))  # Light green
                
                equation_item = QTableWidgetItem(calibration.get('equation', 'N/A'))
                r_squared_item = QTableWidgetItem(f"{calibration.get('r_squared', 'N/A'):.4f}" if calibration.get('r_squared') else 'N/A')
                
                created_date = calibration.get('created_date', 'Unknown')
                if created_date != 'Unknown':
                    try:
                        date_obj = datetime.fromisoformat(created_date)
                        created_item = QTableWidgetItem(date_obj.strftime("%Y-%m-%d %H:%M"))
                    except:
                        created_item = QTableWidgetItem(created_date)
                else:
                    created_item = QTableWidgetItem('Unknown')
                
                standards_used = calibration.get('standards_used', [])
                standards_item = QTableWidgetItem(f"{len(standards_used)} standards" if standards_used else 'N/A')
                
            else:
                # Using default calibration
                status_item = QTableWidgetItem("⚠️ Default")
                status_item.setBackground(QColor(255, 182, 193))  # Light red
                
                default_cal = ELEMENT_DEFINITIONS[element]['default_calibration']
                equation_item = QTableWidgetItem(f"Concentration = {default_cal['slope']:.4f} × Intensity + {default_cal['intercept']:.4f}")
                r_squared_item = QTableWidgetItem('N/A')
                created_item = QTableWidgetItem('Default')
                standards_item = QTableWidgetItem('None')
            
            # Set all items as non-editable
            for item in [status_item, equation_item, r_squared_item, created_item, standards_item]:
                item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            self.calibration_status_table.setItem(i, 1, status_item)
            self.calibration_status_table.setItem(i, 2, equation_item)
            self.calibration_status_table.setItem(i, 3, r_squared_item)
            self.calibration_status_table.setItem(i, 4, created_item)
            self.calibration_status_table.setItem(i, 5, standards_item)
        
        self.calibration_status_table.resizeColumnsToContents()
    
    def export_calibrations(self):
        """Export calibrations to a file"""
        filename, _ = QFileDialog.getSaveFileName(
            self, "Export Calibrations", "xrf_calibrations_export.json", 
            "JSON Files (*.json);;All Files (*)"
        )
        
        if filename:
            if self.calibration_manager.export_calibrations(filename):
                QMessageBox.information(self, "Export Successful", 
                                      f"Calibrations exported to:\n{filename}")
            else:
                QMessageBox.warning(self, "Export Failed", 
                                  "Failed to export calibrations.")
    
    def import_calibrations(self):
        """Import calibrations from a file"""
        filename, _ = QFileDialog.getOpenFileName(
            self, "Import Calibrations", "", 
            "JSON Files (*.json);;All Files (*)"
        )
        
        if filename:
            reply = QMessageBox.question(self, "Import Calibrations", 
                                       "This will overwrite existing calibrations with the same element names.\n\n"
                                       "Do you want to continue?",
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                if self.calibration_manager.import_calibrations(filename):
                    self.load_saved_calibrations()  # Reload into fitters
                    self.refresh_calibration_status()  # Refresh display
                    QMessageBox.information(self, "Import Successful", 
                                          f"Calibrations imported from:\n{filename}")
                else:
                    QMessageBox.warning(self, "Import Failed", 
                                      "Failed to import calibrations.")
    
    def reset_all_calibrations(self):
        """Reset all calibrations to defaults"""
        reply = QMessageBox.question(self, "Reset All Calibrations", 
                                   "This will delete ALL custom calibrations and revert to default values.\n\n"
                                   "This action cannot be undone. Are you sure?",
                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            # Clear all calibrations
            self.calibration_manager.calibrations = {}
            self.calibration_manager.save_calibrations()
            
            # Reset fitters to defaults
            for element in ELEMENT_DEFINITIONS.keys():
                default_cal = ELEMENT_DEFINITIONS[element]['default_calibration']
                self.fitter.update_element_calibration(element, default_cal['slope'], default_cal['intercept'])
                self.peak_fitter.update_element_calibration(element, default_cal['slope'], default_cal['intercept'])
            
            self.refresh_calibration_status()
            QMessageBox.information(self, "Reset Complete", 
                                  "All calibrations have been reset to default values.")

    def display_multi_element_sample_statistics(self, sample_groups):
        """Display statistics for multi-element sample groups"""
        if not sample_groups:
            return
        
        self.results_text.append(f"\n=== MULTI-ELEMENT SAMPLE STATISTICS ===")
        
        for group in sample_groups:
            sample_name = group['sample_name']
            element_groups = group['element_groups']
            selected_elements = group['selected_elements']
            
            self.results_text.append(f"\n{sample_name}:")
            self.results_text.append("-" * 40)
            
            for element in selected_elements:
                if element in element_groups:
                    element_group = element_groups[element]
                    self.results_text.append(f"  {element} ({ELEMENT_DEFINITIONS[element]['name']}):")
                    self.results_text.append(f"    Mean Concentration: {element_group.mean_concentration:.2f} ppm")
                    self.results_text.append(f"    Std Dev: {element_group.std_concentration:.2f} ppm")
                    self.results_text.append(f"    RSD: {element_group.rsd_concentration:.1f}%")
                    self.results_text.append(f"    Spectra: {element_group.n_spectra}")
                else:
                    self.results_text.append(f"  {element}: No successful analyses")
    
    def display_multi_element_summary(self, results, sample_groups):
        """Display summary for multi-element processing"""
        if not results:
            return
        
        selected_elements = results[0].get('selected_elements', [])
        total_files = len(results)
        
        self.results_text.append(f"\n=== MULTI-ELEMENT PROCESSING COMPLETE ===")
        self.results_text.append(f"Elements analyzed: {', '.join(selected_elements)}")
        self.results_text.append(f"Total files processed: {total_files}")
        self.results_text.append(f"Samples analyzed: {len(sample_groups)}")
        
        # Count successful analyses per element
        element_success_counts = {}
        for element in selected_elements:
            success_count = 0
            for result in results:
                element_result = result.get('element_results', {}).get(element, {})
                if 'error' not in element_result and element_result.get('concentration', 0) > 0:
                    success_count += 1
            element_success_counts[element] = success_count
        
        self.results_text.append(f"\nSuccess rates by element:")
        for element in selected_elements:
            success_rate = (element_success_counts[element] / total_files) * 100
            self.results_text.append(f"  {element}: {element_success_counts[element]}/{total_files} ({success_rate:.1f}%)")
        
        # Show concentration ranges
        self.results_text.append(f"\nConcentration ranges:")
        for element in selected_elements:
            concentrations = []
            for result in results:
                element_result = result.get('element_results', {}).get(element, {})
                if 'error' not in element_result:
                    conc = element_result.get('concentration', 0)
                    if conc > 0:
                        concentrations.append(conc)
            
            if concentrations:
                min_conc = min(concentrations)
                max_conc = max(concentrations)
                self.results_text.append(f"  {element}: {min_conc:.2f} - {max_conc:.2f} ppm")
            else:
                self.results_text.append(f"  {element}: No valid results")
        
        QMessageBox.information(self, "Multi-Element Processing Complete", 
                              f"Multi-element processing complete!\n\n"
                              f"Elements: {', '.join(selected_elements)}\n"
                              f"Files: {total_files}\n"
                              f"Samples: {len(sample_groups)}\n\n"
                              f"View results in the text panel and use the Spectrum Browser to examine individual fits.")

    def auto_calibrate_all_elements(self):
        """Automatically create calibrations for all elements with sufficient standards"""
        # Find elements with sufficient standards (at least 2)
        calibratable_elements = []
        
        for element in ELEMENT_DEFINITIONS.keys():
            available_standards = []
            concentrations = []
            
            for material_name, material_data in REFERENCE_MATERIALS.items():
                value = material_data.get(element)
                if value is not None and value != "N/A":
                    try:
                        if isinstance(value, str):
                            if '%' in value:
                                if '<' in value:
                                    continue  # Skip below detection limit values
                                conc = float(value.replace('%', '')) * 10000  # Convert % to ppm
                            else:
                                conc = float(value)
                        else:
                            conc = float(value)
                        
                        available_standards.append(material_name)
                        concentrations.append(conc)
                    except (ValueError, TypeError):
                        continue
            
            if len(available_standards) >= 2:
                calibratable_elements.append((element, available_standards, concentrations))
        
        if not calibratable_elements:
            QMessageBox.warning(self, "No Elements Available", 
                              "No elements have sufficient reference materials (≥2) for calibration.")
            return
        
        # Check for peak interferences
        calibratable_symbols = [elem[0] for elem in calibratable_elements]
        interference_warnings = []
        
        for element in calibratable_symbols:
            if element in PEAK_INTERFERENCES:
                interfering_elements = PEAK_INTERFERENCES[element]
                for interferer in interfering_elements:
                    if interferer in calibratable_symbols:
                        key = (element, interferer)
                        if key in INTERFERENCE_NOTES and key not in [(w[0], w[1]) for w in interference_warnings]:
                            interference_warnings.append((element, interferer, INTERFERENCE_NOTES[key]))
        
        # Show interference warning if needed
        if interference_warnings:
            # Check if Pb-As is among the interferences
            has_pb_as = any((elem1 == 'Pb' and elem2 == 'As') or (elem1 == 'As' and elem2 == 'Pb') 
                           for elem1, elem2, _ in interference_warnings)
            
            warning_msg = "⚠️ PEAK INTERFERENCE DETECTED ⚠️\n\n"
            warning_msg += "The following elements have overlapping emission lines:\n\n"
            for elem1, elem2, note in interference_warnings:
                warning_msg += f"• {elem1} ↔ {elem2}\n  {note}\n\n"
            
            if has_pb_as:
                warning_msg += "✨ SOLUTION AVAILABLE FOR Pb-As:\n"
                warning_msg += "Use the '⚛️ Pb-As Deconvolution' button in the Quant tab\n"
                warning_msg += "for accurate simultaneous quantification of both elements.\n\n"
            
            warning_msg += "Recommendations:\n"
            warning_msg += "1. Calibrate elements separately if possible\n"
            warning_msg += "2. Use standards with only ONE of these elements present\n"
            warning_msg += "3. Consider using alternative peaks (L-beta, K-beta)\n"
            if has_pb_as:
                warning_msg += "4. For Pb+As samples: Use Pb-As Deconvolution (recommended)\n"
            else:
                warning_msg += "4. Apply interference correction (advanced)\n"
            warning_msg += "\nContinue with calibration anyway?"
            
            reply = QMessageBox.question(self, "Peak Interference Warning", warning_msg,
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.No:
                return
        
        # Show dialog for multi-element calibration
        self.show_multi_element_calibration_dialog(calibratable_elements)
    
    def show_multi_element_calibration_dialog(self, calibratable_elements):
        """Show dialog for calibrating multiple elements at once"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Auto-Calibrate All Elements")
        dialog.setGeometry(200, 100, 1000, 700)
        
        layout = QVBoxLayout(dialog)
        
        # Instructions
        info_label = QLabel(f"Multi-Element Automatic Calibration\n\n"
                           f"This tool will calibrate {len(calibratable_elements)} elements simultaneously:\n"
                           f"{', '.join([elem[0] for elem in calibratable_elements])}\n\n"
                           f"For each reference material, select the XRF spectra file.\n"
                           f"The system will analyze all elements in each file automatically.")
        info_label.setWordWrap(True)
        info_label.setStyleSheet("QLabel { background-color: #E8F5E8; padding: 15px; border-radius: 8px; font-size: 12px; }")
        layout.addWidget(info_label)
        
        # Get all unique reference materials needed
        all_materials = set()
        for element, standards, concentrations in calibratable_elements:
            all_materials.update(standards)
        
        all_materials = sorted(list(all_materials))
        
        # File selection table with checkboxes for selection
        file_table = QTableWidget(len(all_materials), 4)
        file_table.setHorizontalHeaderLabels(['Use', 'Reference Material', 'Spectra Folder/Files', 'Status'])
        
        file_paths = {}  # Now stores list of files per material
        status_labels = []
        use_checkboxes = {}
        
        for i, material in enumerate(all_materials):
            # Checkbox to enable/disable this standard
            use_checkbox = QCheckBox()
            use_checkbox.setChecked(True)  # Default: use all standards
            use_checkboxes[material] = use_checkbox
            checkbox_widget = QWidget()
            checkbox_layout = QHBoxLayout(checkbox_widget)
            checkbox_layout.addWidget(use_checkbox)
            checkbox_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            checkbox_layout.setContentsMargins(0, 0, 0, 0)
            file_table.setCellWidget(i, 0, checkbox_widget)
            
            # Material name
            file_table.setItem(i, 1, QTableWidgetItem(material))
            file_table.item(i, 1).setFlags(file_table.item(i, 1).flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            # File/Folder selection buttons
            btn_widget = QWidget()
            btn_layout = QHBoxLayout(btn_widget)
            btn_layout.setContentsMargins(2, 2, 2, 2)
            
            folder_btn = QPushButton(f"📁 Folder")
            folder_btn.setToolTip(f"Select folder containing all {material} spectra")
            folder_btn.clicked.connect(lambda checked, mat=material: self.select_material_folder_for_multi(mat, file_paths, status_labels, all_materials))
            
            files_btn = QPushButton(f"📄 Files")
            files_btn.setToolTip(f"Select individual {material} spectrum files")
            files_btn.clicked.connect(lambda checked, mat=material: self.select_material_files_for_multi(mat, file_paths, status_labels, all_materials))
            
            btn_layout.addWidget(folder_btn)
            btn_layout.addWidget(files_btn)
            file_table.setCellWidget(i, 2, btn_widget)
            
            # Status
            status_item = QTableWidgetItem("No files selected")
            status_item.setFlags(status_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            file_table.setItem(i, 3, status_item)
            status_labels.append(status_item)
        
        file_table.resizeColumnsToContents()
        file_table.setColumnWidth(0, 50)  # Use column
        layout.addWidget(file_table)
        
        # Add Custom Standard button
        custom_std_layout = QHBoxLayout()
        add_custom_btn = QPushButton("➕ Add Custom Standard")
        add_custom_btn.clicked.connect(lambda: self.add_custom_standard_to_calibration(file_table, file_paths, status_labels, use_checkboxes, all_materials))
        add_custom_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                font-weight: bold;
                padding: 8px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #1976D2;
            }
        """)
        custom_std_layout.addWidget(add_custom_btn)
        custom_std_layout.addStretch()
        layout.addLayout(custom_std_layout)
        
        # Element summary table
        summary_group = QGroupBox("Elements to Calibrate")
        summary_layout = QVBoxLayout(summary_group)
        
        element_table = QTableWidget(len(calibratable_elements), 3)
        element_table.setHorizontalHeaderLabels(['Element', 'Standards Available', 'Concentration Range'])
        
        for i, (element, standards, concentrations) in enumerate(calibratable_elements):
            # Element
            element_item = QTableWidgetItem(f"{element} ({ELEMENT_DEFINITIONS[element]['name']})")
            element_item.setFlags(element_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            element_table.setItem(i, 0, element_item)
            
            # Standards count
            standards_item = QTableWidgetItem(f"{len(standards)} standards")
            standards_item.setFlags(standards_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            element_table.setItem(i, 1, standards_item)
            
            # Concentration range
            min_conc = min(concentrations)
            max_conc = max(concentrations)
            range_item = QTableWidgetItem(f"{min_conc:.1f} - {max_conc:.1f} ppm")
            range_item.setFlags(range_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            element_table.setItem(i, 2, range_item)
        
        element_table.resizeColumnsToContents()
        summary_layout.addWidget(element_table)
        layout.addWidget(summary_group)
        
        # Create tabbed view for progress and visualizations
        tabs = QTabWidget()
        
        # Progress tab
        progress_tab = QWidget()
        progress_layout = QVBoxLayout(progress_tab)
        
        self.multi_calibration_progress = QTextEdit()
        self.multi_calibration_progress.setReadOnly(True)
        progress_layout.addWidget(self.multi_calibration_progress)
        
        tabs.addTab(progress_tab, "📊 Progress")
        
        # Peak Fits tab
        fits_tab = QWidget()
        fits_layout = QVBoxLayout(fits_tab)
        
        fits_info = QLabel("Peak fitting visualizations will appear here during calibration")
        fits_info.setStyleSheet("color: gray; font-style: italic;")
        fits_layout.addWidget(fits_info)
        
        self.multi_cal_fits_canvas = PlotCanvas()
        fits_layout.addWidget(self.multi_cal_fits_canvas)
        
        tabs.addTab(fits_tab, "🔬 Peak Fits")
        
        # Calibration Curves tab
        curves_tab = QWidget()
        curves_layout = QVBoxLayout(curves_tab)
        
        curves_info = QLabel("Calibration curves will appear here after analysis")
        curves_info.setStyleSheet("color: gray; font-style: italic;")
        curves_layout.addWidget(curves_info)
        
        self.multi_cal_curves_canvas = PlotCanvas()
        curves_layout.addWidget(self.multi_cal_curves_canvas)
        
        tabs.addTab(curves_tab, "📈 Calibration Curves")
        
        layout.addWidget(tabs)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        analyze_all_btn = QPushButton("🚀 Analyze Selected Standards & Create Calibrations")
        analyze_all_btn.clicked.connect(lambda: self.analyze_all_elements_simultaneously(calibratable_elements, all_materials, file_paths, use_checkboxes, dialog))
        analyze_all_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                font-weight: bold;
                padding: 10px;
                border-radius: 5px;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        
        button_layout.addWidget(analyze_all_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        # Store references
        dialog.file_paths = file_paths
        dialog.status_labels = status_labels
        dialog.all_materials = all_materials
        dialog.use_checkboxes = use_checkboxes
        dialog.file_table = file_table
        dialog.custom_standards = {}  # Store custom standard data
        
        dialog.exec()
    
    def add_custom_standard_to_calibration(self, file_table, file_paths, status_labels, use_checkboxes, all_materials):
        """Add a custom standard to the calibration workflow"""
        # Dialog to get custom standard information
        custom_dialog = QDialog(self)
        custom_dialog.setWindowTitle("Add Custom Standard")
        custom_dialog.setGeometry(300, 200, 600, 500)
        
        layout = QVBoxLayout(custom_dialog)
        
        # Standard name
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel("Standard Name:"))
        name_edit = QLineEdit()
        name_edit.setPlaceholderText("e.g., My Custom Standard 1")
        name_layout.addWidget(name_edit)
        layout.addLayout(name_layout)
        
        # Element concentrations table
        conc_group = QGroupBox("Element Concentrations (ppm)")
        conc_layout = QVBoxLayout(conc_group)
        
        conc_table = QTableWidget(len(ELEMENT_DEFINITIONS), 3)
        conc_table.setHorizontalHeaderLabels(['Include', 'Element', 'Concentration (ppm)'])
        
        element_checkboxes = {}
        concentration_edits = {}
        
        for i, (symbol, data) in enumerate(ELEMENT_DEFINITIONS.items()):
            # Checkbox to include this element
            include_cb = QCheckBox()
            include_cb.setChecked(False)
            element_checkboxes[symbol] = include_cb
            cb_widget = QWidget()
            cb_layout = QHBoxLayout(cb_widget)
            cb_layout.addWidget(include_cb)
            cb_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            cb_layout.setContentsMargins(0, 0, 0, 0)
            conc_table.setCellWidget(i, 0, cb_widget)
            
            # Element name
            conc_table.setItem(i, 1, QTableWidgetItem(f"{symbol} ({data['name']})"))
            conc_table.item(i, 1).setFlags(conc_table.item(i, 1).flags() & ~Qt.ItemFlag.ItemIsEditable)
            
            # Concentration input
            conc_edit = QLineEdit()
            conc_edit.setPlaceholderText("Enter concentration")
            concentration_edits[symbol] = conc_edit
            conc_table.setCellWidget(i, 2, conc_edit)
        
        conc_table.resizeColumnsToContents()
        conc_layout.addWidget(conc_table)
        layout.addWidget(conc_group)
        
        # Buttons
        button_layout = QHBoxLayout()
        add_btn = QPushButton("Add Standard")
        cancel_btn = QPushButton("Cancel")
        
        add_btn.clicked.connect(custom_dialog.accept)
        cancel_btn.clicked.connect(custom_dialog.reject)
        
        button_layout.addWidget(add_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        if custom_dialog.exec() == QDialog.DialogCode.Accepted:
            std_name = name_edit.text().strip()
            if not std_name:
                QMessageBox.warning(self, "Invalid Name", "Please enter a standard name.")
                return
            
            # Collect element concentrations
            custom_concentrations = {}
            for symbol in ELEMENT_DEFINITIONS.keys():
                if element_checkboxes[symbol].isChecked():
                    try:
                        conc = float(concentration_edits[symbol].text())
                        custom_concentrations[symbol] = conc
                    except ValueError:
                        QMessageBox.warning(self, "Invalid Concentration", 
                                          f"Invalid concentration for {symbol}. Please enter a number.")
                        return
            
            if not custom_concentrations:
                QMessageBox.warning(self, "No Elements", 
                                  "Please select at least one element and enter its concentration.")
                return
            
            # Add row to table
            row_count = file_table.rowCount()
            file_table.insertRow(row_count)
            
            # Checkbox
            use_checkbox = QCheckBox()
            use_checkbox.setChecked(True)
            use_checkboxes[std_name] = use_checkbox
            checkbox_widget = QWidget()
            checkbox_layout = QHBoxLayout(checkbox_widget)
            checkbox_layout.addWidget(use_checkbox)
            checkbox_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            checkbox_layout.setContentsMargins(0, 0, 0, 0)
            file_table.setCellWidget(row_count, 0, checkbox_widget)
            
            # Standard name (marked as custom)
            name_item = QTableWidgetItem(f"{std_name} (Custom)")
            name_item.setFlags(name_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            name_item.setBackground(QColor(173, 216, 230))  # Light blue for custom
            file_table.setItem(row_count, 1, name_item)
            
            # File/Folder selection buttons (same as built-in standards)
            btn_widget = QWidget()
            btn_layout = QHBoxLayout(btn_widget)
            btn_layout.setContentsMargins(2, 2, 2, 2)
            
            folder_btn = QPushButton(f"📁 Folder")
            folder_btn.setToolTip(f"Select folder containing all {std_name} spectra")
            folder_btn.clicked.connect(lambda checked, mat=std_name: self.select_material_folder_for_multi(mat, file_paths, status_labels, all_materials))
            
            files_btn = QPushButton(f"📄 Files")
            files_btn.setToolTip(f"Select individual {std_name} spectrum files")
            files_btn.clicked.connect(lambda checked, mat=std_name: self.select_material_files_for_multi(mat, file_paths, status_labels, all_materials))
            
            btn_layout.addWidget(folder_btn)
            btn_layout.addWidget(files_btn)
            file_table.setCellWidget(row_count, 2, btn_widget)
            
            # Status
            status_item = QTableWidgetItem("No files selected")
            status_item.setFlags(status_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            file_table.setItem(row_count, 3, status_item)
            status_labels.append(status_item)
            
            # Store custom standard data
            all_materials.append(std_name)
            # Store concentrations for later use
            if not hasattr(self, 'custom_standards_data'):
                self.custom_standards_data = {}
            self.custom_standards_data[std_name] = custom_concentrations
            
            QMessageBox.information(self, "Success", 
                                  f"Custom standard '{std_name}' added with {len(custom_concentrations)} elements.")
    
    def select_material_folder_for_multi(self, material_name, file_paths, status_labels, all_materials):
        """Select a folder containing multiple spectra files for a specific material"""
        folder_path = QFileDialog.getExistingDirectory(
            self, f"Select Folder Containing {material_name} Spectra Files"
        )
        
        if folder_path:
            # Find all XRF files in the folder
            import glob
            patterns = ['*.txt', '*.csv', '*.xlsx', '*.dat', '*.emsa', '*.spc']
            files = []
            for pattern in patterns:
                files.extend(glob.glob(os.path.join(folder_path, pattern)))
            
            if not files:
                QMessageBox.warning(self, "No Files Found", 
                                  f"No XRF spectrum files found in the selected folder.")
                return
            
            # Store list of files
            file_paths[material_name] = files
            
            # Update status in table
            material_index = all_materials.index(material_name)
            status_labels[material_index].setText(f"✓ {len(files)} files from {os.path.basename(folder_path)}")
            status_labels[material_index].setBackground(QColor(144, 238, 144))  # Light green
    
    def select_material_files_for_multi(self, material_name, file_paths, status_labels, all_materials):
        """Select multiple individual spectra files for a specific material"""
        file_list, _ = QFileDialog.getOpenFileNames(
            self, f"Select {material_name} Spectra Files (multiple selection)", "", 
            "All Files (*);;CSV Files (*.csv);;Text Files (*.txt);;Excel Files (*.xlsx)"
        )
        
        if file_list:
            # Store list of files
            file_paths[material_name] = file_list
            
            # Update status in table
            material_index = all_materials.index(material_name)
            status_labels[material_index].setText(f"✓ {len(file_list)} files selected")
            status_labels[material_index].setBackground(QColor(144, 238, 144))  # Light green
    
    def analyze_all_elements_simultaneously(self, calibratable_elements, all_materials, file_paths, use_checkboxes, dialog):
        """Analyze all elements in all files simultaneously and create calibrations"""
        self.multi_calibration_progress.clear()
        self.multi_calibration_progress.append("🚀 Starting Multi-Element Calibration...\n")
        
        # Filter to only use selected standards
        selected_materials = [mat for mat in all_materials if use_checkboxes.get(mat, QCheckBox()).isChecked()]
        
        if not selected_materials:
            QMessageBox.warning(dialog, "No Standards Selected", 
                              "Please select at least one standard to use for calibration.")
            return
        
        # Check that all selected files have spectra
        missing_files = [mat for mat in selected_materials if mat not in file_paths]
        if missing_files:
            QMessageBox.warning(dialog, "Missing Files", 
                              f"Please select spectra files for: {', '.join(missing_files)}")
            return
        
        self.multi_calibration_progress.append(f"📊 Using {len(selected_materials)} selected standards\n")
        
        # Dictionary to store results for each element
        element_results = {}
        
        # Initialize results structure
        for element, standards, concentrations in calibratable_elements:
            element_results[element] = {
                'intensities': [],  # Averaged intensities per standard
                'concentrations': [],  # Certified concentrations
                'standards': [],  # Standard names
                'fit_quality': [],  # RSD values
                'raw_intensities': {},  # Individual measurements: {standard_name: [intensity1, intensity2, ...]}
                'raw_standards': [],  # List of standard names for each raw measurement
                'example_fit': None  # Store one example fit for visualization
            }
        
        # Analyze each standard's files for all elements (only selected materials)
        self.multi_calibration_progress.append(f"📁 Analyzing {len(selected_materials)} selected reference materials...\n")
        
        for material in selected_materials:
            if material not in file_paths:
                continue
            
            # Get list of files for this material (can be single file or multiple)
            material_files = file_paths[material]
            if not isinstance(material_files, list):
                material_files = [material_files]  # Convert single file to list
            
            self.multi_calibration_progress.append(f"🔬 Analyzing {material} ({len(material_files)} spectra)...")
            QApplication.processEvents()  # Update UI
            
            # Store intensities for each element from all files of this material
            material_element_intensities = {}
            
            # Analyze each file for this material
            for file_idx, file_path in enumerate(material_files):
                try:
                    # Load the file
                    data = self.read_xrf_file(file_path)
                    if data is None:
                        self.multi_calibration_progress.append(f"  ✗ File {file_idx+1}: Failed to load")
                        continue
                    
                    x, y = data
                    
                    # Analyze each element in this file
                    for element, standards, concentrations in calibratable_elements:
                        # Check if this is a custom standard
                        is_custom = hasattr(self, 'custom_standards_data') and material in self.custom_standards_data
                        
                        if is_custom:
                            # Custom standard - check if it has this element
                            if element not in self.custom_standards_data[material]:
                                continue
                            cert_conc = self.custom_standards_data[material][element]
                        else:
                            # Built-in reference material
                            if material not in standards:
                                continue  # This material doesn't have this element
                            
                            # Get the certified concentration for this element in this material
                            if material not in REFERENCE_MATERIALS:
                                continue
                            material_data = REFERENCE_MATERIALS[material]
                            cert_value = material_data.get(element)
                            
                            if cert_value is None or cert_value == "N/A":
                                continue
                            
                            # Parse concentration
                            try:
                                if isinstance(cert_value, str):
                                    if '%' in cert_value:
                                        if '<' in cert_value:
                                            continue
                                        cert_conc = float(cert_value.replace('%', '')) * 10000
                                    else:
                                        cert_conc = float(cert_value)
                                else:
                                    cert_conc = float(cert_value)
                            except (ValueError, TypeError):
                                continue
                        
                        # Create element-specific fitter
                        element_fitter = XRFPeakFitter(element=element)
                        
                        try:
                            # Fit the peak for this element
                            fit_params, fit_curve, r_squared, x_fit, integrated_intensity, concentration = element_fitter.fit_peak(
                                x, y, 
                                peak_region=None,  # Use element defaults
                                background_subtract=True,
                                integration_region=None  # Use element defaults
                            )
                            
                            # Store intensity for this file
                            if element not in material_element_intensities:
                                material_element_intensities[element] = {'intensities': [], 'cert_conc': cert_conc}
                                # Store first fit for visualization
                                material_element_intensities[element]['first_fit'] = {
                                    'x': x, 'y': y, 'x_fit': x_fit, 'fit_curve': fit_curve,
                                    'fit_params': fit_params, 'r_squared': r_squared,
                                    'material': material, 'file_idx': file_idx
                                }
                            material_element_intensities[element]['intensities'].append(integrated_intensity)
                            
                        except Exception as e:
                            continue
                
                except Exception as e:
                    self.multi_calibration_progress.append(f"  ✗ File {file_idx+1}: {str(e)}")
                    continue
            
            # Calculate averages and errors for this material
            for element in material_element_intensities:
                intensities = np.array(material_element_intensities[element]['intensities'])
                cert_conc = material_element_intensities[element]['cert_conc']
                
                if len(intensities) > 0:
                    mean_intensity = np.mean(intensities)
                    std_intensity = np.std(intensities, ddof=1) if len(intensities) > 1 else 0
                    rsd = (std_intensity / mean_intensity * 100) if mean_intensity > 0 else 0
                    
                    # Store averaged results
                    element_results[element]['intensities'].append(mean_intensity)
                    element_results[element]['concentrations'].append(cert_conc)
                    element_results[element]['standards'].append(material)
                    element_results[element]['fit_quality'].append(rsd)  # Store RSD as quality metric
                    
                    # Store first example fit if not already stored
                    if element_results[element]['example_fit'] is None and 'first_fit' in material_element_intensities[element]:
                        element_results[element]['example_fit'] = material_element_intensities[element]['first_fit']
                    
                    # Store raw individual measurements for plotting
                    if material not in element_results[element]['raw_intensities']:
                        element_results[element]['raw_intensities'][material] = []
                    element_results[element]['raw_intensities'][material].extend(intensities.tolist())
                    
                    # Store which standard each raw measurement belongs to
                    for intensity in intensities:
                        element_results[element]['raw_standards'].append(material)
                    
                    if len(intensities) > 1:
                        self.multi_calibration_progress.append(f"    {element}: {mean_intensity:.1f} ± {std_intensity:.1f} cps (RSD={rsd:.1f}%, n={len(intensities)})")
                    else:
                        self.multi_calibration_progress.append(f"    {element}: {mean_intensity:.1f} cps (n=1)")
            
            self.multi_calibration_progress.append("")  # Blank line
        
        # Create calibrations for each element
        self.multi_calibration_progress.append("📊 Creating calibration curves...\n")
        
        successful_calibrations = []
        failed_calibrations = []
        
        for element, standards, concentrations in calibratable_elements:
            results = element_results[element]
            
            if len(results['intensities']) < 2:
                self.multi_calibration_progress.append(f"❌ {element}: Insufficient data ({len(results['intensities'])} points)")
                failed_calibrations.append(element)
                continue
            
            try:
                # Calculate calibration
                slope, intercept, r_value, p_value, std_err = stats.linregress(results['intensities'], results['concentrations'])
                
                # Validate slope is positive (physically required for XRF)
                if slope <= 0:
                    self.multi_calibration_progress.append(f"❌ {element}: NEGATIVE SLOPE ({slope:.6f}) - Data quality issue!")
                    self.multi_calibration_progress.append(f"    Possible causes:")
                    self.multi_calibration_progress.append(f"    - Wrong peak region for this element")
                    self.multi_calibration_progress.append(f"    - Peak overlap/interference")
                    self.multi_calibration_progress.append(f"    - Incorrect standard concentrations")
                    self.multi_calibration_progress.append(f"    - Standards measured in wrong order")
                    failed_calibrations.append(element)
                    continue
                
                # Check for unreasonably large intercept (>20% of lowest concentration)
                min_conc = min(results['concentrations'])
                intercept_pct = abs(intercept / min_conc * 100) if min_conc > 0 else 0
                
                warning_msg = ""
                if intercept_pct > 20:
                    warning_msg = f" ⚠️ Large intercept ({intercept:.1f} ppm = {intercept_pct:.0f}% of lowest std)"
                
                # Update the calibrations in fitters
                self.peak_fitter.update_element_calibration(element, slope, intercept)
                self.fitter.update_element_calibration(element, slope, intercept)
                
                # Save calibration persistently with raw data for plotting
                self.calibration_manager.update_calibration(
                    element, slope, intercept, r_value**2, results['standards'],
                    raw_intensities=results.get('raw_intensities', {}),  # Individual measurements
                    raw_standards=results.get('raw_standards', [])  # Which standard each measurement belongs to
                )
                
                successful_calibrations.append(element)
                
                self.multi_calibration_progress.append(f"✅ {element}: y = {slope:.4f}x + {intercept:.4f}, R² = {r_value**2:.4f}{warning_msg}")
                self.multi_calibration_progress.append(f"    Standards used: {', '.join(results['standards'])}")
                
            except Exception as e:
                self.multi_calibration_progress.append(f"❌ {element}: Calibration failed - {str(e)}")
                failed_calibrations.append(element)
        
        # Summary
        self.multi_calibration_progress.append(f"\n🎯 Multi-Element Calibration Complete!")
        self.multi_calibration_progress.append(f"✅ Successful: {len(successful_calibrations)} elements")
        self.multi_calibration_progress.append(f"❌ Failed: {len(failed_calibrations)} elements")
        
        if successful_calibrations:
            self.multi_calibration_progress.append(f"\nSuccessful calibrations: {', '.join(successful_calibrations)}")
        
        if failed_calibrations:
            self.multi_calibration_progress.append(f"Failed calibrations: {', '.join(failed_calibrations)}")
        
        # Plot peak fits and calibration curves for successful calibrations
        if successful_calibrations:
            self.plot_multi_element_peak_fits(element_results, successful_calibrations)
            self.plot_multi_element_calibration_curves(element_results, successful_calibrations)
        
        # Show summary dialog
        if successful_calibrations:
            summary_msg = (f"Multi-Element Calibration Results:\n\n"
                          f"✅ Successfully calibrated: {len(successful_calibrations)} elements\n"
                          f"   {', '.join(successful_calibrations)}\n\n"
                          f"❌ Failed to calibrate: {len(failed_calibrations)} elements\n"
                          f"   {', '.join(failed_calibrations) if failed_calibrations else 'None'}\n\n"
                          f"Apply all successful calibrations?")
            
            reply = QMessageBox.question(dialog, "Multi-Element Calibration Results", summary_msg,
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                # Refresh calibration status display
                self.refresh_calibration_status()
                
                # Update calibration plot if it exists
                if hasattr(self, 'update_calibration_plot'):
                    self.update_calibration_plot()
                
                QMessageBox.information(dialog, "Success", 
                                      f"Multi-element calibration complete!\n\n"
                                      f"{len(successful_calibrations)} element calibrations have been applied and saved.\n"
                                      f"You can now analyze samples for all calibrated elements.\n\n"
                                      f"View calibration curves in the 'Calibration Plots' tab.")
                dialog.accept()
        else:
            QMessageBox.warning(dialog, "Calibration Failed", 
                              "No elements could be successfully calibrated.\n"
                              "Please check your spectra files and try again.")
    
    def plot_multi_element_peak_fits(self, element_results, successful_elements):
        """Plot example peak fits for all successfully calibrated elements"""
        try:
            # Determine grid layout based on number of elements
            n_elements = len(successful_elements)
            if n_elements == 1:
                rows, cols = 1, 1
            elif n_elements == 2:
                rows, cols = 1, 2
            elif n_elements <= 4:
                rows, cols = 2, 2
            elif n_elements <= 6:
                rows, cols = 2, 3
            elif n_elements <= 9:
                rows, cols = 3, 3
            else:
                rows, cols = 4, 3
            
            # Clear and setup figure
            self.multi_cal_fits_canvas.figure.clear()
            
            # Create subplots
            for idx, element in enumerate(successful_elements):
                if idx >= rows * cols:
                    break  # Don't plot more than grid allows
                
                example_fit = element_results[element].get('example_fit')
                if example_fit is None:
                    continue
                
                ax = self.multi_cal_fits_canvas.figure.add_subplot(rows, cols, idx + 1)
                
                # Extract fit data
                x = example_fit['x']
                y = example_fit['y']
                x_fit = example_fit['x_fit']
                fit_curve = example_fit['fit_curve']
                fit_params = example_fit['fit_params']
                r_squared = example_fit['r_squared']
                material = example_fit['material']
                
                # Calculate background
                background_y = fit_params['background_slope'] * x_fit + fit_params['background_intercept']
                
                # Plot spectrum
                ax.plot(x, y, 'b-', linewidth=0.8, alpha=0.5, label='Spectrum')
                
                # Plot fit
                ax.plot(x_fit, fit_curve, 'r-', linewidth=2, label='Fit')
                
                # Plot background
                ax.plot(x_fit, background_y, 'g--', linewidth=1, alpha=0.7, label='Background')
                
                # Labels and title
                ax.set_xlabel('Energy (keV)', fontsize=8)
                ax.set_ylabel('Counts', fontsize=8)
                ax.set_title(f'{element} - {material}\nR² = {r_squared:.4f}', fontsize=9, fontweight='bold')
                ax.grid(True, alpha=0.3)
                ax.legend(fontsize=7)
                ax.tick_params(labelsize=7)
            
            self.multi_cal_fits_canvas.figure.tight_layout()
            self.multi_cal_fits_canvas.draw()
            
        except Exception as e:
            print(f"Error plotting peak fits: {e}")
    
    def plot_multi_element_calibration_curves(self, element_results, successful_elements):
        """Plot calibration curves for all successfully calibrated elements"""
        try:
            # Determine grid layout based on number of elements
            n_elements = len(successful_elements)
            if n_elements == 1:
                rows, cols = 1, 1
            elif n_elements == 2:
                rows, cols = 1, 2
            elif n_elements <= 4:
                rows, cols = 2, 2
            elif n_elements <= 6:
                rows, cols = 2, 3
            elif n_elements <= 9:
                rows, cols = 3, 3
            else:
                rows, cols = 4, 3
            
            # Clear and setup figure
            self.multi_cal_curves_canvas.figure.clear()
            
            # Create subplots
            for idx, element in enumerate(successful_elements):
                if idx >= rows * cols:
                    break  # Don't plot more than grid allows
                
                ax = self.multi_cal_curves_canvas.figure.add_subplot(rows, cols, idx + 1)
                
                results = element_results[element]
                intensities = np.array(results['intensities'])
                concentrations = np.array(results['concentrations'])
                standards = results['standards']
                
                # Calculate calibration line
                slope, intercept, r_value, p_value, std_err = stats.linregress(intensities, concentrations)
                
                # Plot data points
                ax.scatter(intensities, concentrations, s=100, alpha=0.7, edgecolors='black', linewidth=1.5)
                
                # Plot calibration line
                x_line = np.linspace(0, max(intensities) * 1.1, 100)
                y_line = slope * x_line + intercept
                ax.plot(x_line, y_line, 'r-', linewidth=2, label=f'y = {slope:.3f}x + {intercept:.1f}')
                
                # Annotate points with standard names
                for i, (x, y, std) in enumerate(zip(intensities, concentrations, standards)):
                    ax.annotate(std, (x, y), xytext=(5, 5), textcoords='offset points', 
                               fontsize=8, alpha=0.7)
                
                # Labels and title
                ax.set_xlabel('Integrated Intensity (cps)', fontsize=9)
                ax.set_ylabel('Concentration (ppm)', fontsize=9)
                ax.set_title(f'{element} - R² = {r_value**2:.4f}', fontsize=10, fontweight='bold')
                ax.grid(True, alpha=0.3)
                ax.legend(fontsize=8)
                
                # Force origin to be visible
                ax.set_xlim(left=0)
                ax.set_ylim(bottom=min(0, min(concentrations) * 0.9))
            
            self.multi_cal_curves_canvas.figure.tight_layout()
            self.multi_cal_curves_canvas.draw()
            
        except Exception as e:
            print(f"Error plotting calibration curves: {e}")
    
    def setup_fp_tab(self, tab):
        """Setup the Fundamental Parameters (FP) method tab"""
        layout = QVBoxLayout(tab)
        
        # Info label
        info_text = "⚛️ <b>Fundamental Parameters (FP) Method</b><br>"
        if HAS_XRAYLIB:
            info_text += "Physics-based quantitative XRF analysis using Sherman equation and xraylib atomic data."
        else:
            info_text += "⚠️ <b>xraylib not installed!</b> Install with: <code>pip install xraylib</code>"
        
        info_label = QLabel(info_text)
        info_label.setWordWrap(True)
        info_label.setStyleSheet("""
            QLabel {
                background-color: #FFF3E0;
                padding: 10px;
                border-radius: 5px;
                font-size: 11px;
            }
        """)
        layout.addWidget(info_label)
        
        if not HAS_XRAYLIB:
            # Show installation instructions
            install_group = QGroupBox("Installation Required")
            install_layout = QVBoxLayout(install_group)
            
            install_text = QLabel(
                "The Fundamental Parameters method requires the <b>xraylib</b> library.<br><br>"
                "To install:<br>"
                "1. Open terminal/command prompt<br>"
                "2. Run: <code>pip install xraylib</code><br>"
                "3. Restart this application<br><br>"
                "xraylib provides X-ray cross-sections, fluorescence yields, and other atomic data."
            )
            install_text.setWordWrap(True)
            install_layout.addWidget(install_text)
            
            layout.addWidget(install_group)
            layout.addStretch()
            return
        
        # Instrument parameters
        instrument_group = QGroupBox("Instrument Parameters")
        instrument_layout = QGridLayout(instrument_group)
        
        instrument_layout.addWidget(QLabel("X-ray Tube Voltage (kV):"), 0, 0)
        self.fp_tube_voltage = QDoubleSpinBox()
        self.fp_tube_voltage.setRange(10, 100)
        self.fp_tube_voltage.setValue(50.0)
        self.fp_tube_voltage.setDecimals(1)
        instrument_layout.addWidget(self.fp_tube_voltage, 0, 1)
        
        instrument_layout.addWidget(QLabel("Tube Current (mA):"), 1, 0)
        self.fp_tube_current = QDoubleSpinBox()
        self.fp_tube_current.setRange(0.1, 10.0)
        self.fp_tube_current.setValue(1.0)
        self.fp_tube_current.setDecimals(2)
        instrument_layout.addWidget(self.fp_tube_current, 1, 1)
        
        instrument_layout.addWidget(QLabel("Tube Element:"), 2, 0)
        self.fp_tube_element = QComboBox()
        self.fp_tube_element.addItems(['Rh', 'W', 'Mo', 'Cr', 'Ag'])
        self.fp_tube_element.setCurrentText('Rh')
        instrument_layout.addWidget(self.fp_tube_element, 2, 1)
        
        instrument_layout.addWidget(QLabel("Detector Angle (°):"), 3, 0)
        self.fp_detector_angle = QDoubleSpinBox()
        self.fp_detector_angle.setRange(0, 90)
        self.fp_detector_angle.setValue(45.0)
        self.fp_detector_angle.setDecimals(1)
        instrument_layout.addWidget(self.fp_detector_angle, 3, 1)
        
        instrument_layout.addWidget(QLabel("Takeoff Angle (°):"), 4, 0)
        self.fp_takeoff_angle = QDoubleSpinBox()
        self.fp_takeoff_angle.setRange(0, 90)
        self.fp_takeoff_angle.setValue(45.0)
        self.fp_takeoff_angle.setDecimals(1)
        instrument_layout.addWidget(self.fp_takeoff_angle, 4, 1)
        
        layout.addWidget(instrument_group)
        
        # File selection
        file_group = QGroupBox("Spectrum File")
        file_layout = QHBoxLayout(file_group)
        
        self.fp_file_label = QLabel("No file selected")
        file_layout.addWidget(self.fp_file_label)
        
        fp_load_btn = QPushButton("📁 Load Spectrum")
        fp_load_btn.clicked.connect(self.fp_load_spectrum)
        file_layout.addWidget(fp_load_btn)
        
        layout.addWidget(file_group)
        
        # Elements to analyze
        elements_group = QGroupBox("Elements to Analyze")
        elements_layout = QVBoxLayout(elements_group)
        
        elements_info = QLabel("Select elements present in your sample (detected from Search tab or manual):")
        elements_layout.addWidget(elements_info)
        
        # Button to import from Search tab
        import_search_btn = QPushButton("⬅️ Import Elements from Search Tab")
        import_search_btn.clicked.connect(self.fp_import_from_search)
        import_search_btn.setToolTip("Import detected elements from Search tab")
        import_search_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; }")
        elements_layout.addWidget(import_search_btn)
        
        # List of elements to analyze
        self.fp_elements_list = QListWidget()
        self.fp_elements_list.setSelectionMode(QListWidget.SelectionMode.MultiSelection)
        self.fp_elements_list.setMaximumHeight(120)
        elements_layout.addWidget(self.fp_elements_list)
        
        elements_btn_layout = QHBoxLayout()
        add_element_btn = QPushButton("+ Add Element")
        add_element_btn.clicked.connect(self.fp_add_element_to_list)
        elements_btn_layout.addWidget(add_element_btn)
        
        remove_element_btn = QPushButton("- Remove Selected")
        remove_element_btn.clicked.connect(self.fp_remove_element_from_list)
        elements_btn_layout.addWidget(remove_element_btn)
        
        clear_elements_btn = QPushButton("Clear All")
        clear_elements_btn.clicked.connect(lambda: self.fp_elements_list.clear())
        elements_btn_layout.addWidget(clear_elements_btn)
        
        elements_layout.addLayout(elements_btn_layout)
        layout.addWidget(elements_group)
        
        # Known composition (optional - for spiked samples)
        known_comp_group = QGroupBox("Known Composition (Optional - for spiked samples)")
        known_comp_layout = QVBoxLayout(known_comp_group)
        
        known_info = QLabel("Only fill this if you spiked the sample with known amounts:")
        known_info.setStyleSheet("color: gray; font-style: italic;")
        known_comp_layout.addWidget(known_info)
        
        self.fp_known_comp_table = QTableWidget(0, 2)
        self.fp_known_comp_table.setHorizontalHeaderLabels(['Element', 'Known Mass Fraction'])
        self.fp_known_comp_table.setMaximumHeight(100)
        known_comp_layout.addWidget(self.fp_known_comp_table)
        
        known_btn_layout = QHBoxLayout()
        add_known_btn = QPushButton("+ Add Known")
        add_known_btn.clicked.connect(self.fp_add_known_composition)
        known_btn_layout.addWidget(add_known_btn)
        
        remove_known_btn = QPushButton("- Remove")
        remove_known_btn.clicked.connect(self.fp_remove_known_composition)
        known_btn_layout.addWidget(remove_known_btn)
        
        known_comp_layout.addLayout(known_btn_layout)
        layout.addWidget(known_comp_group)
        
        # Analysis buttons
        analysis_group = QGroupBox("Analysis")
        analysis_layout = QVBoxLayout(analysis_group)
        
        self.fp_fit_btn = QPushButton("⚙️ Fit Composition from Spectrum")
        self.fp_fit_btn.clicked.connect(self.fp_fit_composition)
        self.fp_fit_btn.setEnabled(False)
        self.fp_fit_btn.setStyleSheet("QPushButton { background-color: #2196F3; color: white; font-weight: bold; }")
        self.fp_fit_btn.setToolTip("Use FP method to determine elemental composition from measured intensities")
        analysis_layout.addWidget(self.fp_fit_btn)
        
        self.fp_calculate_btn = QPushButton("🔬 Calculate Theoretical Intensities")
        self.fp_calculate_btn.clicked.connect(self.fp_calculate_intensities)
        self.fp_calculate_btn.setEnabled(False)
        self.fp_calculate_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-weight: bold; }")
        self.fp_calculate_btn.setToolTip("Calculate what intensities SHOULD be for a given composition (requires known composition)")
        analysis_layout.addWidget(self.fp_calculate_btn)
        
        layout.addWidget(analysis_group)
        
        # Results display
        results_group = QGroupBox("FP Results")
        results_layout = QVBoxLayout(results_group)
        
        self.fp_results_text = QTextEdit()
        self.fp_results_text.setReadOnly(True)
        self.fp_results_text.setMaximumHeight(200)
        results_layout.addWidget(self.fp_results_text)
        
        layout.addWidget(results_group)
        layout.addStretch()
    
    def fp_load_spectrum(self):
        """Load spectrum for FP analysis"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Spectrum File", "", "Text Files (*.txt);;All Files (*)"
        )
        
        if file_path:
            try:
                data = self.read_xrf_file(file_path)
                if data is None:
                    raise Exception("Could not parse spectrum file")
                
                self.fp_spectrum_data = {'energy': data[0], 'counts': data[1], 'file_path': file_path}
                self.fp_file_label.setText(os.path.basename(file_path))
                self.fp_calculate_btn.setEnabled(True)
                self.fp_fit_btn.setEnabled(True)
                
                QMessageBox.information(self, "Success", "Spectrum loaded successfully!")
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to load spectrum:\n{str(e)}")
    
    def fp_import_from_search(self):
        """Import detected elements from Search tab"""
        if not hasattr(self, 'search_element_matches') or not self.search_element_matches:
            QMessageBox.warning(self, "No Elements", 
                              "No elements detected in Search tab.\n\n"
                              "Please:\n"
                              "1. Go to Search tab\n"
                              "2. Load a spectrum\n"
                              "3. Click 'Search for Elements'\n"
                              "4. Return here and click this button")
            return
        
        # Clear existing list
        self.fp_elements_list.clear()
        
        # Add detected elements
        for element in sorted(self.search_element_matches.keys()):
            self.fp_elements_list.addItem(element)
        
        n_elements = len(self.search_element_matches)
        QMessageBox.information(self, "Import Successful", 
                              f"Imported {n_elements} elements from Search tab.\n\n"
                              f"These elements will be used to fit composition from your spectrum.")
    
    def fp_add_element_to_list(self):
        """Add element to analysis list"""
        # Show dialog to select element
        element, ok = QInputDialog.getItem(self, "Add Element", "Select element:", 
                                          sorted(ELEMENT_DEFINITIONS.keys()), 0, False)
        if ok and element:
            # Check if already in list
            items = [self.fp_elements_list.item(i).text() for i in range(self.fp_elements_list.count())]
            if element not in items:
                self.fp_elements_list.addItem(element)
            else:
                QMessageBox.information(self, "Already Added", f"{element} is already in the list.")
    
    def fp_remove_element_from_list(self):
        """Remove selected elements from list"""
        for item in self.fp_elements_list.selectedItems():
            self.fp_elements_list.takeItem(self.fp_elements_list.row(item))
    
    def fp_add_known_composition(self):
        """Add known composition for spiked sample"""
        row = self.fp_known_comp_table.rowCount()
        self.fp_known_comp_table.insertRow(row)
        
        # Element combo box
        element_combo = QComboBox()
        for symbol in ELEMENT_DEFINITIONS.keys():
            element_combo.addItem(symbol)
        self.fp_known_comp_table.setCellWidget(row, 0, element_combo)
        
        # Mass fraction spin box
        fraction_spin = QDoubleSpinBox()
        fraction_spin.setRange(0, 1)
        fraction_spin.setValue(0.1)
        fraction_spin.setDecimals(4)
        fraction_spin.setSingleStep(0.01)
        self.fp_known_comp_table.setCellWidget(row, 1, fraction_spin)
    
    def fp_remove_known_composition(self):
        """Remove selected known composition"""
        current_row = self.fp_known_comp_table.currentRow()
        if current_row >= 0:
            self.fp_known_comp_table.removeRow(current_row)
    
    def fp_fit_composition(self):
        """Fit composition from measured spectrum using FP method - THE MAIN FUNCTION"""
        if not hasattr(self, 'fp_spectrum_data'):
            QMessageBox.warning(self, "No Data", "Please load a spectrum first.")
            return
        
        # Get elements to analyze
        elements = [self.fp_elements_list.item(i).text() for i in range(self.fp_elements_list.count())]
        
        if not elements:
            QMessageBox.warning(self, "No Elements", 
                              "Please add elements to analyze.\n\n"
                              "Use 'Import Elements from Search Tab' or '+ Add Element'.")
            return
        
        try:
            # Initialize FP calculator
            fp = XRFFundamentalParameters(
                tube_voltage=self.fp_tube_voltage.value(),
                tube_current=self.fp_tube_current.value(),
                tube_element=self.fp_tube_element.currentText(),
                detector_angle=self.fp_detector_angle.value(),
                takeoff_angle=self.fp_takeoff_angle.value()
            )
            
            # Extract measured intensities from spectrum using existing peak fitter
            measured_intensities = {}
            energy = self.fp_spectrum_data['energy']
            counts = self.fp_spectrum_data['counts']
            
            for element in elements:
                # Use existing peak fitter to get intensity
                element_fitter = XRFPeakFitter(element=element)
                try:
                    fit_params, fit_curve, r_squared, x_fit, integrated_intensity, concentration = element_fitter.fit_peak(
                        energy, counts, peak_region=None, background_subtract=True, integration_region=None
                    )
                    measured_intensities[element] = integrated_intensity
                except:
                    # If fit fails, skip this element
                    print(f"Warning: Could not fit {element}")
                    continue
            
            if not measured_intensities:
                QMessageBox.warning(self, "No Peaks Found", 
                                  "Could not extract peak intensities for any elements.\n\n"
                                  "Make sure the elements are actually present in the spectrum.")
                return
            
            # Fit composition using FP method
            fitted_composition = fp.fit_composition(measured_intensities, normalize=True)
            
            # Display results
            results_text = "⚙️ FP Method: Fitted Composition from Spectrum\n"
            results_text += "=" * 50 + "\n\n"
            results_text += f"Instrument: {self.fp_tube_element.currentText()} @ {self.fp_tube_voltage.value()} kV\n"
            results_text += f"File: {os.path.basename(self.fp_spectrum_data['file_path'])}\n\n"
            
            results_text += "Measured Intensities (cps):\n"
            results_text += "-" * 50 + "\n"
            for elem, intensity in measured_intensities.items():
                results_text += f"  {elem}: {intensity:.1f}\n"
            
            results_text += "\n✨ FITTED COMPOSITION (Mass Fractions):\n"
            results_text += "=" * 50 + "\n"
            for elem in sorted(fitted_composition.keys()):
                frac = fitted_composition[elem]
                ppm = frac * 1e6
                percent = frac * 100
                results_text += f"  {elem}: {frac:.6f} ({percent:.4f}% or {ppm:.1f} ppm)\n"
            
            results_text += "\nNote: This is a physics-based quantification using\n"
            results_text += "fundamental parameters (Sherman equation). Matrix\n"
            results_text += "effects are automatically accounted for.\n"
            
            self.fp_results_text.setText(results_text)
            
            # Plot composition as pie chart
            self.plot_canvas.figure.clear()
            ax = self.plot_canvas.figure.add_subplot(111)
            
            labels = [f"{elem}\n{fitted_composition[elem]*100:.2f}%" for elem in sorted(fitted_composition.keys())]
            values = [fitted_composition[elem] for elem in sorted(fitted_composition.keys())]
            colors = plt.cm.Set3(np.linspace(0, 1, len(values)))
            
            ax.pie(values, labels=labels, colors=colors, autopct='', startangle=90)
            ax.set_title('Fitted Elemental Composition (FP Method)', fontsize=12, fontweight='bold')
            
            self.plot_canvas.figure.tight_layout()
            self.plot_canvas.draw()
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Error", f"FP fitting failed:\n{str(e)}\n\nCheck console for details.")
    
    def fp_calculate_intensities(self):
        """Calculate theoretical intensities for KNOWN composition (for validation/spiked samples)"""
        if not hasattr(self, 'fp_spectrum_data'):
            QMessageBox.warning(self, "No Data", "Please load a spectrum first.")
            return
        
        # Get known composition
        known_comp = {}
        for row in range(self.fp_known_comp_table.rowCount()):
            element_widget = self.fp_known_comp_table.cellWidget(row, 0)
            fraction_widget = self.fp_known_comp_table.cellWidget(row, 1)
            
            if element_widget and fraction_widget:
                element = element_widget.currentText()
                fraction = fraction_widget.value()
                known_comp[element] = fraction
        
        if not known_comp:
            QMessageBox.warning(self, "No Known Composition", 
                              "This function calculates theoretical intensities for a KNOWN composition.\n\n"
                              "Please add elements and mass fractions in the 'Known Composition' table.\n\n"
                              "If you want to DETERMINE composition from your spectrum, use\n"
                              "'Fit Composition from Spectrum' button instead.")
            return
        
        try:
            # Initialize FP calculator
            fp = XRFFundamentalParameters(
                tube_voltage=self.fp_tube_voltage.value(),
                tube_current=self.fp_tube_current.value(),
                tube_element=self.fp_tube_element.currentText(),
                detector_angle=self.fp_detector_angle.value(),
                takeoff_angle=self.fp_takeoff_angle.value()
            )
            
            # Normalize composition
            total = sum(known_comp.values())
            known_comp = {k: v/total for k, v in known_comp.items()}
            
            # Calculate theoretical intensities
            results_text = "🔬 Theoretical FP Intensities for Known Composition\n"
            results_text += "=" * 50 + "\n\n"
            results_text += f"Tube: {self.fp_tube_element.currentText()} @ {self.fp_tube_voltage.value()} kV\n\n"
            results_text += "Input Composition (normalized):\n"
            
            for elem, frac in known_comp.items():
                results_text += f"  {elem}: {frac:.4f} ({frac*100:.2f}%)\n"
            
            results_text += "\nTheoretical Intensities:\n"
            results_text += "-" * 50 + "\n"
            
            for elem in known_comp.keys():
                intensity = fp.calculate_primary_intensity(elem, known_comp[elem], known_comp, 'KA1')
                results_text += f"{elem} K-alpha: {intensity:.6f} (arbitrary units)\n"
                
                # Try L-alpha for heavy elements
                if fp.line_energy(elem, 'LA1') > 0:
                    intensity_l = fp.calculate_primary_intensity(elem, known_comp[elem], known_comp, 'LA1')
                    results_text += f"{elem} L-alpha: {intensity_l:.6f} (arbitrary units)\n"
            
            results_text += "\nUse this to validate your instrument or compare\n"
            results_text += "with measured intensities from spiked samples.\n"
            
            self.fp_results_text.setText(results_text)
            
            # Plot tube spectrum
            energy_range = np.linspace(0, self.fp_tube_voltage.value(), 500)
            tube_spectrum = fp.get_tube_spectrum(energy_range)
            
            self.plot_canvas.figure.clear()
            ax = self.plot_canvas.figure.add_subplot(111)
            ax.plot(energy_range, tube_spectrum, 'b-', linewidth=1.5)
            ax.set_xlabel('Energy (keV)', fontsize=10)
            ax.set_ylabel('Relative Intensity', fontsize=10)
            ax.set_title(f'X-ray Tube Spectrum ({self.fp_tube_element.currentText()} @ {self.fp_tube_voltage.value()} kV)', 
                        fontsize=12, fontweight='bold')
            ax.grid(True, alpha=0.3)
            self.plot_canvas.figure.tight_layout()
            self.plot_canvas.draw()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"FP calculation failed:\n{str(e)}")


class ProtocolDialog(QDialog):
    """Dialog for displaying the XRF SOP with markdown formatting"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("XRF Standard Operating Procedure")
        self.setGeometry(200, 200, 900, 700)
        
        # Create layout
        layout = QVBoxLayout(self)
        
        # Create text browser for markdown display
        self.text_browser = QTextBrowser()
        self.text_browser.setOpenExternalLinks(True)
        self.text_browser.setFont(QFont("Arial", 10))
        layout.addWidget(self.text_browser)
        
        # Add buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Close)
        button_box.rejected.connect(self.accept)
        layout.addWidget(button_box)
        
        # Load and format protocol
        self.load_protocol()
    
    def load_protocol(self):
        """Load and format the XRF SOP text"""
        try:
            with open('xrf_sop_markdown.md', 'r', encoding='utf-8') as f:
                protocol_text = f.read()
            
            # Convert markdown to HTML for display
            html_content = self.markdown_to_html(protocol_text)
            self.text_browser.setHtml(html_content)
            
        except Exception as e:
            self.text_browser.setPlainText(f"Could not load xrf_sop_markdown.md: {e}")
    
    def markdown_to_html(self, markdown_text):
        """Convert markdown text to HTML for display"""
        import re
        
        # Process code blocks first (to avoid interference with other formatting)
        code_blocks = []
        def replace_code_block(match):
            code_blocks.append(match.group(0))
            return f"__CODE_BLOCK_{len(code_blocks)-1}__"
        
        # Find and replace code blocks
        html = re.sub(r'```.*?\n.*?```', replace_code_block, markdown_text, flags=re.DOTALL)
        
        # Process headers
        html = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        
        # Process tables (basic table support)
        def process_table(table_text):
            lines = table_text.strip().split('\n')
            if len(lines) < 3:
                return table_text
            
            html_table = '<table border="1" style="border-collapse: collapse; width: 100%; margin: 15px 0;">'
            
            for i, line in enumerate(lines):
                if line.startswith('|') and line.endswith('|'):
                    # Remove leading/trailing pipes and split
                    cells = [cell.strip() for cell in line[1:-1].split('|')]
                    
                    if i == 0:
                        # Header row
                        html_table += '<tr style="background-color: #f8f9fa; font-weight: bold;">'
                        for cell in cells:
                            html_table += f'<td style="padding: 8px; border: 1px solid #ddd;">{cell}</td>'
                        html_table += '</tr>'
                    elif i == 1 and all(cell.replace('-', '').replace(':', '').replace('|', '').strip() == '' for cell in cells):
                        # Separator row - skip
                        continue
                    else:
                        # Data row
                        html_table += '<tr>'
                        for cell in cells:
                            html_table += f'<td style="padding: 8px; border: 1px solid #ddd;">{cell}</td>'
                        html_table += '</tr>'
            
            html_table += '</table>'
            return html_table
        
        # Find and replace tables
        table_pattern = r'(\|.*\|\n\|[\s\-:|]*\|\n(\|.*\|\n)*)'
        tables = re.findall(table_pattern, html)
        for table in tables:
            html = html.replace(table[0], process_table(table[0]))
        
        # Process lists
        html = re.sub(r'^\s*[-*]\s+(.*?)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        html = re.sub(r'^\s*\d+\.\s+(.*?)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        
        # Process bold and italic (handle nested formatting)
        html = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'__(.*?)__', r'<strong>\1</strong>', html)
        html = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html)
        html = re.sub(r'_(.*?)_', r'<em>\1</em>', html)
        
        # Process inline code
        html = re.sub(r'`(.*?)`', r'<code>\1</code>', html)
        
        # Process horizontal rules
        html = re.sub(r'^---$', r'<hr style="border: none; border-top: 2px solid #3498db; margin: 20px 0;">', html, flags=re.MULTILINE)
        
        # Restore code blocks
        for i, code_block in enumerate(code_blocks):
            # Extract language and content
            lines = code_block.strip().split('\n')
            if lines[0].startswith('```'):
                lang = lines[0][3:].strip()
                content = '\n'.join(lines[1:-1])
                html = html.replace(f"__CODE_BLOCK_{i}__", 
                                  f'<pre><code class="{lang}">{content}</code></pre>')
        
        # Process paragraphs
        html = re.sub(r'\n\n+', '</p>\n<p>', html)
        html = re.sub(r'\n', '<br>', html)
        
        # Wrap in HTML structure with improved styling
        html = f"""
        <html>
        <head>
            <style>
                body {{ 
                    font-family: 'Segoe UI', Arial, sans-serif; 
                    margin: 20px; 
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{ 
                    color: #2c3e50; 
                    border-bottom: 3px solid #3498db; 
                    padding-bottom: 10px; 
                    margin-top: 30px;
                    font-size: 24px;
                }}
                h2 {{ 
                    color: #34495e; 
                    margin-top: 25px; 
                    margin-bottom: 15px;
                    font-size: 20px;
                    border-left: 4px solid #3498db;
                    padding-left: 10px;
                }}
                h3 {{ 
                    color: #7f8c8d; 
                    margin-top: 20px;
                    margin-bottom: 10px;
                    font-size: 16px;
                }}
                code {{ 
                    background-color: #f8f9fa; 
                    padding: 2px 6px; 
                    border-radius: 4px; 
                    font-family: 'Consolas', 'Courier New', monospace;
                    font-size: 0.9em;
                    color: #e74c3c;
                }}
                pre {{ 
                    background-color: #f8f9fa; 
                    padding: 15px; 
                    border-radius: 8px; 
                    border-left: 4px solid #3498db;
                    margin: 15px 0;
                    overflow-x: auto;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                    color: #333;
                }}
                li {{ 
                    margin: 8px 0; 
                    padding-left: 10px;
                }}
                strong {{ 
                    color: #e74c3c; 
                    font-weight: bold;
                }}
                em {{
                    color: #27ae60;
                    font-style: italic;
                }}
                p {{
                    margin: 10px 0;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 15px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #f8f9fa;
                    font-weight: bold;
                }}
                hr {{
                    border: none;
                    border-top: 2px solid #3498db;
                    margin: 20px 0;
                }}
            </style>
        </head>
        <body>
            <p>{html}</p>
        </body>
        </html>
        """
        
        return html

def detect_file_format(file_path):
    """
    Detect the format of an XRF data file by examining its content.
    
    Parameters:
    -----------
    file_path : str
        Path to the file to analyze
        
    Returns:
    --------
    format_type : str
        Detected format: 'emsa', 'nist_standard', 'csv', 'excel', 'space_separated', 'tab_separated', 'unknown'
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            # Read first 50 lines to analyze format
            lines = [f.readline().strip() for _ in range(50)]
        
        # Check for EMSA format
        emsa_indicators = ['#FORMAT', '#VERSION', '#SPECTRUM', '#NPOINTS', '#XUNITS']
        if any(indicator in lines[0] for indicator in emsa_indicators):
            return 'emsa'
        
        # Check for NIST standard format
        nist_indicators = ['B-Baseline', 'Spectral	Data	File', 'Data	Starts	Here']
        if any(indicator in ' '.join(lines[:20]) for indicator in nist_indicators):
            return 'nist_standard'
        
        # Find the first non-comment line to analyze format
        data_line = None
        for line in lines:
            if line and not line.startswith('#'):
                data_line = line
                break
        
        if data_line:
            # Check for CSV format (comma-separated)
            if ',' in data_line:
                return 'csv'
            
            # Check for tab-separated format
            if '\t' in data_line:
                return 'tab_separated'
            
            # Check for space-separated format
            if len(data_line.split()) >= 2:
                return 'space_separated'
        
        return 'unknown'
        
    except Exception as e:
        print(f"Error detecting file format for {file_path}: {e}")
        return 'unknown'

def parse_xrf_file_smart(file_path):
    """
    Smart parser for XRF data files that automatically detects and handles various formats.
    
    Parameters:
    -----------
    file_path : str
        Path to the XRF data file
        
    Returns:
    --------
    x : numpy.array
        Energy values (keV)
    y : numpy.array
        Intensity values (counts)
    format_detected : str
        The detected format type
    """
    try:
        # Detect file format
        format_type = detect_file_format(file_path)
        print(f"Detected format for {os.path.basename(file_path)}: {format_type}")
        
        if format_type == 'emsa':
            # Use existing EMSA parser
            metadata, spectrum_df = parse_emsa_file_pandas(file_path)
            if spectrum_df is not None and len(spectrum_df) > 0:
                x = spectrum_df['energy_kev'].values
                y = spectrum_df['counts'].values
                return x, y, format_type
        
        elif format_type == 'nist_standard':
            # Parse NIST standard format with header and 3 columns
            return parse_nist_standard_format(file_path, format_type)
        
        elif format_type == 'csv':
            # Parse CSV format
            return parse_csv_format(file_path, format_type)
        
        elif format_type == 'tab_separated':
            # Parse tab-separated format
            return parse_tab_separated_format(file_path, format_type)
        
        elif format_type == 'space_separated':
            # Parse space-separated format
            return parse_space_separated_format(file_path, format_type)
        
        else:
            # Try fallback parsing methods
            return parse_fallback_format(file_path, format_type)
            
    except Exception as e:
        print(f"Error in smart parsing of {file_path}: {e}")
        return None, None, 'error'

def parse_nist_standard_format(file_path, format_type):
    """Parse NIST standard files with header and 3 columns"""
    try:
        with open(file_path, 'r') as f:
            lines = f.readlines()
        
        data_lines = []
        in_data_section = False
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines
            if not line:
                continue
            
            # Check for data section start
            if "Data Starts Here" in line or "Data	Starts	Here" in line:
                in_data_section = True
                continue
            
            # Parse data lines (should have 3 columns: energy, intensity, baseline)
            if in_data_section:
                # Split by tabs or spaces
                parts = line.split()
                if len(parts) >= 2:
                    try:
                        energy = float(parts[0])
                        intensity = float(parts[1])
                        data_lines.append([energy, intensity])
                    except ValueError:
                        # Skip lines that can't be parsed as numbers
                        continue
        
        if not data_lines:
            raise ValueError("No valid data found in file")
        
        # Convert to numpy arrays
        data_array = np.array(data_lines)
        x = data_array[:, 0]
        y = data_array[:, 1]
        
        return x, y, format_type
        
    except Exception as e:
        print(f"Error parsing NIST standard format {file_path}: {e}")
        return None, None, format_type

def parse_csv_format(file_path, format_type):
    """Parse CSV format files with mixed header and data content"""
    try:
        # Read file line by line to handle mixed content
        data_lines = []
        in_data_section = False
        
        with open(file_path, 'r') as f:
            for line in f:
                line = line.strip()
                
                # Skip empty lines
                if not line:
                    continue
                
                # Check for data section start
                if "Data begins below" in line or "Data starts below" in line:
                    in_data_section = True
                    continue
                
                # Parse data lines (should have 2 columns: energy, intensity)
                if in_data_section and not line.startswith('#'):
                    # Split by comma and clean up
                    parts = [part.strip() for part in line.split(',')]
                    if len(parts) >= 2:
                        try:
                            energy = float(parts[0])
                            intensity = float(parts[1])
                            data_lines.append([energy, intensity])
                        except ValueError:
                            # Skip lines that can't be parsed as numbers
                            continue
        
        if not data_lines:
            # If no data section found, try to parse the whole file
            try:
                # First try with header
                df = pd.read_csv(file_path)
                
                # Check for common column names
                if 'Energy_keV' in df.columns and 'Intensity' in df.columns:
                    x = df['Energy_keV'].values
                    y = df['Intensity'].values
                    return x, y, format_type
                elif 'Energy' in df.columns and 'Intensity' in df.columns:
                    x = df['Energy'].values
                    y = df['Intensity'].values
                    return x, y, format_type
                elif len(df.columns) >= 2:
                    # Use first two columns
                    x = df.iloc[:, 0].values
                    y = df.iloc[:, 1].values
                    return x, y, format_type
                
                # If that fails, try without header
                df = pd.read_csv(file_path, header=None)
                
                # Find the first two numeric columns
                numeric_cols = []
                for col in df.columns:
                    if df[col].dtype in ['float64', 'int64'] or df[col].apply(lambda x: pd.to_numeric(x, errors='coerce')).notna().any():
                        numeric_cols.append(col)
                        if len(numeric_cols) == 2:
                            break
                
                if len(numeric_cols) >= 2:
                    x = df[numeric_cols[0]].values
                    y = df[numeric_cols[1]].values
                    return x, y, format_type
            except:
                pass
            
            raise ValueError("No valid data found in file")
        
        # Convert to numpy arrays
        data_array = np.array(data_lines)
        x = data_array[:, 0]
        y = data_array[:, 1]
        
        return x, y, format_type
        
    except Exception as e:
        print(f"Error parsing CSV format {file_path}: {e}")
        return None, None, format_type

def parse_tab_separated_format(file_path, format_type):
    """Parse tab-separated format files"""
    try:
        df = pd.read_csv(file_path, delimiter='\t', header=None)
        
        # Find the first two numeric columns
        numeric_cols = []
        for col in df.columns:
            if df[col].dtype in ['float64', 'int64'] or df[col].apply(lambda x: pd.to_numeric(x, errors='coerce')).notna().any():
                numeric_cols.append(col)
                if len(numeric_cols) == 2:
                    break
        
        if len(numeric_cols) >= 2:
            x = df[numeric_cols[0]].values
            y = df[numeric_cols[1]].values
            return x, y, format_type
        else:
            raise ValueError("Could not find two numeric columns")
            
    except Exception as e:
        print(f"Error parsing tab-separated format {file_path}: {e}")
        return None, None, format_type

def parse_space_separated_format(file_path, format_type):
    """Parse space-separated format files"""
    try:
        df = pd.read_csv(file_path, delimiter=r'\s+', header=None)
        
        # Find the first two numeric columns
        numeric_cols = []
        for col in df.columns:
            if df[col].dtype in ['float64', 'int64'] or df[col].apply(lambda x: pd.to_numeric(x, errors='coerce')).notna().any():
                numeric_cols.append(col)
                if len(numeric_cols) == 2:
                    break
        
        if len(numeric_cols) >= 2:
            x = df[numeric_cols[0]].values
            y = df[numeric_cols[1]].values
            return x, y, format_type
        else:
            raise ValueError("Could not find two numeric columns")
            
    except Exception as e:
        print(f"Error parsing space-separated format {file_path}: {e}")
        return None, None, format_type

def parse_fallback_format(file_path, format_type):
    """Fallback parsing for unknown formats"""
    try:
        # Try multiple parsing strategies
        strategies = [
            lambda: pd.read_csv(file_path, header=None),
            lambda: pd.read_csv(file_path, delimiter='\t', header=None),
            lambda: pd.read_csv(file_path, delimiter=r'\s+', header=None),
            lambda: pd.read_csv(file_path, delimiter=',', header=None)
        ]
        
        for i, strategy in enumerate(strategies):
            try:
                df = strategy()
                
                # Find the first two numeric columns
                numeric_cols = []
                for col in df.columns:
                    if df[col].dtype in ['float64', 'int64'] or df[col].apply(lambda x: pd.to_numeric(x, errors='coerce')).notna().any():
                        numeric_cols.append(col)
                        if len(numeric_cols) == 2:
                            break
                
                if len(numeric_cols) >= 2:
                    x = df[numeric_cols[0]].values
                    y = df[numeric_cols[1]].values
                    return x, y, f'fallback_{i}'
                    
            except Exception:
                continue
        
        raise ValueError("No parsing strategy worked")
        
    except Exception as e:
        print(f"Error in fallback parsing of {file_path}: {e}")
        return None, None, format_type

def parse_emsa_file_pandas(filename):
    """
    Parse EMSA/MAS spectral data file and return metadata dict and spectral DataFrame.
    
    Parameters:
    -----------
    filename : str or Path
        Path to the EMSA file
        
    Returns:
    --------
    metadata : dict
        Dictionary containing all header metadata
    spectrum_df : pandas.DataFrame
        DataFrame with columns ['energy_kev', 'counts']
    """
    metadata = {}
    data_lines = []
    
    with open(filename, 'r') as f:
        in_data = False
        for line in f:
            line = line.strip()
            
            # Check for data section start
            if line.startswith('#SPECTRUM'):
                in_data = True
                continue
                
            # Check for data section end
            if line.startswith('#ENDOFDATA'):
                break
                
            # Parse metadata (lines starting with #)
            if not in_data and line.startswith('#'):
                if ':' in line:
                    # Handle standard format: #KEY : VALUE
                    key, value = line[1:].split(':', 1)
                    metadata[key.strip()] = value.strip()
                else:
                    # Handle special format: ##KEY   : VALUE
                    parts = line[1:].split(None, 1)
                    if len(parts) == 2:
                        metadata[parts[0].strip()] = parts[1].strip()
                continue
                
            # Parse spectral data
            if in_data and line and not line.startswith('#'):
                try:
                    x, y = map(float, line.split(','))
                    data_lines.append([x, y])
                except ValueError:
                    continue
    
    # Create DataFrame
    spectrum_df = pd.DataFrame(data_lines, columns=['energy_kev', 'counts'])
    
    return metadata, spectrum_df

def load_multiple_emsa_files(file_pattern="*.txt"):
    """
    Load multiple EMSA files and return a dictionary of metadata and DataFrames.
    
    Parameters:
    -----------
    file_pattern : str
        Glob pattern to match EMSA files
        
    Returns:
    --------
    data_dict : dict
        Dictionary with filename as key, containing 'metadata' and 'spectrum' keys
    """
    data_dict = {}
    
    for file_path in Path('.').glob(file_pattern):
        if file_path.is_file():
            try:
                metadata, spectrum_df = parse_emsa_file_pandas(file_path)
                data_dict[file_path.name] = {
                    'metadata': metadata,
                    'spectrum': spectrum_df
                }
                print(f"Loaded: {file_path.name}")
            except Exception as e:
                print(f"Error loading {file_path.name}: {e}")
    
    return data_dict

def plot_spectrum(spectrum_df, metadata=None, title=None, ax=None):
    """
    Plot a single spectrum with optional metadata annotation.
    
    Parameters:
    -----------
    spectrum_df : pandas.DataFrame
        DataFrame with 'energy_kev' and 'counts' columns
    metadata : dict, optional
        Metadata dictionary for annotation
    title : str, optional
        Plot title
    ax : matplotlib.axes.Axes, optional
        Axes to plot on
    """
    # Apply compact theme if available
    try:
        apply_theme('compact')
    except:
        pass
    
    if ax is None:
        fig, ax = plt.subplots(figsize=(10, 6))
    
    ax.plot(spectrum_df['energy_kev'], spectrum_df['counts'], linewidth=1.5)
    ax.set_xlabel('Energy (keV)')
    ax.set_ylabel('Counts')
    
    if title:
        ax.set_title(title)
    elif metadata and 'TITLE' in metadata:
        ax.set_title(metadata['TITLE'])
    
    # Add metadata annotation if provided
    if metadata:
        info_text = f"Live time: {metadata.get('LIVETIME', 'N/A')} s\n"
        info_text += f"Beam kV: {metadata.get('BEAMKV', 'N/A')} kV\n"
        info_text += f"Date: {metadata.get('DATE', 'N/A')}"
        ax.text(0.02, 0.98, info_text, transform=ax.transAxes, 
                verticalalignment='top', bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
    
    ax.grid(True, alpha=0.3)
    return ax

def plot_multiple_spectra(data_dict, elements_to_highlight=None):
    """
    Plot multiple spectra on the same axes.
    
    Parameters:
    -----------
    data_dict : dict
        Dictionary of loaded spectra data
    elements_to_highlight : list, optional
        List of energy values to highlight (e.g., element peaks)
    """
    # Apply compact theme if available
    try:
        apply_theme('compact')
    except:
        pass
    
    fig, ax = plt.subplots(figsize=(12, 8))
    
    colors = plt.cm.tab10(np.linspace(0, 1, len(data_dict)))
    
    for (filename, data), color in zip(data_dict.items(), colors):
        spectrum_df = data['spectrum']
        metadata = data['metadata']
        
        label = metadata.get('TITLE', filename)
        ax.plot(spectrum_df['energy_kev'], spectrum_df['counts'], 
                label=label, linewidth=1.5, color=color)
    
    ax.set_xlabel('Energy (keV)')
    ax.set_ylabel('Counts')
    ax.set_title('XRF Spectra Comparison')
    ax.legend()
    ax.grid(True, alpha=0.3)
    
    # Highlight specific elements if provided
    if elements_to_highlight:
        for element_energy in elements_to_highlight:
            ax.axvline(x=element_energy, color='red', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    return fig, ax

def main():
    # Suppress Qt warnings
    import warnings
    warnings.filterwarnings("ignore", category=DeprecationWarning)
    
    # Suppress Qt internal drawing messages
    os.environ['QT_LOGGING_RULES'] = 'qt.qpa.drawing=false'
    
    app = QApplication(sys.argv)
    
    # macOS-specific fixes
    if sys.platform == 'darwin':
        # Set application attributes for macOS compatibility
        app.setAttribute(Qt.ApplicationAttribute.AA_DontShowIconsInMenus, True)
        app.setAttribute(Qt.ApplicationAttribute.AA_NativeWindows, True)
        # Fix for macOS dark mode and scaling issues
        # Note: AA_UseHighDpiPixmaps is deprecated in Qt6 but enabled by default
        try:
            app.setAttribute(Qt.ApplicationAttribute.AA_UseHighDpiPixmaps, True)
        except AttributeError:
            pass  # Not needed in Qt6, already enabled by default
    
    # Set application style
    app.setStyle('Fusion')
    
    # Create and show main window
    window = XRFPeakFittingGUI()
    window.show()
    
    sys.exit(app.exec())

if __name__ == '__main__':
    main()